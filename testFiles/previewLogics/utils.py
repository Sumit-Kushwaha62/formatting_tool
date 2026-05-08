import re
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ═══════════════════════════
# DRAWING / IMAGE DETECTION
# ═══════════════════════════

WP_NS  = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
MC_NS  = 'http://schemas.openxmlformats.org/markup-compatibility/2006'
W_NS   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def has_drawing(para):
    """Return True if paragraph contains any image/chart/drawing/object element."""
    p = para._p
    for tag in [qn('w:drawing'), qn('w:pict'), qn('w:object')]:
        if p.find('.//' + tag) is not None:
            return True
    if p.find(f'{{{MC_NS}}}AlternateContent') is not None:
        return True
    return False


# ═══════════════════════════
# PRE-CLEANING
# ═══════════════════════════

def clean_runs_in_para(para):
    for run in para.runs:
        cleaned = run.text.replace('\t', '').replace('\n', ' ')
        cleaned = re.sub(r' {2,}', ' ', cleaned)
        run.text = cleaned


def remove_proof_errors(para):
    """Remove <w:proofErr> elements that cause word splits during run iteration."""
    p = para._p
    for proof_err in p.findall(qn('w:proofErr')):
        p.remove(proof_err)


def run_has_drawing(run):
    """Return True if this run contains any drawing/image element."""
    r = run._r
    if r.find(qn('w:drawing')) is not None:
        return True
    if r.find(qn('w:pict')) is not None:
        return True
    if r.find(qn('w:object')) is not None:
        return True
    return False


def merge_runs_in_para(para):
    """Merge adjacent runs with identical formatting. NEVER merges runs that contain drawings."""
    if len(para.runs) <= 1:
        return

    i = 0
    while i < len(para.runs) - 1:
        r1 = para.runs[i]
        r2 = para.runs[i + 1]

        if run_has_drawing(r1) or run_has_drawing(r2):
            i += 1
            continue

        def fmt_sig(run):
            rPr = run._element.find(qn('w:rPr'))
            if rPr is None:
                return ('', None, None, None, None)
            bold   = rPr.find(qn('w:b'))
            italic = rPr.find(qn('w:i'))
            sz     = rPr.find(qn('w:sz'))
            color  = rPr.find(qn('w:color'))
            rFonts = rPr.find(qn('w:rFonts'))
            b_val     = None if bold   is None else bold.get(qn('w:val'), 'true')
            i_val     = None if italic is None else italic.get(qn('w:val'), 'true')
            sz_val    = None if sz     is None else sz.get(qn('w:val'))
            color_val = None if color  is None else color.get(qn('w:val'))
            font_val  = None if rFonts is None else rFonts.get(qn('w:ascii'))
            return (b_val, i_val, sz_val, color_val, font_val)

        if fmt_sig(r1) == fmt_sig(r2):
            r1.text = (r1.text or '') + (r2.text or '')
            r2._r.getparent().remove(r2._r)
        else:
            i += 1


def is_all_bold(para):
    runs = [r for r in para.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)


def is_bullet_para(para):
    """True if paragraph has a numbering/bullet list marker in XML."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return False
    return pPr.find(qn('w:numPr')) is not None


def apply_bold_before_colon(para, font_name, krutidev_mode):
    """If para has 'Label: rest' pattern, make text before ':' bold."""
    text = para.text
    colon_idx = text.find(':')
    if colon_idx <= 0 or colon_idx > 80:
        return

    label = text[:colon_idx + 1]
    rest  = text[colon_idx + 1:]

    first_run = para.runs[0] if para.runs else None
    size_pt = None
    if first_run:
        size_pt = first_run.font.size

    for run in list(para.runs):
        run._r.getparent().remove(run._r)

    r_bold = para.add_run(label)
    r_bold.bold = True
    if not krutidev_mode:
        set_font_properly(r_bold, font_name)
        if size_pt:
            r_bold.font.size = size_pt
        r_bold.font.color.rgb = RGBColor(0, 0, 0)

    if rest:
        r_rest = para.add_run(rest)
        r_rest.bold = False
        if not krutidev_mode:
            set_font_properly(r_rest, font_name)
            if size_pt:
                r_rest.font.size = size_pt
            r_rest.font.color.rgb = RGBColor(0, 0, 0)


def merge_split_paragraphs(doc):
    pass  # Disabled — causes duplicate rendering with mixed fonts


def clear_pPr_sz(para):
    """Remove font size override from paragraph-level rPr."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        return
    for tag in [qn('w:sz'), qn('w:szCs')]:
        el = rPr.find(tag)
        if el is not None:
            rPr.remove(el)


def set_pPr_sz(para, half_pts):
    """Set font size at paragraph-level rPr."""
    pPr = para._p.get_or_add_pPr()
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)
    for tag_name in ['w:sz', 'w:szCs']:
        el = rPr.find(qn(tag_name))
        if el is None:
            el = OxmlElement(tag_name)
            rPr.append(el)
        el.set(qn('w:val'), str(half_pts))


def clear_para_indent(para):
    """Remove all left/first-line indent from paragraph XML."""
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is not None:
        pPr.remove(ind)
    para.paragraph_format.left_indent        = None
    para.paragraph_format.first_line_indent  = None


def preprocess_document(doc):
    for para in doc.paragraphs:
        if has_drawing(para):
            continue
        remove_proof_errors(para)
        clean_runs_in_para(para)
        merge_runs_in_para(para)
    merge_split_paragraphs(doc)


# ═══════════════════════════
# HELPERS
# ═══════════════════════════

KRUTIDEV_FONTS = {'Kruti Dev 010', 'Kruti Dev 011', 'Krutidev010', 'Krutidev011',
                  'KrutiDev010', 'KrutiDev011', 'Kruti Dev010', 'Kruti Dev011'}

FONT_NAME_MAP = {
    'Krutidev010': 'Kruti Dev 010',
    'Krutidev011': 'Kruti Dev 011',
    'Mangal':      'Mangal',
    'Kokila':      'Kokila',
    'Utsaah':      'Utsaah',
    'Aparajita':   'Aparajita',
    'Nirmala UI':  'Nirmala UI',
}


# ═══════════════════════════
# HINDI CONVERSION
# ═══════════════════════════

def unicode_to_krutidev(text):
    """Convert Unicode Devanagari text to Kruti Dev 010 ASCII encoding."""
    if not text:
        return ""
    if not re.search(r'[\u0900-\u097F]', text):
        return text

    HALANT = '\u094D'

    CONJUNCTS = [
        ('\u0915\u094D\u0937', '{k'),
        ('\u0924\u094D\u0930', '\u00d8'),
        ('\u091C\u094D\u091E', 'K'),
        ('\u0936\u094D\u0930', 'J'),
        ('\u092A\u094D\u0930', 'iz'),
        ('\u0917\u094D\u0930', 'xz'),
        ('\u0915\u094D\u0930', 'dz'),
        ('\u092C\u094D\u0930', 'cz'),
        ('\u092E\u094D\u0930', 'ez'),
        ('\u091F\u094D\u0930', 'Vz'),
        ('\u0921\u094D\u0930', 'Mz'),
        ('\u0927\u094D\u0930', '/z'),
        ('\u0939\u094D\u0930', 'gz'),
        ('\u092D\u094D\u0930', 'Hkz'),
        ('\u0926\u094D\u092F', '|'),
        ('\u0926\u094D\u0927', '/~/k'),
        ('\u0926\u094D\u0935', 'n~o'),
        ('\u0924\u094D\u0924', '\u00d9k'),
        ('\u0924\u094D\u0915', 'Rd'),
        ('\u0924\u094D\u092A', 'Ri'),
        ('\u0924\u094D\u0938', 'Rl'),
        ('\u0938\u094D\u0924', 'Lr'),
        ('\u0938\u094D\u0925', 'LFk'),
        ('\u0938\u094D\u0928', 'Lu'),
        ('\u0928\u094D\u0924', 'Ur'),
        ('\u0928\u094D\u0926', 'Un'),
        ('\u0928\u094D\u0928', 'Uu'),
        ('\u0937\u094D\u091F', '"V'),
        ('\u0937\u094D\u0920', '"B'),
        ('\u0936\u094D\u0935', "'o"),
        ('\u0936\u094D\u0928', "'u"),
        ('\u0932\u094D\u0932', 'Yy'),
    ]
    for uni, kd in CONJUNCTS:
        text = text.replace(uni, kd)

    C = {
        'अ': 'v',  'आ': 'vk', 'इ': 'b',  'ई': 'bZ',
        'उ': 'm',  'ऊ': 'Å',  'ए': ',',  'ऐ': ',s',
        'ओ': 'vks','औ': 'vkS',
        'ा': 'k',  'ि': 'f',  'ी': 'h',  'ु': 'q',
        'ू': 'w',  'ृ': '`',  'े': 's',  'ै': 'S',
        'ो': 'ks', 'ौ': 'kS', 'ं': 'a',  'ः': '%',  'ँ': '\u00a1',
        'क': 'd',  'ख': '[k', 'ग': 'x',  'घ': '?k', 'ङ': 'M~',
        'च': 'p',  'छ': 'N',  'ज': 't',  'झ': '>k', 'ञ': '\u00a5',
        'ट': 'V',  'ठ': 'B',  'ड': 'M',  'ढ': '<',  'ण': '.k',
        'त': 'r',  'थ': 'Fk', 'द': 'n',  'ध': '/k', 'न': 'u',
        'प': 'i',  'फ': 'Q',  'ब': 'c',  'भ': 'Hk', 'म': 'e',
        'य': ';',  'र': 'j',  'ल': 'y',  'व': 'o',
        'श': "'k", 'ष': '"k', 'स': 'l',  'ह': 'g',
        'क़': 'd+','ख़': '[k+','ग़': 'x+','ज़': 't+',
        'ड़': 'M+','ढ़': '<+','फ़': 'Q+',
        '।': 'A',  '॥': 'AA',
        '०': ')',  '१': '!',  '२': '@',  '३': '#',  '४': '$',
        '५': '%',  '६': '^',  '७': '&',  '८': '*',  '९': '(',
    }

    HALF = {
        'क': 'D',  'ख': '[',  'ग': 'X',  'घ': '?',
        'च': 'P',  'ज': 'T',  'झ': '>',
        'ट': 'V~', 'ड': 'M~', 'ण': '.k~',
        'त': 'R',  'थ': 'F',  'द': 'n~', 'ध': '/',
        'न': 'U',  'प': 'I',  'ब': 'C',  'भ': 'H',
        'म': 'E',  'य': 'Y',  'र': 'z',
        'ल': 'y~', 'व': 'O',
        'श': "'",  'ष': '"',  'स': 'L',  'ह': 'g~',
        'ञ': '\u00a5~',
    }

    VOWELS = set('अआइईउऊएऐओऔ')
    MATRAS = set('ािीुूृेैोौंःँ')

    result = []
    chars  = list(text)
    n      = len(chars)
    i      = 0

    while i < n:
        c = chars[i]

        if ord(c) < 0x900 or ord(c) > 0x97F:
            result.append(c)
            i += 1
            continue

        if c == 'र' and i + 1 < n and chars[i + 1] == HALANT:
            if i + 2 < n and chars[i + 2] in C and chars[i + 2] not in VOWELS:
                i += 2
                syl = []
                nc  = chars[i]
                if i + 1 < n and chars[i + 1] == HALANT and nc in HALF:
                    syl.append(HALF[nc])
                    i += 2
                elif i + 1 < n and chars[i + 1] == 'ि':
                    syl.append('f')
                    syl.append(C.get(nc, nc))
                    i += 2
                else:
                    syl.append(C.get(nc, nc))
                    i += 1
                while i < n and chars[i] in MATRAS:
                    syl.append(C.get(chars[i], chars[i]))
                    i += 1
                syl.append('Z')
                result.extend(syl)
                continue

        if c in HALF and i + 1 < n and chars[i + 1] == HALANT:
            result.append(HALF[c])
            i += 2
            continue

        if c in C and c not in VOWELS and c not in MATRAS:
            if i + 1 < n and chars[i + 1] == 'ि':
                result.append('f')
                result.append(C.get(c, c))
                i += 2
                continue

        result.append(C.get(c, c))
        i += 1

    return ''.join(result)


def is_krutidev(font_name):
    return font_name and any(k.lower() in font_name.lower() for k in ['kruti', 'krutidev'])


def has_unicode_hindi(text):
    return bool(re.search(r'[\u0900-\u097F]', text))


def set_font_properly(run, font_name, size_pt=None):
    formal_name = FONT_NAME_MAP.get(font_name, font_name)
    run.font.name = formal_name

    r    = run._element
    rPr  = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()

    if is_krutidev(formal_name):
        rFonts.set(qn('w:hint'), 'default')
    else:
        rFonts.set(qn('w:hint'), 'complex')

    for attr in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
        rFonts.set(qn(f'w:{attr}'), formal_name)

    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        rPr.append(lang)

    if is_krutidev(formal_name):
        lang.set(qn('w:val'),   'en-US')
        lang.set(qn('w:ascii'), 'en-US')
        lang.set(qn('w:hAnsi'), 'en-US')
        lang.set(qn('w:bidi'),  'hi-IN')
    else:
        lang.set(qn('w:val'), 'hi-IN')
        lang.set(qn('w:cs'),  'hi-IN')

    if size_pt:
        run.font.size = Pt(size_pt)
        sz_cs = rPr.find(qn('w:szCs'))
        if sz_cs is None:
            sz_cs = OxmlElement('w:szCs')
            rPr.append(sz_cs)
        sz_cs.set(qn('w:val'), str(int(size_pt * 2)))


def set_para_font(para, font_name):
    """Set font at paragraph-level rPr."""
    formal_name = FONT_NAME_MAP.get(font_name, font_name)
    pPr  = para._p.get_or_add_pPr()
    rPr  = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)

    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)

    if is_krutidev(formal_name):
        rFonts.set(qn('w:hint'), 'default')

    for attr in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
        rFonts.set(qn(f'w:{attr}'), formal_name)

    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        rPr.append(lang)
    if is_krutidev(formal_name):
        lang.set(qn('w:val'), 'en-US')
    else:
        lang.set(qn('w:val'), 'hi-IN')


def add_run_with_font(para, text, font_name, size_pt, bold=False, color=None):
    run = para.add_run(text)
    run.bold = bold
    set_font_properly(run, font_name)
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def add_fld_char(run, fld_type):
    fc = OxmlElement('w:fldChar')
    fc.set(qn('w:fldCharType'), fld_type)
    run._r.append(fc)


def add_instr_text(run, instr):
    it = OxmlElement('w:instrText')
    it.text = instr
    run._r.append(it)


# ═══════════════════════════
# JUSTIFY HELPER
# ═══════════════════════════

def apply_clean_justify(para):
    """Justify only long paragraphs. Short lines stay LEFT."""
    text  = para.text.strip()
    words = text.split()
    if len(words) < 12 or len(text) < 100 or text.endswith(('?', ':', '!', ';')):
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pPr = para._p.get_or_add_pPr()
    for jc in pPr.findall(qn('w:jc')):
        pPr.remove(jc)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'both')
    pPr.append(jc)


def get_original_alignment(para):
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return None
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        return None
    val = jc.get(qn('w:val'))
    mapping = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right':  WD_ALIGN_PARAGRAPH.RIGHT,
        'both':   WD_ALIGN_PARAGRAPH.JUSTIFY,
        'left':   WD_ALIGN_PARAGRAPH.LEFT,
    }
    return mapping.get(val)


# ═══════════════════════════
# CORE FORMATTING ENGINE
# ═══════════════════════════

def apply_para_formatting(para, etype, font_name, font_size_pt, bold, color, align,
                           space_before_pt, space_after_pt,
                           first_indent=None, left_indent=None,
                           line_spacing=1.15):
    is_chapter_type = etype in ('chapter_heading', 'chapter_title', 'book_title')

    set_para_font(para, font_name)
    clear_pPr_sz(para)
    set_pPr_sz(para, int(font_size_pt * 2))

    para.paragraph_format.space_before = Pt(space_before_pt)
    para.paragraph_format.space_after  = Pt(space_after_pt)
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'),            str(int(space_before_pt * 20)))
    spacing.set(qn('w:after'),             str(int(space_after_pt  * 20)))
    spacing.set(qn('w:beforeAutospacing'), '0')
    spacing.set(qn('w:afterAutospacing'),  '0')

    try:
        ls = float(line_spacing)
    except Exception:
        ls = 1.15

    if ls == 1.0:
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif ls == 2.0:
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    else:
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para.paragraph_format.line_spacing = ls

    ind = pPr.find(qn('w:ind'))
    if ind is not None:
        pPr.remove(ind)

    if first_indent is not None or left_indent is not None:
        ind = OxmlElement('w:ind')
        if left_indent is not None:
            ind.set(qn('w:left'), str(int(left_indent * 1440)))
        if first_indent is not None:
            twips = int(first_indent * 1440) if isinstance(first_indent, float) else int(first_indent.inches * 1440)
            ind.set(qn('w:firstLine'), str(twips))
        pPr.append(ind)
    else:
        para.paragraph_format.first_line_indent = None
        para.paragraph_format.left_indent       = None

    para.alignment = align
    pPr2 = para._p.get_or_add_pPr()
    for jc_el in pPr2.findall(qn('w:jc')):
        pPr2.remove(jc_el)
    jc_new = OxmlElement('w:jc')
    align_val_map = {
        WD_ALIGN_PARAGRAPH.JUSTIFY: 'both',
        WD_ALIGN_PARAGRAPH.CENTER:  'center',
        WD_ALIGN_PARAGRAPH.LEFT:    'left',
        WD_ALIGN_PARAGRAPH.RIGHT:   'right',
    }
    jc_new.set(qn('w:val'), align_val_map.get(align, 'both'))
    pPr2.append(jc_new)

    for run in para.runs:
        if run_has_drawing(run):
            continue
        run.bold = bold
        set_font_properly(run, font_name, font_size_pt)
        run.font.color.rgb = color


# ═══════════════════════════
# TABLE HELPERS
# ═══════════════════════════

def format_table_cells(doc, font_name, base_size, line_spacing, black):
    """Apply font/size to all table cell content + 5pt spacing after each table."""
    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tbl_parent = tbl.getparent()
        tbl_siblings = list(tbl_parent)
        tbl_idx = tbl_siblings.index(tbl)
        if tbl_idx + 1 < len(tbl_siblings):
            next_el = tbl_siblings[tbl_idx + 1]
            if next_el.tag == qn('w:p'):
                nPr = next_el.get_or_add_pPr() if hasattr(next_el, 'get_or_add_pPr') else next_el.find(qn('w:pPr'))
                if nPr is None:
                    nPr = OxmlElement('w:pPr')
                    next_el.insert(0, nPr)
                sp = nPr.find(qn('w:spacing'))
                if sp is None:
                    sp = OxmlElement('w:spacing')
                    nPr.append(sp)
                sp.set(qn('w:before'), '100')
                sp.set(qn('w:beforeAutospacing'), '0')

        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if not para.text.strip() and not has_drawing(para):
                        continue
                    if has_drawing(para):
                        continue
                    set_para_font(para, font_name)
                    clear_pPr_sz(para)
                    set_pPr_sz(para, int(base_size * 2))
                    for run in para.runs:
                        if run_has_drawing(run):
                            continue
                        was_bold   = run.bold
                        was_italic = run.italic
                        set_font_properly(run, font_name, base_size)
                        run.bold   = was_bold
                        run.italic = was_italic
                        run.font.color.rgb = black


def center_all_tables(doc):
    for table in doc.tables:
        tbl   = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        jc = tblPr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            tblPr.append(jc)
        jc.set(qn('w:val'), 'center')


def set_para_text_formatted(para, new_text, font_size_pt, bold, color, font_name=None):
    """Set paragraph text while preserving run-level formatting."""
    p = para._p
    for r in p.findall(qn('w:r')):
        p.remove(r)
    run = para.add_run(new_text)
    run.bold = True if bold else False
    run.font.size = Pt(font_size_pt)
    run.font.color.rgb = color
    if font_name:
        set_font_properly(run, font_name, font_size_pt)


def strip_list_numbering(para):
    """Remove w:numPr from paragraph so Word doesn't render list number prefix."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return
    numPr = pPr.find(qn('w:numPr'))
    if numPr is not None:
        pPr.remove(numPr)


def apply_caps_upper(para, krutidev_mode=False):
    if krutidev_mode:
        return
    for run in para.runs:
        if run.text:
            run.text = run.text.upper()

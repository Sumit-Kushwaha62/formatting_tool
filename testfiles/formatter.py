import sys
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Mm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Arguments ──────────────────────────────────────────────
input_path   = sys.argv[1]
output_path  = sys.argv[2]
doc_type     = sys.argv[3] if len(sys.argv) > 3 else 'book'
options_file = sys.argv[4] if len(sys.argv) > 4 else None

options = {}
if options_file:
    try:
        with open(options_file, 'r', encoding='utf-8') as f:
            options = json.load(f)
    except Exception as e:
        print(f"Options load error: {e}")

print(f"DocType: {doc_type}")
print(f"Options: {options}")

# ── Defaults ────────────────────────────────────────────────
DEFAULT_FONTS = {
    'book':     'Garamond',
    'thesis':   'Times New Roman',
    'research': 'Arial',
    'letter':   'Calibri',
}

PAGE_SIZES = {
    'A4':     (210, 297),
    'A5':     (148, 210),
    'A3':     (297, 420),
    'Letter': (216, 279),
    'Legal':  (216, 356),
}

ALIGNMENT_MAP = {
    'left':    WD_ALIGN_PARAGRAPH.LEFT,
    'center':  WD_ALIGN_PARAGRAPH.CENTER,
    'right':   WD_ALIGN_PARAGRAPH.RIGHT,
    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# Word spacing map — twips value (1 pt = 20 twips)
WORD_SPACING_MAP = {
    'normal': 0,
    'wide':   40,
    'wider':  80,
    'widest': 120,
}

def get_font(opts, doc_type):
    return opts.get('font_style') or DEFAULT_FONTS.get(doc_type, 'Calibri')

def get_alignment(opts, default='justify'):
    key = opts.get('alignment', default).lower()
    return ALIGNMENT_MAP.get(key, WD_ALIGN_PARAGRAPH.JUSTIFY)

def get_word_spacing(opts):
    key = opts.get('word_spacing', 'normal').lower()
    return WORD_SPACING_MAP.get(key, 0)

def apply_page_size(doc, size_key='A4'):
    size = PAGE_SIZES.get(size_key, PAGE_SIZES['A4'])
    for section in doc.sections:
        section.page_width  = Mm(size[0])
        section.page_height = Mm(size[1])

# ── Helper: Page Border ─────────────────────────────────────
def add_page_border(doc, color='2E4057'):
    for section in doc.sections:
        sectPr = section._sectPr
        for existing in sectPr.findall(qn('w:pgBorders')):
            sectPr.remove(existing)
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for side in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '18')
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), color)
            pgBorders.append(border)
        sectPr.append(pgBorders)

# ── Helper: Header ──────────────────────────────────────────
def set_header(doc, text, font_name='Calibri', color_hex='6a5e4e'):
    if not text:
        return
    r, g, b = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
    for section in doc.sections:
        section.different_first_page_header_footer = False
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.clear()
        # Left side — text
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(r, g, b)
        run.font.name = font_name
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Bottom border on header
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'D1D5DB')
        pBdr.append(bottom)
        pPr.append(pBdr)

# ── Helper: Footer with optional page numbers ───────────────
def set_footer(doc, text, show_page_numbers=False, page_position='center', font_name='Calibri'):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()

        align_map = {
            'left':   WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right':  WD_ALIGN_PARAGRAPH.RIGHT,
        }
        p.alignment = align_map.get(page_position, WD_ALIGN_PARAGRAPH.CENTER)

        if text:
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x6a, 0x5e, 0x4e)
            run.font.name = font_name

        if show_page_numbers:
            if text:
                sep = p.add_run('  |  ')
                sep.font.size = Pt(9)
                sep.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)

            # Page number field
            run_pg = p.add_run()
            run_pg.font.size = Pt(9)
            run_pg.font.color.rgb = RGBColor(0x6a, 0x5e, 0x4e)
            run_pg.font.name = font_name
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.text = ' PAGE '
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run_pg._r.append(fldChar1)
            run_pg._r.append(instrText)
            run_pg._r.append(fldChar2)

            # " of X" part
            run_of = p.add_run(' of ')
            run_of.font.size = Pt(9)
            run_of.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)

            run_total = p.add_run()
            run_total.font.size = Pt(9)
            run_total.font.color.rgb = RGBColor(0x6a, 0x5e, 0x4e)
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'begin')
            instrText2 = OxmlElement('w:instrText')
            instrText2.text = ' NUMPAGES '
            fldChar4 = OxmlElement('w:fldChar')
            fldChar4.set(qn('w:fldCharType'), 'end')
            run_total._r.append(fldChar3)
            run_total._r.append(instrText2)
            run_total._r.append(fldChar4)

        # Top border on footer
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '4')
        top.set(qn('w:space'), '1')
        top.set(qn('w:color'), 'D1D5DB')
        pBdr.append(top)
        pPr.append(pBdr)

# ── Helper: Margins ─────────────────────────────────────────
def set_margins(doc, top=1.0, bottom=1.0, left=1.2, right=1.0):
    for section in doc.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)

# ── Helper: Apply word spacing to a run ─────────────────────
def apply_word_spacing(run, spacing_twips):
    if spacing_twips == 0:
        return
    rPr = run._r.get_or_add_rPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:val'), str(spacing_twips))
    rPr.append(spacing)

# ── Helper: Smart alignment ────────────────────────────────────────
import re

def get_heading_level(para):
    """Return heading level 1-9, or 0 if not a heading."""
    style = para.style
    if not style or not style.name:
        return 0
    name = style.name
    if name.startswith('Heading'):
        try:
            return int(name.split()[-1])
        except:
            return 1
    return 0

def smart_body_align(para, body_alignment):
    """Only apply smart logic when user chose JUSTIFY."""
    if body_alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
        return body_alignment
    text = para.text.strip()
    if not text:
        return body_alignment
    word_count = len(text.split())
    char_count = len(text)
    # Tab character — "Answers:   text" pattern, Word stretches badly
    if '\t' in text:
        return WD_ALIGN_PARAGRAPH.LEFT
    # List items — A. B. C. or 1. 2. 3.
    if re.match(r'^[A-Da-d]\.\s', text) or re.match(r'^\d+\.\s', text):
        return WD_ALIGN_PARAGRAPH.LEFT
    # Label lines ending with colon
    if text.endswith(':'):
        return WD_ALIGN_PARAGRAPH.LEFT
    # Short paragraph — likely single line, avoid justify gaps
    # 130 chars ~ 2 full lines at 12pt on A4 with margins
    if word_count < 20 or char_count < 130:
        return WD_ALIGN_PARAGRAPH.LEFT
    return WD_ALIGN_PARAGRAPH.JUSTIFY

# ── Helper: Format Paragraphs ────────────────────────────────────────
def format_paragraphs(doc, body_font='Calibri', body_size=11,
                       heading_size=14, heading_color='2E4057',
                       body_alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       word_spacing=0):
    r, g, b = tuple(int(heading_color[i:i+2], 16) for i in (0, 2, 4))

    # Skip tables — MCQ, data tables must keep original formatting
    for para in doc.paragraphs:
        if not para.text.strip():
            continue

        level = get_heading_level(para)

        if level == 1:
            # Chapter title — always CENTER, large, bold
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(heading_size + 4)
                run.font.color.rgb = RGBColor(r, g, b)
                run.font.name = body_font

        elif level >= 2:
            # Subheadings (2.3, 4.5 etc) — LEFT aligned
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(heading_size - (level - 2))
                run.font.color.rgb = RGBColor(r, g, b)
                run.font.name = body_font

        elif para.style and para.style.name and para.style.name.startswith('Heading'):
            # Any other heading style — LEFT
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(heading_size)
                run.font.color.rgb = RGBColor(r, g, b)
                run.font.name = body_font

        else:
            # Body text — smart alignment
            final_align = smart_body_align(para, body_alignment)
            para.alignment = final_align
            # Force via XML to override inherited styles
            pPr = para._p.get_or_add_pPr()
            for jc in pPr.findall(qn('w:jc')):
                pPr.remove(jc)
            jc_el = OxmlElement('w:jc')
            val_map = {
                WD_ALIGN_PARAGRAPH.JUSTIFY: 'both',
                WD_ALIGN_PARAGRAPH.CENTER:  'center',
                WD_ALIGN_PARAGRAPH.RIGHT:   'right',
                WD_ALIGN_PARAGRAPH.LEFT:    'left',
            }
            jc_el.set(qn('w:val'), val_map.get(final_align, 'left'))
            pPr.append(jc_el)
            for run in para.runs:
                run.font.size = Pt(body_size)
                run.font.name = body_font
                # Word spacing only on non-justify — JUSTIFY + spacing = double gap
                if final_align != WD_ALIGN_PARAGRAPH.JUSTIFY:
                    apply_word_spacing(run, word_spacing)

# ══════════════════════════════════════════════════════════════
#  BOOK
# ══════════════════════════════════════════════════════════════
def format_book(doc, opts):
    apply_page_size(doc, opts.get('page_size', 'A4'))
    set_margins(doc, top=1.0, bottom=1.0, left=1.5, right=1.0)
    add_page_border(doc, '2E4057')

    font          = get_font(opts, 'book')
    alignment     = get_alignment(opts, 'justify')
    word_spacing  = get_word_spacing(opts)
    show_pg       = opts.get('page_numbers', False)
    pg_pos        = opts.get('page_number_position', 'center')

    format_paragraphs(doc, body_font=font, body_size=12,
                      heading_size=16, body_alignment=alignment,
                      word_spacing=word_spacing)

    header_text = opts.get('header') or opts.get('title') or ''
    footer_parts = []
    if opts.get('footer'):      footer_parts.append(opts['footer'])
    if opts.get('volume'):      footer_parts.append(opts['volume'])
    if opts.get('website_url'): footer_parts.append(opts['website_url'])
    if opts.get('isbn'):        footer_parts.append('ISBN: ' + opts['isbn'])

    set_header(doc, header_text, font_name=font)
    set_footer(doc, '  |  '.join(footer_parts) if footer_parts else '',
               show_page_numbers=show_pg, page_position=pg_pos, font_name=font)

# ══════════════════════════════════════════════════════════════
#  THESIS
# ══════════════════════════════════════════════════════════════
def format_thesis(doc, opts):
    apply_page_size(doc, opts.get('page_size', 'A4'))
    set_margins(doc, top=1.2, bottom=1.0, left=1.5, right=1.0)
    add_page_border(doc, '1a1a5e')

    font         = get_font(opts, 'thesis')
    alignment    = get_alignment(opts, 'justify')
    word_spacing = get_word_spacing(opts)
    show_pg      = opts.get('page_numbers', False)
    pg_pos       = opts.get('page_number_position', 'center')

    format_paragraphs(doc, body_font=font, body_size=12,
                      heading_size=14, heading_color='1a1a5e',
                      body_alignment=alignment, word_spacing=word_spacing)

    header_parts = []
    if opts.get('university'): header_parts.append(opts['university'])
    if opts.get('department'):  header_parts.append(opts['department'])
    header_text = opts.get('header') or ' — '.join(header_parts)

    footer_parts = []
    if opts.get('footer'):     footer_parts.append(opts['footer'])
    if opts.get('supervisor'): footer_parts.append('Supervisor: ' + opts['supervisor'])
    if opts.get('year'):       footer_parts.append(opts['year'])

    set_header(doc, header_text, font_name=font)
    set_footer(doc, '  |  '.join(footer_parts) if footer_parts else '',
               show_page_numbers=show_pg, page_position=pg_pos, font_name=font)

# ══════════════════════════════════════════════════════════════
#  RESEARCH
# ══════════════════════════════════════════════════════════════
def format_research(doc, opts):
    apply_page_size(doc, opts.get('page_size', 'A4'))
    set_margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0)
    add_page_border(doc, '1a4a2a')

    font         = get_font(opts, 'research')
    alignment    = get_alignment(opts, 'justify')
    word_spacing = get_word_spacing(opts)
    show_pg      = opts.get('page_numbers', False)
    pg_pos       = opts.get('page_number_position', 'center')

    format_paragraphs(doc, body_font=font, body_size=11,
                      heading_size=13, heading_color='1a4a2a',
                      body_alignment=alignment, word_spacing=word_spacing)

    header_text = opts.get('header') or opts.get('journal') or ''
    footer_parts = []
    if opts.get('footer'): footer_parts.append(opts['footer'])
    if opts.get('volume'): footer_parts.append(opts['volume'])
    if opts.get('doi'):    footer_parts.append('DOI: ' + opts['doi'])

    set_header(doc, header_text, font_name=font)
    set_footer(doc, '  |  '.join(footer_parts) if footer_parts else '',
               show_page_numbers=show_pg, page_position=pg_pos, font_name=font)

# ══════════════════════════════════════════════════════════════
#  LETTER
# ══════════════════════════════════════════════════════════════
def format_letter(doc, opts):
    apply_page_size(doc, opts.get('page_size', 'A4'))
    set_margins(doc, top=1.2, bottom=1.0, left=1.2, right=1.0)
    add_page_border(doc, '5a3010')

    font         = get_font(opts, 'letter')
    alignment    = get_alignment(opts, 'left')
    word_spacing = get_word_spacing(opts)
    show_pg      = opts.get('page_numbers', False)
    pg_pos       = opts.get('page_number_position', 'center')

    format_paragraphs(doc, body_font=font, body_size=11,
                      heading_size=13, heading_color='5a3010',
                      body_alignment=alignment, word_spacing=word_spacing)

    header_text = opts.get('header') or opts.get('org_name') or ''
    footer_parts = []
    if opts.get('footer'):      footer_parts.append(opts['footer'])
    if opts.get('website_url'): footer_parts.append(opts['website_url'])
    if opts.get('ref_no'):      footer_parts.append('Ref: ' + opts['ref_no'])

    set_header(doc, header_text, font_name=font)
    set_footer(doc, '  |  '.join(footer_parts) if footer_parts else '',
               show_page_numbers=show_pg, page_position=pg_pos, font_name=font)

# ══════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════
def main():
    doc = Document(input_path)

    if doc_type == 'book':       format_book(doc, options)
    elif doc_type == 'thesis':   format_thesis(doc, options)
    elif doc_type == 'research': format_research(doc, options)
    elif doc_type == 'letter':   format_letter(doc, options)
    else:                        format_book(doc, options)

    doc.save(output_path)
    print(f"Done! Saved to: {output_path}")

main()

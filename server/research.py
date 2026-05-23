import re
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from utils import (
    has_drawing, run_has_drawing, is_all_bold, is_bullet_para,
    apply_para_formatting, set_para_text_formatted, set_font_properly,
    format_table_cells, add_run_with_font, is_krutidev, inject_heading_number
)

# ═══════════════════════════════════════════════
# CONSTANTS
# ═══════════════════════════════════════════════

RESEARCH_FONT     = 'Times New Roman'
BLACK             = RGBColor(0, 0, 0)
BASE_SIZE         = 14.0
TITLE_SIZE        = 14.0
LINE_SPACING      = 1.15

INTRO_TRIGGER_WORDS = {'introduction', 'background', 'overview', 'motivation'}
PRE_INTRO_SECTIONS  = {'abstract', 'keywords', 'keyword', 'key words'}
COMMON_HEADINGS     = {
    'discussion', 'conclusion', 'methodology', 'methods', 'results',
    'findings', 'analysis', 'recommendations', 'limitations',
    'future work', 'acknowledgements', 'acknowledgment',
    'research objectives', 'research methodology', 'questionnaire',
    'case study', 'data analysis', 'literature review',
    'hypothesis', 'scope', 'significance', 'problem statement',
    'research design', 'sampling', 'data collection', 'ethical considerations',
    'appendix', 'appendices',
}
REFERENCE_TRIGGERS  = {'references', 'bibliography', 'works cited', 'संदर्भ', 'ग्रंथसूची'}

QUESTIONNAIRE_TRIGGERS = {
    'questionnaire', 'survey', 'survey questions',
    'section a', 'section b', 'section c', 'section d',
    'part a', 'part b', 'part c', 'part d',
}

# ═══════════════════════════════════════════════
# TITLE PAGE / HEADER
# ═══════════════════════════════════════════════

def insert_research_title_page(doc, opts, font_name):
    """Insert title page when user provides explicit metadata"""
    font = font_name if font_name else RESEARCH_FONT
    title = opts.get('title', '').strip()
    if not title:
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    set_font_properly(r, font)
    r.font.size = Pt(TITLE_SIZE)
    r.font.color.rgb = BLACK

    for key in ['author', 'institution']:
        val = opts.get(key, '').strip()
        if val:
            ap = doc.add_paragraph()
            ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ar = ap.add_run(val)
            ar.bold = True
            set_font_properly(ar, font)
            ar.font.size = Pt(12)

    body = doc.element.body
    for p in reversed(doc.paragraphs[-3:]):
        if p.text.strip() == title:
            body.remove(p._element)
            body.insert(0, p._element)

# ═══════════════════════════════════════════════
# DETECTION LOGIC
# ═══════════════════════════════════════════════

def _normalize(text):
    return text.lower().strip().rstrip(':').strip()

def _get_existing_number(text):
    """Extract existing section number from text like '2. Discussion' -> 2"""
    m = re.match(r'^(\d+)\.?\s+', text)
    if m:
        return int(m.group(1))
    return None

def _is_abstract_heading(text):
    """Check if text is the Abstract heading"""
    return _normalize(text) == 'abstract'

def _is_introduction_heading(text):
    """Check if text is the Introduction heading (numbering starts here)"""
    return _normalize(text) in INTRO_TRIGGER_WORDS

def _is_reference_heading_text(text):
    """Check if text is a reference heading (including variations like 'References (APA)')"""
    clean = _normalize(text)
    base = re.sub(r'\s*\(.*?\)\s*', '', clean).strip()
    return base in REFERENCE_TRIGGERS

def _split_reference_blob(text):
    """Split a reference blob into individual reference entries."""
    if not text.strip():
        return []

    apa_start = r'(?:[A-Z][a-zà-ü]+(?:-[A-Z][a-zà-ü]+)?,\s+[A-Z]\.(?:\s*[A-Z]\.)*(?:\s*&?\s*[A-Z][a-zà-ü]+,\s+[A-Z]\.(?:\s*[A-Z]\.)*)?\s*\(\d{4}[a-z]?\))'
    org_start = r'(?:[A-Z][a-zà-ü]+(?:\s+[A-Z][a-zà-ü]+){2,}(?:\s+\([^)]+\))?\s*\(\d{4}[a-z]?\))'
    combined = f'({apa_start}|{org_start})'

    matches = list(re.finditer(combined, text))

    if len(matches) <= 1:
        return [text.strip()]

    entries = []
    for i, match in enumerate(matches):
        start = match.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        entry = text[start:end].strip()
        if len(entry) > 20:
            entries.append(entry)
        elif entries:
            entries[-1] = entries[-1] + ' ' + entry

    return entries


def detect_research_structure(para, in_numbered_zone, in_reference_zone, in_questionnaire_zone=False, past_abstract=False):
    text = para.text.strip()
    words = text.split()
    wc = len(words)

    if wc == 0:
        return 'empty'
    if has_drawing(para):
        return 'drawing'

    style_name = para.style.name if para.style else ''
    is_bold = is_all_bold(para)
    is_heading_style = style_name in ('Heading 1', 'Heading 2', 'Heading 3')

    # Reference Heading
    if _is_reference_heading_text(text):
        return 'reference_heading'
    if in_reference_zone:
        return 'reference_entry'

    # Questionnaire section headers (Section A:, Section B:, etc.)
    if re.match(r'^Section\s+[A-Z]\s*:', text, re.IGNORECASE) and wc <= 5:
        return 'questionnaire_section_header'

    # Questionnaire parent heading
    if _normalize(text) in QUESTIONNAIRE_TRIGGERS and wc <= 5:
        return 'numbered_section_heading'

    # Numbered List / Questionnaire items (any length)
    if re.match(r'^\d+[\.\)]\s+', text):
        if in_questionnaire_zone:
            return 'questionnaire_item'
        # Distinguish numbered section headings from list items:
        # Heading clues: short (<=8 words), no trailing period, not a sentence fragment
        stripped = re.sub(r'^\d+[\.\)]\s+', '', text).strip()
        stripped_norm = _normalize(stripped)
        is_likely_heading = (
            wc <= 8
            and not text.endswith('.')
            and not text.endswith(',')
            and len(stripped) > 0
        )
        if _is_introduction_heading(stripped) and len(stripped.split()) <= 5:
            return 'numbered_section_heading'
        if stripped_norm in COMMON_HEADINGS:
            return 'numbered_section_heading'
        if is_likely_heading and (is_bold or is_heading_style or past_abstract):
            return 'numbered_section_heading'
        return 'numbered_list_item'

    # Heading 2 + "Table" → table heading (no numbering)
    if style_name == 'Heading 2' and re.match(r'^Table\s*\d*:?', text, re.IGNORECASE):
        return 'section_heading'

    # Keywords line (already merged by pre-pass)
    if re.match(r'^keywords?\s*:', text, re.IGNORECASE):
        return 'keywords_line'

    # Abstract heading → section_heading (NO numbering ever)
    if _is_abstract_heading(text) and wc <= 3:
        return 'section_heading'

    # Introduction heading → always numbered_section_heading (even if not past_abstract yet)
    if _is_introduction_heading(text) and wc <= 5:
        return 'numbered_section_heading'

    # Heading 2/3 styles → numbered section heading (only if past Abstract)
    if is_heading_style and style_name != 'Heading 1' and wc <= 15:
        if past_abstract:
            return 'numbered_section_heading'
        else:
            return 'section_heading'

    # Section Headings (by content)
    if _normalize(text) in COMMON_HEADINGS and wc <= 10:
        if past_abstract:
            return 'numbered_section_heading'
        else:
            return 'section_heading'

    if _normalize(text) in PRE_INTRO_SECTIONS and wc <= 4:
        return 'section_heading'

    # Existing numbered heading (1.1, 2.0 etc)
    if re.match(r'^\d+(\.\d+)+\s+', text):
        return 'numbered_subsection_heading'
    if re.match(r'^\d+\.?\s+', text) and wc <= 10:
        return 'numbered_section_heading'

    # Short bold line as heading
    if is_bold and wc <= 12 and not text.endswith('.'):
        if in_numbered_zone:
            return 'numbered_subsection_heading'
        return 'section_heading'

    return 'body'

# ═══════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════

# Namespaces for separator detection
_VML_NS = 'urn:schemas-microsoft-com:vml'
_OO_NS  = 'urn:schemas-microsoft-com:office:office'
_WP_NS  = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
_MC_NS  = 'http://schemas.openxmlformats.org/markup-compatibility/2006'

def _is_gpt_hr_paragraph(para):
    """
    Returns True if paragraph is a GPT-inserted horizontal rule (safe to delete).

    Detection:
      - Contains a VML <v:rect o:hr="t"> element  (the actual line)
      AND
      - Does NOT contain any real visual content:
          * <wp:inline> / <wp:anchor>  → inline/floating image
          * <mc:AlternateContent>      → chart, SmartArt, equation
          * <v:imagedata>              → embedded VML image

    This means a paragraph that is BOTH a HR AND has a real image is kept.
    In practice GPT never mixes these, but the check makes deletion bulletproof.
    """
    p = para._p
    # Must have VML horizontal rule
    has_vml_hr = any(
        rect.get(f'{{{_OO_NS}}}hr') == 't'
        for rect in p.findall(f'.//{{{_VML_NS}}}rect')
    )
    if not has_vml_hr:
        return False
    # Must NOT have real visual content
    has_real_content = (
        p.find(f'.//{{{_WP_NS}}}inline')          is not None or
        p.find(f'.//{{{_WP_NS}}}anchor')          is not None or
        p.find(f'.//{{{_MC_NS}}}AlternateContent') is not None or
        p.find(f'.//{{{_VML_NS}}}imagedata')       is not None
    )
    return not has_real_content


def remove_gpt_hr_lines(doc):
    """
    Remove all GPT-inserted horizontal rule paragraphs from document.
    Safe: never removes paragraphs containing real images, charts, or SmartArt.
    Returns count of removed paragraphs.
    """
    to_remove = [
        para._p for para in doc.paragraphs
        if _is_gpt_hr_paragraph(para)
    ]
    for p_elem in to_remove:
        parent = p_elem.getparent()
        if parent is not None:
            parent.remove(p_elem)
    return len(to_remove)


def _clear_all_indents(para):
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is not None:
        pPr.remove(ind)
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.left_indent = None

def _set_hanging_indent(para, indent=0.25):
    _clear_all_indents(para)
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    twips = int(indent * 1440)
    ind.set(qn('w:left'), str(twips))
    ind.set(qn('w:hanging'), str(twips))
    pPr.append(ind)

def _set_first_line_indent(para, indent=0.25):
    _clear_all_indents(para)
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    twips = int(indent * 1440)
    ind.set(qn('w:firstLine'), str(twips))
    pPr.append(ind)

def _split_merged_numbered_items(doc):
    """
    Pre-pass: Split paragraphs where numbered items or section headers
    are merged into a single paragraph.
    """
    paragraphs = list(doc.paragraphs)
    i = 0

    while i < len(paragraphs):
        para = paragraphs[i]
        text = para.text.strip()

        if not text or has_drawing(para):
            i += 1
            continue

        style_name = para.style.name if para.style else ''

        # Skip headings
        if style_name in ('Heading 1', 'Heading 2', 'Heading 3'):
            i += 1
            continue

        # Pattern: "Section X: something 1. first item 2. second item..."
        section_with_items = re.match(
            r'^(Section\s+[A-Z]\s*:.*?)\s+(\d+\.\s+.+)$',
            text, re.IGNORECASE
        )

        # Pattern: multiple numbered items "1. ... 2. ... 3. ..."
        has_multiple_items = len(re.findall(r'\d+\.\s+', text)) >= 2

        # Pattern: "5 To provide..." (missing dot after number)
        has_numbered_with_missing_dot = re.search(r'[.;]\s+(\d+)\s+(To\s+|to\s+|[A-Z])', text)

        if section_with_items:
            section_header = section_with_items.group(1).strip()
            items_text = section_with_items.group(2).strip()

            para.clear()
            para.add_run(section_header)

            _insert_split_items(doc, para, items_text)
            paragraphs = list(doc.paragraphs)
            i += 1
            continue

        elif has_multiple_items:
            items_text = text
            _insert_split_items(doc, para, items_text, replace_original=True)
            paragraphs = list(doc.paragraphs)
            i += 1
            continue

        elif has_numbered_with_missing_dot:
            fixed_text = re.sub(
                r'([.;])\s+(\d+)\s+(To\s+|to\s+|[A-Z])',
                r'\1\n\2. \3', text
            )
            parts = fixed_text.split('\n')

            parent = para._element.getparent()
            ref_index = list(parent).index(para._element)

            para.clear()
            para.add_run(parts[0].strip())

            for part in parts[1:]:
                part = part.strip()
                if part:
                    new_p = doc.add_paragraph(part)
                    if para.style:
                        new_p.style = para.style
                    ref_index += 1
                    parent.insert(ref_index, new_p._element)

            paragraphs = list(doc.paragraphs)
            i += 1
            continue

        i += 1


def _insert_split_items(doc, current_para, items_text, replace_original=False):
    """Split items_text into individual paragraphs and insert after current_para."""
    parts = re.split(r'(\d+\.\s+)', items_text)

    items = []
    current_num = None
    for part in parts:
        if re.match(r'^\d+\.\s+$', part):
            current_num = part
        elif part.strip() and current_num is not None:
            items.append(current_num + part.strip())
            current_num = None

    if not items:
        return

    parent = current_para._element.getparent()
    ref_index = list(parent).index(current_para._element)

    if replace_original:
        current_para.clear()
        current_para.add_run(items[0])
        ref_index += 1
        for item in items[1:]:
            new_p = doc.add_paragraph(item)
            if current_para.style:
                new_p.style = current_para.style
            parent.insert(ref_index, new_p._element)
            ref_index += 1
    else:
        ref_index += 1
        for item in items:
            new_p = doc.add_paragraph(item)
            if current_para.style:
                new_p.style = current_para.style
            parent.insert(ref_index, new_p._element)
            ref_index += 1

# ═══════════════════════════════════════════════
# PRE-PASS: Fix Keywords + References Structure
# ═══════════════════════════════════════════════

def _handle_keywords_and_refs_prepass(doc, font_name):
    """
    Pre-process document to:
    0. Split merged numbered items and section headers
    1. Merge "Keywords" Heading 2 + next Normal para into "Keywords: content"
    2. Split reference blob into individual paragraphs
    Returns: (title_para_element, table_heading_elements, heading1_elements, 
              abstract_para_element, intro_para_element)
    """
    font = font_name if font_name else RESEARCH_FONT

    # === STEP 0: Split merged numbered items ===
    _split_merged_numbered_items(doc)

    # === STEP 0b: Convert tabular text to actual tables ===
    _convert_tabular_text_to_tables(doc)

    paragraphs = list(doc.paragraphs)

    heading1_elements = set()
    table_heading_elements = set()
    title_para_element = None
    abstract_para_element = None
    intro_para_element = None
    found_title = False
    found_abstract = False
    found_intro = False

    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        text = para.text.strip()
        style_name = para.style.name if para.style else ''

        # Track Heading 1 (title)
        if style_name == 'Heading 1' and not found_title:
            heading1_elements.add(para._p)
            if text:
                title_para_element = para._p
                found_title = True

        # Track Abstract paragraph
        if _is_abstract_heading(text) and not found_abstract:
            abstract_para_element = para._p
            found_abstract = True

        # Track Introduction paragraph
        if _is_introduction_heading(text) and len(text.split()) <= 5 and not found_intro:
            intro_para_element = para._p
            found_intro = True

        # Track Table headings
        if style_name == 'Heading 2' and re.match(r'^Table\s*\d*:?', text, re.IGNORECASE):
            table_heading_elements.add(para._p)

        # Merge Keywords: if Heading 2 "Keywords" + next paragraph
        if style_name == 'Heading 2' and _normalize(text) in ('keywords', 'keyword', 'key words'):
            keyword_para = para
            next_idx = i + 1
            while next_idx < len(paragraphs) and not paragraphs[next_idx].text.strip():
                next_idx += 1

            if next_idx < len(paragraphs):
                next_para = paragraphs[next_idx]
                if next_para.style.name == 'Normal' or next_para.style.name == 'List Paragraph':
                    kw_text = next_para.text.strip()
                    keyword_para.clear()
                    r1 = keyword_para.add_run('Keywords: ')
                    r1.bold = True
                    set_font_properly(r1, font)
                    r2 = keyword_para.add_run(kw_text)
                    set_font_properly(r2, font)
                    keyword_para.style = doc.styles['Normal']

                    next_para._element.getparent().remove(next_para._element)
                    paragraphs = list(doc.paragraphs)
                    continue

        i += 1

    # Split reference blob
    paragraphs = list(doc.paragraphs)
    ref_start_idx = None

    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if _is_reference_heading_text(text) and len(text.split()) <= 6:
            for j in range(i + 1, len(paragraphs)):
                if paragraphs[j].text.strip():
                    ref_start_idx = j
                    break
            if ref_start_idx is not None:
                break

    if ref_start_idx is not None and ref_start_idx < len(paragraphs):
        ref_blob = paragraphs[ref_start_idx]
        blob_text = ref_blob.text.strip()

        year_matches = re.findall(r'\(\d{4}[a-z]?\)', blob_text)
        if len(year_matches) >= 2:
            entries = _split_reference_blob(blob_text)
            if len(entries) > 1:
                parent = ref_blob._element.getparent()
                ref_index = list(parent).index(ref_blob._element)

                for entry in entries:
                    new_p = doc.add_paragraph(entry)
                    if ref_blob.style:
                        new_p.style = ref_blob.style
                    parent.insert(ref_index + 1, new_p._element)
                    ref_index += 1

                parent.remove(ref_blob._element)

    return title_para_element, table_heading_elements, heading1_elements, abstract_para_element, intro_para_element



# ADD this new function anywhere above format_research_body (e.g. after _insert_split_items):

def _convert_tabular_text_to_tables(doc):
    """
    Detect paragraphs that look like tabular data (tab-separated or pipe-separated)
    and convert them into actual Word tables.
    """
    paragraphs = list(doc.paragraphs)
    i = 0

    while i < len(paragraphs):
        para = paragraphs[i]
        text = para.text.strip()

        if not text or has_drawing(para):
            i += 1
            continue

        # Detect separator: tab or pipe
        has_tab  = '\t' in text
        has_pipe = '|' in text and text.count('|') >= 2

        if not has_tab and not has_pipe:
            i += 1
            continue

        sep = '\t' if has_tab else '|'

        # Collect consecutive tabular lines
        group = []
        j = i
        while j < len(paragraphs):
            t = paragraphs[j].text.strip()
            if not t:
                j += 1
                break
            if sep == '\t' and '\t' not in t:
                break
            if sep == '|' and '|' not in t:
                break
            group.append((paragraphs[j], t))
            j += 1

        # Need at least 2 rows to make a table
        if len(group) < 2:
            i += 1
            continue

        # Parse rows and determine column count
        rows_data = []
        for _, row_text in group:
            if sep == '|':
                cells = [c.strip() for c in row_text.strip('|').split('|')]
            else:
                cells = [c.strip() for c in row_text.split('\t')]
            rows_data.append(cells)

        col_count = max(len(r) for r in rows_data)
        if col_count < 2:
            i += 1
            continue

        # Pad rows to same column count
        for row in rows_data:
            while len(row) < col_count:
                row.append('')

        # Insert table before first paragraph of group
        ref_para = group[0][0]
        parent = ref_para._p.getparent()
        ref_idx = list(parent).index(ref_para._p)

        table = doc.add_table(rows=len(rows_data), cols=col_count)
        table.style = 'Table Grid'

        for r_idx, row_cells in enumerate(rows_data):
            for c_idx, cell_text in enumerate(row_cells):
                cell = table.cell(r_idx, c_idx)
                cell.text = cell_text
                # Bold first row (header)
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

        # Insert table XML before first group paragraph
        parent.insert(ref_idx, table._tbl)

        # Remove original paragraphs
        for orig_para, _ in group:
            orig_para._p.getparent().remove(orig_para._p)

        # Refresh paragraph list
        paragraphs = list(doc.paragraphs)
        i = 0  # restart scan — indices shifted


# END of _convert_tabular_text_to_tables

# ═══════════════════════════════════════════════
# MAIN FORMATTER
# ═══════════════════════════════════════════════

def _apply_ref_para_spacing(para, before, after, line, font, size, black):
    """Apply exact reference-file spacing + font to a paragraph.
    before/after/line are twip strings or None to skip. size is pt float."""
    from docx.oxml import OxmlElement as _OE
    from docx.oxml.ns import qn as _qn
    from docx.shared import Pt as _Pt

    # Spacing XML
    _pPr = para._p.get_or_add_pPr()
    _sp = _pPr.find(_qn('w:spacing'))
    if _sp is None:
        _sp = _OE('w:spacing')
        _pPr.append(_sp)
    if before is not None:
        _sp.set(_qn('w:before'), before)
    if after is not None:
        _sp.set(_qn('w:after'), after)
    if line is not None:
        _sp.set(_qn('w:line'), line)
        _sp.set(_qn('w:lineRule'), 'auto')
    for _attr in [_qn('w:beforeLines'), _qn('w:afterLines'),
                  _qn('w:beforeAutospacing'), _qn('w:afterAutospacing')]:
        if _sp.get(_attr) is not None:
            del _sp.attrib[_attr]

    # Fix paragraph mark font (rPr inside pPr)
    _rPr = _pPr.find(_qn('w:rPr'))
    if _rPr is None:
        _rPr = _OE('w:rPr')
        _pPr.append(_rPr)
    _rFonts = _rPr.find(_qn('w:rFonts'))
    if _rFonts is None:
        _rFonts = _OE('w:rFonts')
        _rPr.insert(0, _rFonts)
    _rFonts.set(_qn('w:ascii'), font)
    _rFonts.set(_qn('w:hAnsi'), font)
    _rFonts.set(_qn('w:cs'), font)
    for _sz_name in ['w:sz', 'w:szCs']:
        _sz_el = _rPr.find(_qn(_sz_name))
        if _sz_el is None:
            _sz_el = _OE(_sz_name)
            _rPr.append(_sz_el)
        _sz_el.set(_qn('w:val'), str(int(size * 2)))

    # Fix runs font+size
    for r in para.runs:
        set_font_properly(r, font)
        r.font.size = _Pt(size)
        r.font.color.rgb = black


def _fix_normal_style(doc, font, base_size, line_spacing):
    """
    Fix the doc-level Normal style so pressing Enter creates a paragraph
    with correct font (Times New Roman), size (12pt), and 1.15 line spacing.
    This prevents Mangal/24pt appearing on new paragraphs.
    """
    from docx.oxml import OxmlElement as _OE
    from docx.oxml.ns import qn as _qn
    from docx.shared import Pt as _Pt

    try:
        normal = doc.styles['Normal']
    except KeyError:
        return

    # Fix font
    nf = normal.font
    nf.name = font
    nf.size = _Pt(base_size)
    nf.bold = False
    nf.italic = False
    nf.color.rgb = BLACK

    # Fix paragraph format
    npf = normal.paragraph_format
    npf.space_before = _Pt(0)
    npf.space_after = _Pt(0)
    npf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    npf.line_spacing = line_spacing

    # Fix via XML on the style element too
    el = normal.element
    pPr = el.find(_qn('w:pPr'))
    if pPr is None:
        pPr = _OE('w:pPr')
        el.append(pPr)
    sp = pPr.find(_qn('w:spacing'))
    if sp is None:
        sp = _OE('w:spacing')
        pPr.append(sp)
    sp.set(_qn('w:before'), '0')
    sp.set(_qn('w:after'), '0')
    sp.set(_qn('w:line'), str(int(line_spacing * 240)))
    sp.set(_qn('w:lineRule'), 'auto')

    # Fix rPr (run properties) on style — sets font for paragraph mark
    rPr = el.find(_qn('w:rPr'))
    if rPr is None:
        rPr = _OE('w:rPr')
        el.append(rPr)
    rFonts = rPr.find(_qn('w:rFonts'))
    if rFonts is None:
        rFonts = _OE('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(_qn('w:ascii'), font)
    rFonts.set(_qn('w:hAnsi'), font)
    rFonts.set(_qn('w:cs'), font)
    for _sz_name in ['w:sz', 'w:szCs']:
        _sz_el = rPr.find(_qn(_sz_name))
        if _sz_el is None:
            _sz_el = _OE(_sz_name)
            rPr.append(_sz_el)
        _sz_el.set(_qn('w:val'), str(int(base_size * 2)))


def format_research_body(doc, opts, font_name):
    font = font_name if font_name else RESEARCH_FONT
    base_size = 12.0  # Body/heading font size always 12pt (title uses TITLE_SIZE=14pt)
    line_spacing = float(opts.get('line_spacing', LINE_SPACING))
    krutidev_mode = is_krutidev(font)

    # === STEP -1: Fix doc-level Normal style so Enter gives correct font/size ===
    _fix_normal_style(doc, font, base_size, line_spacing)

    # === STEP 0: Remove GPT horizontal rule separators ===
    remove_gpt_hr_lines(doc)

    # === PRE-PASS ===
    title_elem, table_heading_elems, heading1_elems, abstract_elem, intro_elem = _handle_keywords_and_refs_prepass(doc, font)

    in_numbered_zone = False
    in_reference_zone = False
    in_questionnaire_zone = False
    past_abstract = False
    numbering_started = False
    sec_counter = 0
    sub_counter = 0
    ref_counter = 0
    q_counter = 0
    prev_was_heading = False

    paragraphs = list(doc.paragraphs)

    # 1. Find first paragraph (title)
    first_p = None
    for p in paragraphs:
        text = p.text.strip()
        if not text or has_drawing(p):
            continue
        style_name = p.style.name if p.style else ''
        if style_name == 'Heading 1' or (first_p is None and text):
            first_p = p
            if style_name == 'Heading 1':
                break

    # 2. CENTER everything BEFORE Abstract (FIX: use element index comparison)
    if abstract_elem is not None:
        # Get index of abstract in the body's children
        body = doc.element.body
        body_children = list(body)
        abstract_idx = None
        for idx, child in enumerate(body_children):
            if child == abstract_elem:
                abstract_idx = idx
                break
        
        if abstract_idx is not None:
            for child in body_children[:abstract_idx]:
                # Check if this is a paragraph element
                if child.tag == qn('w:p'):
                    # Find corresponding paragraph object
                    for p in paragraphs:
                        if p._p == child:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                            p.paragraph_format.line_spacing = 1.0
                            for r in p.runs:
                                set_font_properly(r, font)
                                r.font.size = Pt(base_size)
                                r.font.color.rgb = BLACK
                            break

    # 3. Format title
    if first_p:
        first_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _clear_all_indents(first_p)
        first_p.paragraph_format.space_before = Pt(0)
        first_p.paragraph_format.space_after = Pt(0)
        first_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        first_p.paragraph_format.line_spacing = 1.0
        # Force via XML
        from docx.oxml import OxmlElement as _OE
        from docx.oxml.ns import qn as _qn
        _pPr = first_p._p.get_or_add_pPr()
        _sp = _pPr.find(_qn('w:spacing'))
        if _sp is None:
            _sp = _OE('w:spacing')
            _pPr.append(_sp)
        _sp.set(_qn('w:before'), '0')
        _sp.set(_qn('w:after'), '0')
        _sp.set(_qn('w:beforeLines'), '0')
        _sp.set(_qn('w:afterLines'), '0')
        _sp.set(_qn('w:line'), '240')
        _sp.set(_qn('w:lineRule'), 'auto')
        for attr in [_qn('w:beforeAutospacing'), _qn('w:afterAutospacing')]:
            if _sp.get(attr) is not None:
                del _sp.attrib[attr]
        for r in first_p.runs:
            set_font_properly(r, font)
            r.bold = True
            r.font.size = Pt(TITLE_SIZE)
            r.font.color.rgb = BLACK

        txt = first_p.text.strip()
        if re.match(r'^\d+\.?\s+', txt):
            new_txt = re.sub(r'^\d+\.?\s+', '', txt)
            if first_p.runs:
                first_p.runs[0].text = new_txt

    # 4. Sync counter with already-numbered headings (only after Introduction)
    for para in paragraphs:
        text = para.text.strip()
        if not text or para == first_p:
            continue
        if intro_elem is not None and para._p == intro_elem:
            numbering_started = True
            sec_counter = 0  # Reset — Introduction will set to 1
        if not numbering_started:
            continue
        existing = _get_existing_number(text)
        if existing is not None:
            etype = detect_research_structure(para, True, False, False, True)
            if etype in ('numbered_section_heading', 'numbered_subsection_heading'):
                if existing > sec_counter:
                    sec_counter = existing

    # 5. Body Processing
    prev_body_para = None
    numbering_started = False

    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text or has_drawing(para):
            if text:
                prev_body_para = None
            continue
        if para == first_p:
            prev_body_para = None
            continue

        para_elem = para._p
        style_name = para.style.name if para.style else ''

        # Check if we've crossed Abstract
        if abstract_elem is not None and para_elem == abstract_elem:
            past_abstract = True

        # Check if we've hit Introduction (FIX: numbering starts at 1)
        if intro_elem is not None and para_elem == intro_elem:
            numbering_started = True
            sec_counter = 0  # Reset to 0 so next increment gives 1

        etype = detect_research_structure(para, in_numbered_zone, in_reference_zone, in_questionnaire_zone, past_abstract)

        # --- REFERENCE ZONE ENTRY/EXIT ---
        if etype == 'reference_heading':
            in_reference_zone = True
            in_numbered_zone = False
            in_questionnaire_zone = False
            ref_counter = 0

        if in_reference_zone and etype not in ('reference_heading', 'reference_entry', 'empty'):
            in_reference_zone = False

        if etype == 'numbered_section_heading':
            in_numbered_zone = True
            if _normalize(text) in QUESTIONNAIRE_TRIGGERS:
                in_questionnaire_zone = True
                q_counter = 0

        # --- QUESTIONNAIRE SECTION HEADER ---
        if etype == 'questionnaire_section_header':
            _clear_all_indents(para)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.line_spacing = 1.15
            for r in para.runs:
                set_font_properly(r, font)
                r.bold = True
                r.font.size = Pt(base_size)
                r.font.color.rgb = BLACK
            prev_body_para = None
            continue

        # --- QUESTIONNAIRE ITEMS ---
        if etype == 'questionnaire_item':
            q_counter += 1
            _set_hanging_indent(para, 0.35)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing = 1.15
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.space_before = Pt(0)

            for r in para.runs:
                set_font_properly(r, font)
                r.bold = False
                r.font.size = Pt(base_size)
                r.font.color.rgb = BLACK
            prev_body_para = None
            continue

        # --- EMPTY PARAGRAPHS: reference values before=100 after=100 line=360 ---
        if etype == 'empty':
            _apply_ref_para_spacing(para, '100', '100', '360', font, base_size, BLACK)
            continue

        # --- HEADINGS ---
        if etype in ('section_heading', 'numbered_section_heading', 'numbered_subsection_heading', 'reference_heading'):
            is_title = (style_name == 'Heading 1' or para_elem in heading1_elems)
            is_table = (style_name == 'Heading 2' and re.match(r'^Table\s*\d*:?', text, re.IGNORECASE)) or (para_elem in table_heading_elems)
            is_abstract = (abstract_elem is not None and para_elem == abstract_elem)
            is_intro = (intro_elem is not None and para_elem == intro_elem)
            already_numbered = _get_existing_number(text) is not None
            is_ref = (etype == 'reference_heading')

            if is_ref or is_title or is_table or is_abstract:
                # No numbering for these
                pass
            elif is_intro:
                # Introduction always gets number 1
                if not already_numbered:
                    sec_counter = 1
                    sub_counter = 0
                    inject_heading_number(para, sec_counter, krutidev_mode=krutidev_mode)
            elif etype == 'numbered_section_heading' and numbering_started:
                if not already_numbered:
                    sec_counter += 1
                    sub_counter = 0
                    inject_heading_number(para, sec_counter, krutidev_mode=krutidev_mode)
                else:
                    existing = _get_existing_number(text)
                    if existing is not None:
                        sec_counter = existing
                        sub_counter = 0
            elif etype == 'numbered_subsection_heading' and numbering_started:
                if not already_numbered:
                    sub_counter += 1
                    inject_heading_number(para, sec_counter, sub_counter, krutidev_mode=krutidev_mode)
                else:
                    existing = _get_existing_number(text)
                    if existing is not None:
                        sub_counter = existing

            # _clear_all_indents(para)
            # para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # para.paragraph_format.space_before = Pt(12)
            # para.paragraph_format.space_after = Pt(6)
            # para.paragraph_format.line_spacing = 1.5
            # for r in para.runs:
            #     set_font_properly(r, font)
            #     r.bold = True
            #     r.font.size = Pt(base_size)
            #     r.font.color.rgb = BLACK
            # prev_body_para = None
            # continue


            _clear_all_indents(para)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            para.paragraph_format.line_spacing = line_spacing
            # Remove Word's auto-spacing via XML directly
            pPr = para._p.get_or_add_pPr()
            from docx.oxml import OxmlElement as _OE
            from docx.oxml.ns import qn as _qn
            spacing_el = pPr.find(_qn('w:spacing'))
            if spacing_el is None:
                spacing_el = _OE('w:spacing')
                pPr.append(spacing_el)
            # Reference file exact values: before=100, after=100, line=360 (1.5)
            spacing_el.set(_qn('w:before'), '100')
            spacing_el.set(_qn('w:after'), '100')
            spacing_el.set(_qn('w:line'), '360')
            spacing_el.set(_qn('w:lineRule'), 'auto')
            for attr in [_qn('w:beforeLines'), _qn('w:afterLines'),
                         _qn('w:beforeAutospacing'), _qn('w:afterAutospacing')]:
                if spacing_el.get(attr) is not None:
                    del spacing_el.attrib[attr]
            # Remove contextualSpacing (Word adds gap between heading and body)
            ctx = pPr.find(_qn('w:contextualSpacing'))
            if ctx is not None:
                pPr.remove(ctx)
            # Override style-level spacing so Heading 1/2/3 style defaults don't bleed through
            pStyle = pPr.find(_qn('w:pStyle'))
            if pStyle is not None:
                style_val = pStyle.get(_qn('w:val'), '')
                if style_val and style_val in doc.styles:
                    try:
                        sobj = doc.styles[style_val]
                        sobj.paragraph_format.space_before = Pt(0)
                        sobj.paragraph_format.space_after  = Pt(0)
                    except Exception:
                        pass
            # Flag: next body para should get space_before
            prev_was_heading = True
            for r in para.runs:
                set_font_properly(r, font)
                r.bold = True
                r.font.size = Pt(base_size)
                r.font.color.rgb = BLACK
            prev_body_para = None
            continue



        # --- KEYWORDS ---
        if etype == 'keywords_line':
            _clear_all_indents(para)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            para.paragraph_format.line_spacing = line_spacing
            full_txt = para.text
            para.clear()
            m = re.match(r'^(keywords?\s*:)(.*)', full_txt, re.IGNORECASE)
            if m:
                r1 = para.add_run(m.group(1))
                r1.bold = True
                r1.italic = True
                set_font_properly(r1, font)
                r1.font.size = Pt(11)
                r2 = para.add_run(m.group(2))
                r2.italic = True
                r2.bold = False
                set_font_properly(r2, font)
                r2.font.size = Pt(11)
            # Force XML
            from docx.oxml import OxmlElement as _OE
            from docx.oxml.ns import qn as _qn
            _pPr = para._p.get_or_add_pPr()
            _sp  = _pPr.find(_qn('w:spacing'))
            if _sp is None:
                _sp = _OE('w:spacing')
                _pPr.append(_sp)
            _sp.set(_qn('w:before'), '0')
            _sp.set(_qn('w:after'), '0')
            _sp.set(_qn('w:line'), str(int(line_spacing * 240)))
            _sp.set(_qn('w:lineRule'), 'auto')
            prev_body_para = None
            continue

        # --- LIST ITEMS / BULLETS ---
        if etype == 'numbered_list_item' or is_bullet_para(para):
            _set_hanging_indent(para, 0.25)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Reference file exact values: before=100, after=100, no line override
            _apply_ref_para_spacing(para, '100', '100', None, font, base_size, BLACK)
            prev_body_para = None
            continue

        # --- REFERENCE ENTRIES ---
        if etype == 'reference_entry':
            ref_counter += 1
            _set_hanging_indent(para, 0.3)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing = 1.15
            para.paragraph_format.space_after = Pt(12)

            current_text = para.text.strip()
            current_text = re.sub(r'^[•\-\*]\s*', '', current_text)
            current_text = re.sub(r'^\[\d+\]\s*', '', current_text)

            para.clear()
            r_num = para.add_run(f'[{ref_counter}] ')
            set_font_properly(r_num, font)
            r_num.font.size = Pt(11)
            r_num.font.color.rgb = BLACK

            r_text = para.add_run(current_text)
            set_font_properly(r_text, font)
            r_text.font.size = Pt(11)
            r_text.font.color.rgb = BLACK

            prev_body_para = None
            continue

        # --- BODY TEXT ---
        _clear_all_indents(para)

        prev_body = None
        for j in range(i - 1, -1, -1):
            prev_para = paragraphs[j]
            prev_text = prev_para.text.strip()
            if not prev_text or has_drawing(prev_para):
                continue
            prev_past_abstract = past_abstract
            if abstract_elem is not None:
                for k in range(j, -1, -1):
                    if paragraphs[k]._p == abstract_elem:
                        prev_past_abstract = True
                        break
            prev_estyle = detect_research_structure(prev_para, in_numbered_zone, in_reference_zone, in_questionnaire_zone, prev_past_abstract)
            if prev_estyle == 'body':
                prev_body = prev_para
            break

        # if prev_body is not None:
        #     _set_first_line_indent(para, 0.25)
        #     prev_body_para = para
        # else:
        #     prev_body_para = para

        # para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # para.paragraph_format.line_spacing = 1.5
        # para.paragraph_format.space_after = Pt(12)
        # for r in para.runs:
        #     set_font_properly(r, font)
        #     r.bold = False
        #     r.font.size = Pt(base_size)
        #     r.font.color.rgb = BLACK


        prev_body_para = para

        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Reference file exact values: before=100, after=100, no explicit line (inherited)
        _apply_ref_para_spacing(para, '100', '100', None, font, base_size, BLACK)
        # Ensure bold is off for body text
        for r in para.runs:
            r.bold = False
        prev_was_heading = False









    format_table_cells(doc, font, base_size, line_spacing, BLACK)

    # No post-processing needed — reference values applied directly


def _apply_transition_spacing(doc, in_numbered_zone=False, in_reference_zone=False, in_questionnaire_zone=False):
    """
    Separate post-processing pass — adds 1.5 line (360 twips) space_before to any paragraph
    that immediately follows a heading, OR is the first bullet after a body paragraph.
    Does NOT touch any other paragraph spacing. Safe to call after main formatting loop.
    """
    from docx.oxml import OxmlElement as _OE
    from docx.oxml.ns import qn as _qn
    from docx.shared import Pt

    # Word line spacing: 1.5 lines = w:line of 360 (240 * 1.5)
    # Space before paragraph: reference doc uses 100 twips (5pt) before heading/body/bullet transitions
    SPACING_BEFORE  = '360'   # 18pt space before transitions (1.5 line visual gap)
    LINE_SPACING_1_5 = '360'  # 1.5 line spacing = 360 twips

    paragraphs = [p for p in doc.paragraphs if p._element.getparent() is not None]
    total = len(paragraphs)

    # Build a simple label list: 'heading', 'bullet', 'body', 'other'
    def _label(para):
        txt = para.text.strip()
        if not txt:
            return 'empty'
        if is_bullet_para(para):
            return 'bullet'
        pPr = para._p.find(_qn('w:pPr'))
        if pPr is not None:
            pStyle = pPr.find(_qn('w:pStyle'))
            if pStyle is not None:
                val = pStyle.get(_qn('w:val'), '').lower()
                if 'heading' in val:
                    return 'heading'
        # Detect bold-only paragraph (subheading / heading style)
        if txt and all(r.bold for r in para.runs if r.text.strip()):
            if para.runs:
                return 'heading'
        return 'body'

    labels = [_label(p) for p in paragraphs]

    def _set_before(para, before_twips, line_twips=None):
        """Set ONLY w:before on paragraph — does NOT touch line spacing."""
        from docx.shared import Pt as _Pt
        pt_val = int(before_twips) / 20.0
        para.paragraph_format.space_before = _Pt(pt_val)
        _pPr = para._p.get_or_add_pPr()
        _sp = _pPr.find(_qn('w:spacing'))
        if _sp is None:
            _sp = _OE('w:spacing')
            _pPr.append(_sp)
        _sp.set(_qn('w:before'), before_twips)
        # Do NOT set w:line here — preserve existing 1.15 line spacing set by main loop
        for _attr in [_qn('w:beforeAutospacing')]:
            if _sp.get(_attr) is not None:
                del _sp.attrib[_attr]

    for i in range(1, total):
        cur  = labels[i]
        prev = labels[i - 1]

        if cur in ('empty', 'other'):
            continue

        # Case 1: anything right after a heading → space_before + 1.5 line spacing
        if prev == 'heading' and cur in ('body', 'bullet'):
            _set_before(paragraphs[i], SPACING_BEFORE, LINE_SPACING_1_5)

        # Case 2: first bullet after a body paragraph → space_before + 1.5 line spacing
        elif prev == 'body' and cur == 'bullet':
            _set_before(paragraphs[i], SPACING_BEFORE, LINE_SPACING_1_5)

        # Case 3: heading right after another heading → space_before + 1.5 line spacing
        elif prev == 'heading' and cur == 'heading':
            _set_before(paragraphs[i], SPACING_BEFORE, LINE_SPACING_1_5)










""" 

yaha pr bolna dashboard ke call automation agents zoho crm se connect hai aur abhi thanks call agent me ek changes karna hai -

mere pass 7 different websites hai aur jab bhi user x website se form fill karta hai to vo zoho crm me store hota hai aur
than thanks call tirgger hota hai uss number ke liye to abhi jab bhi thanks call jata hai to greetings me wo edwin incorpartion 
means parent copany ka naam leta hai but mai chahalta hu ki jiss plateform se form fill kiya gya hai thanks call me greeting ussi
organization ke name se ho like agar form xyz page se fill hua hai to - hi i'm calling from xyz organization aur yese hi sabhi ke liye 
batao kya approch hoga isko solve karne ka aur kaha aur kya changes karne padenge? 

-----------

iss repo ko dekho yaha pr maine ek stand alone api calculator banaya tha but now mujhe yesa chahiye ki sift ugc wala part hai ho aur usse me deploye
kru aur current zohosite pr bani ek website ke navbar me connect kru 

but yaha mujhe logic logout feature chahiye jiska id and password ek hi hogna fixed aur koi bhi sift ussi he login kr payega

to mai soch rha hu ki ek html, css, aur js ka use kr ke banou usse vercel ya kahi bhi deploye kr du aur uska link zoho sitesme 
ke nav me link kr du to batao ki kya approch hona chahiye page to thik hai lekin login wale feature ka kya karenge yaha database to nhi 
chahiye pr result calculate hone ke baad pdf me export hone wala feature chahiye batao mujhe kya approach hona chahiye? 


"""














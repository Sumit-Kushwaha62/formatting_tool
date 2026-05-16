import re
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from utils import (
    has_drawing, run_has_drawing, is_all_bold, is_bullet_para,
    apply_para_formatting, set_font_properly, format_table_cells,
    is_krutidev, CHAPTER_HEADING_RE, CHAPTER_HEADING_LOOSE_RE
)


# ═══════════════════════════
# TITLE PAGE — THESIS
# ═══════════════════════════

def insert_thesis_title_page(doc, opts, font_name):
    black      = RGBColor(0, 0, 0)
    title      = opts.get('title', '').strip()
    author     = opts.get('author', '').strip()
    university = opts.get('university', '').strip()
    department = opts.get('department', '').strip()
    supervisor = opts.get('supervisor', '').strip()
    year       = opts.get('year', '').strip()

    if not title and not author:
        return

    insert_paras = []

    def make_para(text, align, size, bold=False, italic=False,
                  space_before=0, space_after=10, color=None):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        if text:
            r = p.add_run(text)
            r.bold   = bold
            r.italic = italic
            set_font_properly(r, font_name)
            r.font.size = Pt(size)
            r.font.color.rgb = color or black
        return p

    def add_horizontal_rule(thick=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        pPr   = p._p.get_or_add_pPr()
        pBdr  = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '12' if thick else '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(36)
    spacer.paragraph_format.space_after  = Pt(0)
    insert_paras.append(spacer)

    if university:
        insert_paras.append(make_para(university, WD_ALIGN_PARAGRAPH.CENTER, 16, bold=True, space_after=4))
    if department:
        insert_paras.append(make_para(department, WD_ALIGN_PARAGRAPH.CENTER, 12, space_after=20))

    insert_paras.append(add_horizontal_rule(thick=True))

    sp2 = doc.add_paragraph()
    sp2.paragraph_format.space_before = Pt(24)
    sp2.paragraph_format.space_after  = Pt(0)
    insert_paras.append(sp2)

    insert_paras.append(make_para('A Thesis Submitted in Partial Fulfillment of the',
                                   WD_ALIGN_PARAGRAPH.CENTER, 11, italic=True, space_after=2))
    insert_paras.append(make_para('Requirements for the Degree',
                                   WD_ALIGN_PARAGRAPH.CENTER, 11, italic=True, space_after=24))

    if title:
        insert_paras.append(make_para(title, WD_ALIGN_PARAGRAPH.CENTER, 22, bold=True,
                                       space_before=8, space_after=28))

    insert_paras.append(add_horizontal_rule(thick=False))

    insert_paras.append(make_para('Submitted by', WD_ALIGN_PARAGRAPH.CENTER, 10, italic=True,
                                   space_before=20, space_after=4))
    if author:
        insert_paras.append(make_para(author, WD_ALIGN_PARAGRAPH.CENTER, 15, bold=True, space_after=4))

    if supervisor:
        insert_paras.append(make_para('Under the Supervision of', WD_ALIGN_PARAGRAPH.CENTER,
                                       10, italic=True, space_before=16, space_after=4))
        insert_paras.append(make_para(supervisor, WD_ALIGN_PARAGRAPH.CENTER, 13, bold=True, space_after=4))

    if year:
        insert_paras.append(make_para(year, WD_ALIGN_PARAGRAPH.CENTER, 12, space_before=20, space_after=0))

    pb_para = doc.add_paragraph()
    pb_para.paragraph_format.space_before = Pt(0)
    pb_para.paragraph_format.space_after  = Pt(0)
    run = pb_para.add_run()
    br  = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    insert_paras.append(pb_para)

    body = doc.element.body
    for p in reversed(insert_paras):
        body.remove(p._element)
        body.insert(0, p._element)


# ═══════════════════════════
# THESIS STRUCTURE DETECTION
# ═══════════════════════════

def detect_thesis_structure(para, index, doc):
    """Thesis-aware structure detection."""
    text  = para.text.strip()
    words = text.split()
    wc    = len(words)

    if wc == 0:
        return 'empty'

    if has_drawing(para):
        return 'drawing'

    if is_bullet_para(para):
        return 'bullet'

    is_bold = is_all_bold(para)

    # Caption keyword match: only SHORT lines (<=15 words) to avoid catching
    # body sentences like 'Chart 11.3 illustrates the ethical trade-offs...'
    if wc <= 15 and re.match(r'^(table|figure|fig|chart|graph|diagram|image|photo|plate|'
                              r'तालिका|चित्र|आकृति|ग्राफ)\s*[\.\-–—:1-9]', text, re.IGNORECASE):
        return 'figure_caption'

    # Para BEFORE a drawing: only treat as caption if short or starts with caption keyword
    # Look ahead skipping empty paragraphs
    look_ahead = index + 1
    while look_ahead < len(doc.paragraphs) and not doc.paragraphs[look_ahead].text.strip():
        look_ahead += 1
    if look_ahead < len(doc.paragraphs) and has_drawing(doc.paragraphs[look_ahead]):
        is_caption_text = (
            re.match(r'^(table|figure|fig|chart|graph|diagram|image|photo|plate|'
                     r'तालिका|चित्र|आकृति|ग्राफ)\b', text, re.IGNORECASE)
            or wc <= 20
        )
        if is_caption_text:
            return 'figure_caption'

    # Para AFTER a drawing: only short/caption-keyword lines are captions
    # Look back skipping empty paragraphs
    look_back = index - 1
    while look_back >= 0 and not doc.paragraphs[look_back].text.strip():
        look_back -= 1
    if look_back >= 0 and has_drawing(doc.paragraphs[look_back]):
        is_caption_text = (
            re.match(r'^(table|figure|fig|chart|graph|diagram|image|photo|plate|'
                     r'तालिका|चित्र|आकृति|ग्राफ)\b', text, re.IGNORECASE)
            or wc <= 25
        )
        if is_caption_text:
            return 'figure_caption'
        # Long paragraph after drawing → fall through to body detection

    if CHAPTER_HEADING_RE.match(text) and wc <= 20:
        return 'chapter_heading'

    # 'षष्ठम अध्याय: ...' style — ordinal word before अध्याय/chapter
    if CHAPTER_HEADING_LOOSE_RE.match(text) and wc <= 20:
        return 'chapter_heading'

    if index > 0:
        prev_text = doc.paragraphs[index - 1].text.strip()
        if (CHAPTER_HEADING_RE.match(prev_text) or CHAPTER_HEADING_LOOSE_RE.match(prev_text)) and wc <= 15:
            return 'chapter_heading'

    special_sections = {
        'abstract', 'introduction', 'references', 'bibliography',
        'acknowledgement', 'acknowledgements', 'appendix', 'keywords',
        'methodology', 'discussion', 'results', 'preface', 'index',
        'conclusion', 'conclusions', 'summary', 'recommendations',
        'निष्कर्ष', 'सारांश', 'अनुशंसाएँ', 'संदर्भ', 'ग्रंथसूची',
    }
    if text.lower().strip('.').strip() in special_sections and wc <= 3:
        return 'section_heading'

    # if re.match(r'^\d+\.\d+\.\d+', text) and (is_bold or text == text.upper()):
    #     return 'subheading'

    # if re.match(r'^\d+\.\d+\.?\s', text) and (is_bold or text == text.upper()):
    #     return 'section_heading'


    if re.match(r'^\d+\.\d+\.\d+', text) and (is_bold or wc <= 20):
        return 'subheading'

    if re.match(r'^\d+\.\d+\.?\s', text) and (is_bold or wc <= 20):
        return 'section_heading'


    if re.match(r'^\d+\.?\s+\S', text) and is_bold:
        return 'section_heading'

    # Devanagari numeral heading (१. heading, २. heading)
    if re.match(r'^[१-९][०-९]*\.?\s+\S', text) and is_bold:
        return 'section_heading'

    if re.match(r'^[A-Z]\.\s', text) and is_bold:
        return 'section_heading'

    if text.endswith(':') and is_bold and wc <= 20:
        return 'subheading_colon'

    if text.isupper() and is_bold:
        return 'section_heading'

    if text.isupper() and 2 <= wc <= 6:
        return 'section_heading'

    if is_bold and wc <= 15:
        return 'subheading'

    return 'body'


# ═══════════════════════════
# THESIS BODY FORMATTING
# ═══════════════════════════

def format_thesis_body(doc, opts, font_name):
    black         = RGBColor(0, 0, 0)
    krutidev_mode = is_krutidev(font_name)

    if krutidev_mode:
        base_size        = 14.0
        ch_heading_size  = 18.0
        ch_title_size    = 18.0
        sec_heading_size = 17.0
        sub_heading_size = 15.0
    else:
        base_size        = 12.0
        ch_heading_size  = 16.0
        ch_title_size    = 16.0
        sec_heading_size = 14.0
        sub_heading_size = 12.0

    line_spacing = float(opts.get('line_spacing', 1.15))

    if opts.get('font_size'):
        base_size = float(opts['font_size'])

    heading_font = font_name if krutidev_mode else 'Times New Roman'

    def apply_caps_upper(para):
        if krutidev_mode:
            return
        for run in para.runs:
            if run.text and run.text.strip():
                run.text = run.text.upper()

    def set_widow_orphan(para):
        pPr = para._p.get_or_add_pPr()
        wc  = pPr.find(qn('w:widowControl'))
        if wc is None:
            wc = OxmlElement('w:widowControl')
            pPr.append(wc)
        wc.set(qn('w:val'), '1')

    def set_keep_next(para):
        pPr = para._p.get_or_add_pPr()
        kn  = pPr.find(qn('w:keepNext'))
        if kn is None:
            kn = OxmlElement('w:keepNext')
            pPr.append(kn)
        kn.set(qn('w:val'), '1')

    i          = 0
    prev_etype = None

    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        text = para.text.strip()

        if has_drawing(para):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pPr_d = para._p.get_or_add_pPr()
            for jc_el in pPr_d.findall(qn('w:jc')):
                pPr_d.remove(jc_el)
            jc_draw = OxmlElement('w:jc')
            jc_draw.set(qn('w:val'), 'center')
            pPr_d.append(jc_draw)
            sp_d = pPr_d.find(qn('w:spacing'))
            if sp_d is None:
                sp_d = OxmlElement('w:spacing')
                pPr_d.append(sp_d)
            sp_d.set(qn('w:before'), '100')
            sp_d.set(qn('w:after'), '100')
            sp_d.set(qn('w:beforeAutospacing'), '0')
            sp_d.set(qn('w:afterAutospacing'), '0')
            i += 1
            continue

        if not text:
            pPr_e = para._p.get_or_add_pPr()
            sp_e  = pPr_e.find(qn('w:spacing'))
            if sp_e is None:
                sp_e = OxmlElement('w:spacing')
                pPr_e.append(sp_e)
            sp_e.set(qn('w:before'), '0')
            sp_e.set(qn('w:after'),  '0')
            sp_e.set(qn('w:beforeAutospacing'), '0')
            sp_e.set(qn('w:afterAutospacing'),  '0')
            sp_e.set(qn('w:line'), '240')
            sp_e.set(qn('w:lineRule'), 'auto')
            i += 1
            continue

        etype = detect_thesis_structure(para, i, doc)

        if etype in ('empty', 'drawing'):
            i += 1
            continue

        if etype == 'figure_caption':
            apply_para_formatting(para, etype, font_name,
                font_size_pt=12.0,
                bold=True, color=black,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                space_before_pt=4, space_after_pt=4,
                line_spacing=1.0)
            set_widow_orphan(para)
            prev_etype = etype
            i += 1
            continue

        # --- Spacing defaults ---
        space_after  = 5.0
        space_before = 10.0

        next_etype = None
        if i < len(doc.paragraphs) - 1:
            next_para = doc.paragraphs[i + 1]
            if next_para.text.strip() and not has_drawing(next_para):
                next_etype = detect_thesis_structure(next_para, i + 1, doc)

        # Heading after heading → tight 5pt gap
        # if etype in ['section_heading', 'subheading', 'subheading_colon']:
        #     if prev_etype in ['chapter_heading', 'chapter_title', 'section_heading', 'subheading', 'subheading_colon']:
        #         space_before = 5.0
        #     else:
        #         space_before = 10.0
        #     space_after = 3.0


        # Heading after heading → 10pt gap, tight only after chapter heading
        if etype in ['section_heading', 'subheading', 'subheading_colon']:
            if prev_etype in ['chapter_heading', 'chapter_title']:
                space_before = 5.0
            elif prev_etype in ['section_heading', 'subheading', 'subheading_colon']:
                space_before = 10.0
            else:
                space_before = 10.0
            space_after = 3.0



        # Body/bullet after body/bullet → no extra gap
        if etype in ['body', 'bullet']:
            space_before = 0.0
            space_after  = 5.0

        # Body before a heading → small gap
        if etype in ['body', 'bullet'] and next_etype in ['section_heading', 'subheading', 'subheading_colon', 'chapter_heading']:
            space_after = 2.0

        if etype == 'chapter_heading':
            if ':' in text and re.match(r'^chapter\s+\S+', text, re.IGNORECASE):
                parts         = text.split(':', 1)
                chapter_label = parts[0].strip()
                chapter_title = parts[1].strip()

                label_text = chapter_label.upper() if not krutidev_mode else chapter_label
                for r in list(para.runs):
                    r._r.getparent().remove(r._r)
                r_new = para.add_run(label_text)
                r_new.bold = True
                set_font_properly(r_new, heading_font)
                apply_para_formatting(para, etype, heading_font,
                    font_size_pt=ch_heading_size, bold=True, color=black,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before_pt=15, space_after_pt=0,
                    line_spacing=line_spacing)
                set_widow_orphan(para)

                title_para = doc.add_paragraph()
                para._p.addnext(title_para._p)
                title_run = title_para.add_run(chapter_title.upper() if not krutidev_mode else chapter_title)
                title_run.bold = True
                set_font_properly(title_run, heading_font)
                apply_para_formatting(title_para, 'chapter_title', heading_font,
                    font_size_pt=ch_title_size, bold=True, color=black,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before_pt=0, space_after_pt=10,
                    line_spacing=line_spacing)
                set_widow_orphan(title_para)
                i += 2
                prev_etype = 'chapter_heading'
                continue
            else:
                next_is_title = False
                if i + 1 < len(doc.paragraphs):
                    nxt = doc.paragraphs[i + 1]
                    nxt_text = nxt.text.strip()
                    if nxt_text and not has_drawing(nxt):
                        nxt_etype = detect_thesis_structure(nxt, i + 1, doc)
                        if nxt_etype == 'chapter_heading' and not CHAPTER_HEADING_RE.match(nxt_text):
                            next_is_title = True

                if not krutidev_mode:
                    apply_caps_upper(para)
                apply_para_formatting(para, etype, heading_font,
                    font_size_pt=ch_heading_size, bold=True, color=black,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before_pt=15, space_after_pt=0 if next_is_title else 10,
                    line_spacing=line_spacing)
                set_widow_orphan(para)
                set_keep_next(para)

                if next_is_title and i + 1 < len(doc.paragraphs):
                    title_para = doc.paragraphs[i + 1]
                    if not krutidev_mode:
                        for run in title_para.runs:
                            if run.text and run.text.strip():
                                run.text = run.text.upper()
                    apply_para_formatting(title_para, 'chapter_title', heading_font,
                        font_size_pt=ch_title_size, bold=True, color=black,
                        align=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before_pt=0, space_after_pt=10,
                        line_spacing=line_spacing)
                    set_widow_orphan(title_para)
                    prev_etype = 'chapter_title'
                    i += 2
                    continue

        elif etype == 'section_heading':
            if not krutidev_mode:
                apply_caps_upper(para)
            apply_para_formatting(para, etype, heading_font,
                font_size_pt=sec_heading_size,
                bold=True,
                color=black,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_before_pt=space_before, space_after_pt=3.0,
                left_indent=0.0, first_indent=0.0,
                line_spacing=line_spacing)
            set_widow_orphan(para)
            set_keep_next(para)

        elif etype == 'subheading':
            if not krutidev_mode:
                apply_caps_upper(para)
            apply_para_formatting(para, etype, heading_font,
                font_size_pt=sub_heading_size,
                bold=True,
                color=black,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_before_pt=space_before, space_after_pt=3.0,
                left_indent=0.0, first_indent=0.0,
                line_spacing=line_spacing)
            set_widow_orphan(para)
            set_keep_next(para)

        elif etype == 'subheading_colon':
            if not krutidev_mode:
                apply_caps_upper(para)
            apply_para_formatting(para, 'subheading', heading_font,
                font_size_pt=sub_heading_size,
                bold=True,
                color=black,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_before_pt=space_before, space_after_pt=3.0,
                left_indent=0.0, first_indent=0.0,
                line_spacing=line_spacing)
            set_widow_orphan(para)

        elif etype == 'bullet':
            is_bold_para = is_all_bold(para)
            apply_para_formatting(para, etype, font_name,
                font_size_pt=base_size, bold=is_bold_para, color=black,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before_pt=0, space_after_pt=space_after,
                left_indent=0.25, first_indent=-0.25,
                line_spacing=line_spacing)
            set_widow_orphan(para)

        else:  # body
            words_in_para = text.split()
            should_justify = len(words_in_para) >= 12 and len(text) >= 100 and not text.endswith(('?', ':', '!', ';'))
            final_align = WD_ALIGN_PARAGRAPH.JUSTIFY if should_justify else WD_ALIGN_PARAGRAPH.LEFT

            apply_para_formatting(para, etype, font_name,
                font_size_pt=base_size, bold=False, color=black,
                align=final_align,
                space_before_pt=0, space_after_pt=5.0,
                left_indent=0.0, first_indent=0.0,
                line_spacing=line_spacing)

            if should_justify:
                pPr = para._p.get_or_add_pPr()
                for jc in pPr.findall(qn('w:jc')):
                    pPr.remove(jc)
                jc_el = OxmlElement('w:jc')
                jc_el.set(qn('w:val'), 'both')
                pPr.append(jc_el)
            set_widow_orphan(para)

        prev_etype = etype
        i += 1

    format_table_cells(doc, font_name, base_size, line_spacing, black)

    # Add 5pt space after every table so next paragraph isn't cramped
    for table in doc.tables:
        tbl = table._tbl
        next_sib = tbl.getnext()
        if next_sib is not None and next_sib.tag == qn('w:p'):
            from docx.oxml import OxmlElement as _OE
            pPr = next_sib.find(qn('w:pPr'))
            if pPr is None:
                pPr = _OE('w:pPr')
                next_sib.insert(0, pPr)
            sp = pPr.find(qn('w:spacing'))
            if sp is None:
                sp = _OE('w:spacing')
                pPr.append(sp)
            existing_before = sp.get(qn('w:before'))
            if existing_before is None or int(existing_before) < 100:
                sp.set(qn('w:before'), '100')
                sp.set(qn('w:beforeAutospacing'), '0')
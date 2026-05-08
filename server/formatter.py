import sys
import json
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from utils import (
    preprocess_document, is_krutidev, has_unicode_hindi,
    unicode_to_krutidev, has_drawing, set_font_properly,
    center_all_tables, add_fld_char, add_instr_text,
    convert_doc_runs
)
from book import insert_title_page, format_book_body
from thesis import insert_thesis_title_page, format_thesis_body
from letter import insert_letter_header, has_existing_letter_header, format_letter_body
from research import insert_research_title_page, format_research_body

PAGE_SIZE_MAP = {
    'A4':     (Mm(210), Mm(297)),
    'A5':     (Mm(148), Mm(210)),
    'A3':     (Mm(297), Mm(420)),
    'Letter': (Mm(215.9), Mm(279.4)),
    'Legal':  (Mm(215.9), Mm(355.6)),
}


def format_document(input_file, output_file, opts, doc_type='book'):
    doc       = Document(input_file)
    font_name = opts.get('font_style') or 'Garamond'
    black     = RGBColor(0, 0, 0)
    gray      = RGBColor(100, 100, 100)

    # 1. Pre-clean
    preprocess_document(doc)

    # 2. Page Size & Margins
    page_size_key = opts.get('page_size', 'A4')
    page_w, page_h = PAGE_SIZE_MAP.get(page_size_key, PAGE_SIZE_MAP['A4'])
    for section in doc.sections:
        section.page_width  = page_w
        section.page_height = page_h
        if doc_type == 'thesis':
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin   = Inches(1.5)
            section.right_margin  = Inches(1.0)
        elif doc_type == 'letter':
            if opts.get('page_size') and opts.get('page_size') != 'A4':
                section.top_margin    = Inches(0.8)
                section.bottom_margin = Inches(0.8)
                section.left_margin   = Inches(1.2)
                section.right_margin  = Inches(1.0)
        else:
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin   = Inches(1.0)
            section.right_margin  = Inches(1.0)

    # 2b. Center all tables
    center_all_tables(doc)

    # 3. Title page — by doc_type
    if doc_type == 'thesis':
        insert_thesis_title_page(doc, opts, font_name)
    elif doc_type == 'letter':
        has_user_header = opts.get('org_name') or opts.get('subject')
        if has_user_header and not has_existing_letter_header(doc):
            insert_letter_header(doc, opts, font_name)
    elif doc_type == 'research':
        insert_research_title_page(doc, opts, font_name)
    else:
        insert_title_page(doc, opts, font_name)

    # 4. Body formatting — by doc_type
    if doc_type == 'thesis':
        format_thesis_body(doc, opts, font_name)
    elif doc_type == 'letter':
        format_letter_body(doc, opts, font_name)
    elif doc_type == 'research':
        format_research_body(doc, opts, font_name)
    else:
        # book / research paper
        format_book_body(doc, opts, font_name)

    # 5. Headers & Footers
    header_text  = opts.get('header', '').strip()
    footer_text  = opts.get('footer', '').strip()
    page_numbers = opts.get('page_numbers', False)
    page_num_pos = opts.get('page_number_position', 'center')
    start_page   = opts.get('start_page_number', 1)
    try:
        start_page = int(start_page)
    except (ValueError, TypeError):
        start_page = 1

    ALIGN_MAP = {
        'left':   WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right':  WD_ALIGN_PARAGRAPH.RIGHT,
    }
    num_align = ALIGN_MAP.get(page_num_pos, WD_ALIGN_PARAGRAPH.CENTER)

    if doc_type == 'thesis':
        page_numbers = True
        num_align    = WD_ALIGN_PARAGRAPH.CENTER

    for section in doc.sections:
        section.header_distance = Mm(12.5)
        section.footer_distance = Mm(12.5)
        if page_numbers and start_page != 1:
            sectPr    = section._sectPr
            pgNumType = sectPr.find(qn('w:pgNumType'))
            if pgNumType is None:
                pgNumType = OxmlElement('w:pgNumType')
                sectPr.append(pgNumType)
            pgNumType.set(qn('w:start'), str(start_page))

        if header_text:
            hdr_para = section.header.paragraphs[0]
            hdr_para.clear()
            hdr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = hdr_para.add_run(header_text)
            set_font_properly(r, font_name)
            r.font.size = Pt(9)
            r.font.color.rgb = gray

        ftr = section.footer
        for fp in ftr.paragraphs:
            fp.clear()

        if footer_text:
            ft_para = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
            ft_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = ft_para.add_run(footer_text)
            set_font_properly(r, font_name)
            r.font.size = Pt(9)
            r.font.color.rgb = gray

        if page_numbers:
            pn_para = ftr.add_paragraph()
            pn_para.alignment = num_align
            r1 = pn_para.add_run()
            set_font_properly(r1, font_name)
            r1.font.size = Pt(10 if doc_type == 'thesis' else 9)
            r1.font.color.rgb = RGBColor(0, 0, 0) if doc_type == 'thesis' else gray
            add_fld_char(r1, 'begin')
            add_instr_text(r1, ' PAGE \\* ARABIC ')
            add_fld_char(r1, 'end')

    # 6. Kruti Dev Unicode → ASCII conversion (must run AFTER all formatting
    #    so runs added by title-page and body formatters are also converted)
    convert_doc_runs(doc, font_name)

    doc.save(output_file)


if __name__ == '__main__':
    in_p   = sys.argv[1]
    out_p  = sys.argv[2]
    type_d = sys.argv[3]
    opts_f = sys.argv[4]

    options = {}
    if os.path.exists(opts_f) and os.path.getsize(opts_f) > 0:
        with open(opts_f, 'r', encoding='utf-8') as f:
            options = json.load(f)

    format_document(in_p, out_p, options, doc_type=type_d)
    print(f'Success: {out_p}')


    

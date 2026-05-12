import re
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ═══════════════════════════
# DRAWING / IMAGE DETECTION
# ═══════════════════════════

WP_NS  = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
MC_NS  = 'http://schemas.openxmlformats.org/markup-compatibility/2006'
W_NS   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def has_drawing(para):
    """Return True if paragraph contains any image/chart/drawing/object element."""
    p = para._p
    for tag in [qn('w:drawing'), qn('w:pict'), qn('w:object')]:
        if p.find('.//' + tag) is not None:
            return True
    if p.find(f'{{{MC_NS}}}AlternateContent') is not None:
        return True
    return False


# ═══════════════════════════
# PRE-CLEANING
# ═══════════════════════════

def clean_runs_in_para(para):
    for run in para.runs:
        cleaned = run.text.replace('\t', '').replace('\n', ' ')
        cleaned = re.sub(r' {2,}', ' ', cleaned)
        run.text = cleaned


def remove_proof_errors(para):
    """Remove <w:proofErr> elements that cause word splits during run iteration."""
    p = para._p
    for proof_err in p.findall(qn('w:proofErr')):
        p.remove(proof_err)


def run_has_drawing(run):
    """Return True if this run contains any drawing/image element."""
    r = run._r
    if r.find(qn('w:drawing')) is not None:
        return True
    if r.find(qn('w:pict')) is not None:
        return True
    if r.find(qn('w:object')) is not None:
        return True
    return False


def merge_runs_in_para(para):
    """Merge adjacent runs with identical formatting. NEVER merges runs that contain drawings."""
    if len(para.runs) <= 1:
        return

    i = 0
    while i < len(para.runs) - 1:
        r1 = para.runs[i]
        r2 = para.runs[i + 1]

        if run_has_drawing(r1) or run_has_drawing(r2):
            i += 1
            continue

        def fmt_sig(run):
            rPr = run._element.find(qn('w:rPr'))
            if rPr is None:
                return ('', None, None, None, None)
            bold   = rPr.find(qn('w:b'))
            italic = rPr.find(qn('w:i'))
            sz     = rPr.find(qn('w:sz'))
            color  = rPr.find(qn('w:color'))
            rFonts = rPr.find(qn('w:rFonts'))
            b_val     = None if bold   is None else bold.get(qn('w:val'), 'true')
            i_val     = None if italic is None else italic.get(qn('w:val'), 'true')
            sz_val    = None if sz     is None else sz.get(qn('w:val'))
            color_val = None if color  is None else color.get(qn('w:val'))
            font_val  = None if rFonts is None else rFonts.get(qn('w:ascii'))
            return (b_val, i_val, sz_val, color_val, font_val)

        if fmt_sig(r1) == fmt_sig(r2):
            r1.text = (r1.text or '') + (r2.text or '')
            r2._r.getparent().remove(r2._r)
        else:
            i += 1


def is_all_bold(para):
    runs = [r for r in para.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)


def is_bullet_para(para):
    """True if paragraph has a numbering/bullet list marker in XML."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return False
    return pPr.find(qn('w:numPr')) is not None


def apply_bold_before_colon(para, font_name, krutidev_mode):
    """If para has 'Label: rest' pattern, make text before ':' bold."""
    text = para.text
    colon_idx = text.find(':')
    if colon_idx <= 0 or colon_idx > 80:
        return

    label = text[:colon_idx + 1]
    rest  = text[colon_idx + 1:]

    first_run = para.runs[0] if para.runs else None
    size_pt = None
    if first_run:
        size_pt = first_run.font.size

    for run in list(para.runs):
        run._r.getparent().remove(run._r)

    r_bold = para.add_run(label)
    r_bold.bold = True
    if not krutidev_mode:
        set_font_properly(r_bold, font_name)
        if size_pt:
            r_bold.font.size = size_pt
        r_bold.font.color.rgb = RGBColor(0, 0, 0)

    if rest:
        r_rest = para.add_run(rest)
        r_rest.bold = False
        if not krutidev_mode:
            set_font_properly(r_rest, font_name)
            if size_pt:
                r_rest.font.size = size_pt
            r_rest.font.color.rgb = RGBColor(0, 0, 0)


def merge_split_paragraphs(doc):
    pass  # Disabled — causes duplicate rendering with mixed fonts


def clear_pPr_sz(para):
    """Remove font size override from paragraph-level rPr."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        return
    for tag in [qn('w:sz'), qn('w:szCs')]:
        el = rPr.find(tag)
        if el is not None:
            rPr.remove(el)


def set_pPr_sz(para, half_pts):
    """Set font size at paragraph-level rPr."""
    pPr = para._p.get_or_add_pPr()
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)
    for tag_name in ['w:sz', 'w:szCs']:
        el = rPr.find(qn(tag_name))
        if el is None:
            el = OxmlElement(tag_name)
            rPr.append(el)
        el.set(qn('w:val'), str(half_pts))


def clear_para_indent(para):
    """Remove all left/first-line indent from paragraph XML."""
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is not None:
        pPr.remove(ind)
    para.paragraph_format.left_indent        = None
    para.paragraph_format.first_line_indent  = None


def preprocess_document(doc):
    for para in doc.paragraphs:
        if has_drawing(para):
            continue
        remove_proof_errors(para)
        clean_runs_in_para(para)
        merge_runs_in_para(para)
    merge_split_paragraphs(doc)


# ═══════════════════════════
# HELPERS
# ═══════════════════════════

KRUTIDEV_FONTS = {'Kruti Dev 010', 'Kruti Dev 011', 'Krutidev010', 'Krutidev011',
                  'KrutiDev010', 'KrutiDev011', 'Kruti Dev 010', 'Kruti Dev 011'}

# ═══════════════════════════
# SHARED DETECTION CONSTANTS
# ═══════════════════════════

# Hindi + English chapter/unit words, ASCII + Devanagari numerals
CHAPTER_HEADING_RE = re.compile(
    r'^(chapter|unit|part|lesson|अध्याय|इकाई|भाग|पाठ)'
    r'\s*[-–—]?\s*(\d+|[ivxlcdmIVXLCDM]+|[०-९]+)\b',
    re.IGNORECASE
)

# Matches headings where chapter word appears ANYWHERE in first 3 words
# e.g. 'षष्ठम अध्याय: ...' or 'तृतीय अध्याय: ...'
CHAPTER_HEADING_LOOSE_RE = re.compile(
    r'^[\w\u0900-\u097F]+\s+(अध्याय|chapter|unit|part|lesson|इकाई|भाग|पाठ)'
    r'\s*[-:–—]',
    re.IGNORECASE
)


def inject_heading_number(para, sec, sub=None, krutidev_mode=False):
    """
    Inject section number prefix into heading para.
    Skipped in krutidev_mode because Hindi documents already carry
    Devanagari numerals in the text — adding an ASCII counter would
    produce garbage like '0१' or '1२'.
    Also skipped if the paragraph already starts with a digit.
    """
    if krutidev_mode:
        return
    text = para.text.strip()
    # Skip if already starts with ASCII digit, Devanagari digit,
    # OR KrutiDev-encoded digit chars (! @ # $ % ^ & * ( for १-९)
    KRUTIDEV_DIGIT_CHARS = '!@#$%^&*('
    if re.match(r'^[\d०-९]', text) or (text and text[0] in KRUTIDEV_DIGIT_CHARS):
        return
    prefix = f"{sec}. " if sub is None else f"{sec}.{sub} "
    if para.runs:
        para.runs[0].text = prefix + para.runs[0].text.lstrip()
    else:
        para.add_run(prefix)

FONT_NAME_MAP = {
    'Krutidev010': 'Kruti Dev 010',
    'Krutidev011': 'Kruti Dev 011',
    'Mangal':      'Mangal',
    'Kokila':      'Kokila',
    'Utsaah':      'Utsaah',
    'Aparajita':   'Aparajita',
    'Nirmala UI':  'Nirmala UI',
}


# ═══════════════════════════
# HINDI CONVERSION
# ═══════════════════════════

def unicode_to_krutidev(text):
    """
    Convert Unicode Devanagari text to Kruti Dev 010 ASCII encoding.
    Uses syllable-based reordering + comprehensive conjunct table.

    Fixes over previous version:
    - ह्र conjunct corrected: 'gz' -> 'g`'
    - Added missing conjuncts: व्र, फ्र, क्त, त्व, स्व, न्त, न्ध,
      ल्ल, न्न, ह्न, ह्म, क्क, त्न, स्न, स्म
    - Removed incorrect न्द conjunct (caused हिन्दी to break)
    - Extended syll_pattern to cover full Devanagari consonant range
    - Added nukta variants to C map
    """
    if not text:
        return ""
    if not re.search(r'[\u0900-\u097F]', text):
        return text

    import unicodedata
    text = unicodedata.normalize('NFC', text)

    # Normalize special punctuation that may appear inside Hindi runs
    # (en-dash, em-dash, curly quotes, NBSP, etc.)
    SPECIAL_NORM = [
        ('–', '-'), ('—', '--'), ('…', '...'),
        (' ', ' '), ('•', '*'),  ('·', '*'),
        ('“', '"'), ('”', '"'),
        ('‘', "'"), ('’', "'"),
        ('―', '-'), ('‒', '-'),
    ]
    for uni, repl in SPECIAL_NORM:
        text = text.replace(uni, repl)

    # Handle consonant+nukta pairs (NFC does NOT compose these in Python)
    NUKTA = '\u093c'
    NUKTA_PAIRS = [
        ('\u0921' + NUKTA, 'M'),   # ड़ -> M
        ('\u0922' + NUKTA, '<'),   # ढ़ -> <
        ('\u091c' + NUKTA, 't'),   # ज़ -> t
        ('\u092b' + NUKTA, 'Q'),   # फ़ -> Q
        ('\u0915' + NUKTA, 'd'),   # क़ -> d
        ('\u0916' + NUKTA, '[k'),  # ख़ -> [k
        ('\u0917' + NUKTA, 'x'),   # ग़ -> x
        ('\u092f' + NUKTA, ';'),   # य़ -> ;
        (NUKTA, ''),                # strip any remaining nukta
    ]
    for pair, kd in NUKTA_PAIRS:
        text = text.replace(pair, kd)

    halant = '\u094D'
    reph = '\u0930' + halant

    syll_pattern = (
        r'((?:[\u0900-\u0939\u0958-\u0961\u0972]\u094D)*'
        r'[\u0900-\u0939\u0958-\u0961\u0972]'
        r'[\u093E-\u094D\u0901-\u0903\u094E-\u094F\u0955-\u0957]*)'
    )

    def process_syllable(m):
        s = m.group(0)
        has_reph = False
        if s.startswith(reph) and len(s) > 2:
            has_reph = True
            s = s[2:]
        has_short_i = '\u093f' in s
        if has_short_i:
            s = s.replace('\u093f', '')
        res = ('f' if has_short_i else '') + s
        if has_reph:
            idx = next((i for i, c in enumerate(res) if c in '\u0902\u0903\u0901'), -1)
            if idx >= 0:
                res = res[:idx] + 'Z' + res[idx:]
            else:
                res += 'Z'
        return res

    text = re.sub(syll_pattern, process_syllable, text)

    C = {
        '\u0905': 'v',   '\u0906': 'vk',  '\u0907': 'b',   '\u0908': 'bZ',
        '\u0909': 'm',   '\u090a': '\xc5',  '\u090b': '_',
        '\u090f': ',',   '\u0910': ',s',  '\u0913': 'vks', '\u0914': 'vkS',
        '\u0911': 'vkW',
        '\u093e': 'k',   '\u093f': 'f',   '\u0940': 'h',   '\u0941': 'q',
        '\u0942': 'w',   '\u0943': '`',   '\u0947': 's',   '\u0948': 'S',
        '\u094b': 'ks',  '\u094c': 'kS',  '\u094e': 'W',   '\u094f': 'V',  '\u0949': 'W',
        '\u0902': 'a',   '\u0903': '%',   '\u0901': 'i',
        '\u0915': 'd',   '\u0916': '[k',  '\u0917': 'x',   '\u0918': '?k',  '\u0919': 'M~',
        '\u091a': 'p',   '\u091b': 'N',   '\u091c': 't',   '\u091d': '>k',  '\u091e': '\xa5',
        '\u091f': 'V',   '\u0920': 'B',   '\u0921': 'M',   '\u0922': '<',   '\u0923': '.k',
        '\u0924': 'r',   '\u0925': 'Fk',  '\u0926': 'n',   '\u0927': '/k',  '\u0928': 'u',
        '\u092a': 'i',   '\u092b': 'Q',   '\u092c': 'c',   '\u092d': 'Hk',  '\u092e': 'e',
        '\u092f': ';',   '\u0930': 'j',   '\u0932': 'y',   '\u0935': 'o',
        '\u0936': "'k",  '\u0937': '"k',  '\u0938': 'l',   '\u0939': 'g',
        '\u0958': 'd',   '\u0959': '[k',  '\u095a': 'x',   '\u095b': 't',
        '\u095c': 'M',   '\u095d': '<',   '\u095e': 'Q',   '\u095f': ';',
        '\u0960': '__',  '\u0961': '_',   '\u0950': 'ks',  '\u093d': "'",
        '\u0964': 'A',   '\u0965': 'AA',
        '\u0966': ')',   '\u0967': '!',   '\u0968': '@',   '\u0969': '#',   '\u096a': '$',
        '\u096b': '%',   '\u096c': '^',   '\u096d': '&',   '\u096e': '*',   '\u096f': '(',
    }

    HALF = {
        '\u0915': 'D',   '\u0916': '[',   '\u0917': 'X',   '\u0918': '?',
        '\u091a': 'P',   '\u091c': 'T',   '\u091d': '>',
        '\u091f': 'V~',  '\u0920': 'B~',  '\u0921': 'M~',  '\u0922': '<~',  '\u0923': '.k~',
        '\u0924': 'R',   '\u0925': 'F',   '\u0926': 'n~',  '\u0927': '/',   '\u0928': 'U',
        '\u092a': 'I',   '\u092b': 'Q~',  '\u092c': 'C',   '\u092d': 'H',   '\u092e': 'E',
        '\u092f': 'Y',   '\u0930': 'z',   '\u0932': 'y~',  '\u0935': 'O',
        '\u0936': "'",    '\u0937': '"',   '\u0938': 'L',   '\u0939': 'g~',
    }

    # ORDER MATTERS: longer conjuncts before shorter
    CONJUNCTS = [
        # 3-consonant clusters first
        ('\u0938\u094d\u0924\u094d\u0930', 'L='),    # स्त्र
        # 2-consonant clusters
        ('\u0915\u094d\u0937', '{k'),     # क्ष
        ('\u0924\u094d\u0930', '='),      # त्र
        ('\u091c\u094d\u091e', 'K'),      # ज्ञ
        ('\u0936\u094d\u0930', "'J"),     # श्र
        ('\u092a\u094d\u0930', 'iz'),     # प्र
        ('\u0917\u094d\u0930', 'xz'),     # ग्र
        ('\u0915\u094d\u0930', 'dz'),     # क्र
        ('\u092c\u094d\u0930', 'cz'),     # ब्र
        ('\u092e\u094d\u0930', 'ez'),     # म्र
        ('\u0926\u094d\u0930', 'nz'),     # द्र
        ('\u0927\u094d\u0930', '/z'),     # ध्र
        ('\u092d\u094d\u0930', 'Hkz'),    # भ्र
        ('\u0939\u094d\u0930', 'g`'),     # ह्र  *** FIXED (was gz) ***
        ('\u0935\u094d\u0930', 'oz'),     # व्र  *** NEW ***
        ('\u092b\u094d\u0930', 'Qz'),     # फ्र  *** NEW ***
        ('\u091f\u094d\u0930', 'Vz'),     # ट्र
        ('\u0921\u094d\u0930', 'Mz'),     # ड्र
        ('\u0926\u094d\u092f', '|'),      # द्य
        ('\u0926\u094d\u0927', '/~/k'),   # द्ध
        ('\u0926\u094d\u0935', 'n~o'),    # द्व
        ('\u0924\u094d\u0924', '\xd9k'),  # त्त
        ('\u0915\u094d\u0924', 'DR'),     # क्त  *** NEW ***
        ('\u0924\u094d\u0935', 'Ro'),     # त्व  *** NEW ***
        ('\u0938\u094d\u0935', 'Lo'),     # स्व  *** NEW ***
        ('\u0928\u094d\u0924', 'UR'),     # न्त  *** NEW ***
        ('\u0928\u094d\u0927', 'U/k'),    # न्ध  *** NEW ***
        ('\u0932\u094d\u0932', 'y~y'),    # ल्ल  *** NEW ***
        ('\u0928\u094d\u0928', 'UU'),     # न्न  *** NEW ***
        ('\u0939\u094d\u0928', 'g~u'),    # ह्न  *** NEW ***
        ('\u0939\u094d\u092e', 'g~e'),    # ह्म  *** NEW ***
        ('\u0915\u094d\u0915', 'DD'),     # क्क  *** NEW ***
        ('\u0924\u094d\u0928', 'Ru'),     # त्न  *** NEW ***
        ('\u0938\u094d\u0928', 'Lu'),     # स्न  *** NEW ***
        ('\u0938\u094d\u092e', 'Le'),     # स्म  *** NEW ***
        # ('\u0930\u0942', 'tw'),
        # ('\u0930\u0941', 'rq'),
    ]

    for uni, kd in CONJUNCTS:
        text = text.replace(uni, kd)

    for uni, kd in HALF.items():
        text = text.replace(uni + halant, kd)

    # Strip any remaining unconsumed halant (trailing virama)
    result = ''.join(C.get(char, char) for char in text)
    result = result.replace('\u094d', '')
    return result

def is_krutidev(font_name):
    return font_name and any(k.lower() in font_name.lower() for k in ['kruti', 'krutidev'])


def has_unicode_hindi(text):
    return bool(re.search(r'[\u0900-\u097F]', text))


def set_font_properly(run, font_name, size_pt=None):
    formal_name = FONT_NAME_MAP.get(font_name, font_name)
    run.font.name = formal_name

    r    = run._element
    rPr  = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()

    if is_krutidev(formal_name):
        # KrutiDev is a legacy ASCII font — treat it like any ASCII font.
        # DO NOT set w:hint (removing it lets Word use ascii/hAnsi path).
        # DO NOT set eastAsia (causes Word to switch to CJK/Indic font).
        # Set only ascii + hAnsi so Word uses KrutiDev for all Latin runs.
        hint_attr = qn('w:hint')
        if rFonts.get(hint_attr):
            del rFonts.attrib[hint_attr]
        rFonts.set(qn('w:ascii'), formal_name)
        rFonts.set(qn('w:hAnsi'), formal_name)
        # Remove eastAsia — causes fallback to system Devanagari font
        ea_attr = qn('w:eastAsia')
        if rFonts.get(ea_attr):
            del rFonts.attrib[ea_attr]
        # Remove cs — KrutiDev has no complex-script glyphs
        cs_attr = qn('w:cs')
        if rFonts.get(cs_attr):
            del rFonts.attrib[cs_attr]
        # Remove theme font overrides
        for theme_attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme']:
            ta = qn(theme_attr)
            if rFonts.get(ta):
                del rFonts.attrib[ta]
        # Remove bidi/rtl — KrutiDev is LTR ASCII
        for cs_tag in ['w:rtl', 'w:cs', 'w:bidi']:
            el = rPr.find(qn(cs_tag))
            if el is not None:
                rPr.remove(el)
        # Language: x-none disables spell-check without triggering font substitution.
        # en-US is wrong — it causes Word to suggest "fix" the font on Hindi text.
        lang = rPr.find(qn('w:lang'))
        if lang is None:
            lang = OxmlElement('w:lang')
            rPr.append(lang)
        lang.set(qn('w:val'),   'x-none')
        lang.set(qn('w:ascii'), 'x-none')
        lang.set(qn('w:hAnsi'), 'x-none')
        # Remove bidi lang attr if present
        bidi_lang = qn('w:bidi')
        if lang.get(bidi_lang):
            del lang.attrib[bidi_lang]
        # NoProof: prevents spell-check red underlines
        no_proof = rPr.find(qn('w:noProof'))
        if no_proof is None:
            no_proof = OxmlElement('w:noProof')
            rPr.append(no_proof)
    else:
        rFonts.set(qn('w:hint'), 'complex')
        for attr in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
            rFonts.set(qn(f'w:{attr}'), formal_name)
            
        lang = rPr.find(qn('w:lang'))
        if lang is None:
            lang = OxmlElement('w:lang')
            rPr.append(lang)
        lang.set(qn('w:val'), 'hi-IN')
        lang.set(qn('w:cs'),  'hi-IN')

    if size_pt:
        run.font.size = Pt(size_pt)
        if not is_krutidev(formal_name):
            sz_cs = rPr.find(qn('w:szCs'))
            if sz_cs is None:
                sz_cs = OxmlElement('w:szCs')
                rPr.append(sz_cs)
            sz_cs.set(qn('w:val'), str(int(size_pt * 2)))


def set_para_font(para, font_name):
    """Set font at paragraph-level rPr."""
    formal_name = FONT_NAME_MAP.get(font_name, font_name)
    pPr  = para._p.get_or_add_pPr()
    
    # CT_PPr does not have get_or_add_rPr, use manual creation
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)

    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)

    if is_krutidev(formal_name):
        # Same logic as set_font_properly: no hint, no eastAsia, no cs
        hint_attr = qn('w:hint')
        if rFonts.get(hint_attr):
            del rFonts.attrib[hint_attr]
        rFonts.set(qn('w:ascii'), formal_name)
        rFonts.set(qn('w:hAnsi'), formal_name)
        ea_attr = qn('w:eastAsia')
        if rFonts.get(ea_attr):
            del rFonts.attrib[ea_attr]
        cs_attr = qn('w:cs')
        if rFonts.get(cs_attr):
            del rFonts.attrib[cs_attr]
        for theme_attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme']:
            ta = qn(theme_attr)
            if rFonts.get(ta):
                del rFonts.attrib[ta]
        for cs_tag in ['w:rtl', 'w:cs', 'w:bidi']:
            el = rPr.find(qn(cs_tag))
            if el is not None:
                rPr.remove(el)
    else:
        for attr in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
            rFonts.set(qn(f'w:{attr}'), formal_name)

    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        rPr.append(lang)

    if is_krutidev(formal_name):
        lang.set(qn('w:val'),   'x-none')
        lang.set(qn('w:ascii'), 'x-none')
        lang.set(qn('w:hAnsi'), 'x-none')
        bidi_lang = qn('w:bidi')
        if lang.get(bidi_lang):
            del lang.attrib[bidi_lang]
    else:
        lang.set(qn('w:val'), 'hi-IN')


def add_run_with_font(para, text, font_name, size_pt, bold=False, color=None):
    run = para.add_run(text)
    run.bold = bold
    set_font_properly(run, font_name)
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def add_fld_char(run, fld_type):
    fc = OxmlElement('w:fldChar')
    fc.set(qn('w:fldCharType'), fld_type)
    run._r.append(fc)


def add_instr_text(run, instr):
    it = OxmlElement('w:instrText')
    it.text = instr
    run._r.append(it)


# ═══════════════════════════
# JUSTIFY HELPER
# ═══════════════════════════

def apply_clean_justify(para):
    """Justify only long paragraphs. Short lines stay LEFT."""
    text  = para.text.strip()
    words = text.split()
    if len(words) < 12 or len(text) < 100 or text.endswith(('?', ':', '!', ';')):
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pPr = para._p.get_or_add_pPr()
    for jc in pPr.findall(qn('w:jc')):
        pPr.remove(jc)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'both')
    pPr.append(jc)


def get_original_alignment(para):
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return None
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        return None
    val = jc.get(qn('w:val'))
    mapping = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right':  WD_ALIGN_PARAGRAPH.RIGHT,
        'both':   WD_ALIGN_PARAGRAPH.JUSTIFY,
        'left':   WD_ALIGN_PARAGRAPH.LEFT,
    }
    return mapping.get(val)


# ═══════════════════════════
# CORE FORMATTING ENGINE
# ═══════════════════════════

def apply_para_formatting(para, etype, font_name, font_size_pt, bold, color, align,
                           space_before_pt, space_after_pt,
                           first_indent=None, left_indent=None,
                           line_spacing=1.15):
    set_para_font(para, font_name)
    clear_pPr_sz(para)
    set_pPr_sz(para, int(font_size_pt * 2))

    # Set paragraph-level bold in pPr>rPr so all runs inherit it
    pPr_b = para._p.get_or_add_pPr()
    rPr_b = pPr_b.find(qn('w:rPr'))
    if rPr_b is None:
        rPr_b = OxmlElement('w:rPr')
        pPr_b.append(rPr_b)
    # Set/clear w:b at paragraph level
    b_ppr = rPr_b.find(qn('w:b'))
    if bold:
        if b_ppr is None:
            b_ppr = OxmlElement('w:b')
            rPr_b.insert(0, b_ppr)
        b_ppr.attrib.pop(qn('w:val'), None)
    else:
        if b_ppr is not None:
            rPr_b.remove(b_ppr)
    # Always remove bCs at paragraph level — overrides run bold=False otherwise
    bcs_ppr = rPr_b.find(qn('w:bCs'))
    if bcs_ppr is not None:
        rPr_b.remove(bcs_ppr)

    para.paragraph_format.space_before = Pt(space_before_pt)
    para.paragraph_format.space_after  = Pt(space_after_pt)
    
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'),            str(int(space_before_pt * 20)))
    spacing.set(qn('w:after'),             str(int(space_after_pt  * 20)))
    spacing.set(qn('w:beforeAutospacing'), '0')
    spacing.set(qn('w:afterAutospacing'),  '0')

    try:
        ls = float(line_spacing)
    except Exception:
        ls = 1.15

    if ls == 1.0:
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif ls == 2.0:
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    else:
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para.paragraph_format.line_spacing = ls

    ind = pPr.find(qn('w:ind'))
    if ind is not None:
        pPr.remove(ind)

    if first_indent is not None or left_indent is not None:
        ind = OxmlElement('w:ind')
        if left_indent is not None:
            ind.set(qn('w:left'), str(int(left_indent * 1440)))
        if first_indent is not None:
            twips = int(first_indent * 1440) if isinstance(first_indent, float) else int(first_indent.inches * 1440)
            ind.set(qn('w:firstLine'), str(twips))
        pPr.append(ind)

    para.alignment = align
    pPr2 = para._p.get_or_add_pPr()
    for jc_el in pPr2.findall(qn('w:jc')):
        pPr2.remove(jc_el)
    jc_new = OxmlElement('w:jc')
    align_val_map = {
        WD_ALIGN_PARAGRAPH.JUSTIFY: 'both',
        WD_ALIGN_PARAGRAPH.CENTER:  'center',
        WD_ALIGN_PARAGRAPH.LEFT:    'left',
        WD_ALIGN_PARAGRAPH.RIGHT:   'right',
    }
    jc_new.set(qn('w:val'), align_val_map.get(align, 'left'))
    pPr2.append(jc_new)

    for run in para.runs:
        if run_has_drawing(run):
            continue
        run.bold = bold
        run.italic = False
        run.underline = False
        r = run._element
        rPr = r.get_or_add_rPr()
        # Force bold XML — remove bCs always to prevent bold bleed from original doc
        b_el = rPr.find(qn('w:b'))
        if bold:
            if b_el is None:
                b_el = OxmlElement('w:b')
                rPr.insert(0, b_el)
            b_el.attrib.pop(qn('w:val'), None)  # clear val=false if present
        else:
            if b_el is not None:
                rPr.remove(b_el)
        # Always remove bCs — it independently forces bold for complex scripts
        bcs_el = rPr.find(qn('w:bCs'))
        if bcs_el is not None:
            rPr.remove(bcs_el)
        for tag in ['w:strike', 'w:dstrike', 'w:highlight', 'w:shd',
                    'w:em', 'w:outline', 'w:shadow', 'w:emboss', 'w:imprint']:
            el = rPr.find(qn(tag))
            if el is not None:
                rPr.remove(el)
        set_font_properly(run, font_name, font_size_pt)
        run.font.color.rgb = color


# ═══════════════════════════
# TABLE HELPERS
# ═══════════════════════════


def _apply_table_borders(table):
    """Apply Table Grid borders (all sides, black, 1pt) to a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Remove tblCellSpacing (causes gaps between cells)
    cell_spacing = tblPr.find(qn('w:tblCellSpacing'))
    if cell_spacing is not None:
        tblPr.remove(cell_spacing)

    # Set tblW to 100% auto
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.insert(0, tblW)
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')

    # Remove old borders if any, then add fresh ones
    old_borders = tblPr.find(qn('w:tblBorders'))
    if old_borders is not None:
        tblPr.remove(old_borders)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')       # 0.5pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Also set cell margins tight
    tblCellMar = tblPr.find(qn('w:tblCellMar'))
    if tblCellMar is not None:
        tblPr.remove(tblCellMar)
    tblCellMar = OxmlElement('w:tblCellMar')
    for side in ['top', 'left', 'bottom', 'right']:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), '80')
        m.set(qn('w:type'), 'dxa')
        tblCellMar.append(m)
    tblPr.append(tblCellMar)


def format_table_cells(doc, font_name, base_size, line_spacing, black):
    for table in doc.tables:
        # Apply borders to every table
        _apply_table_borders(table)

        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if not para.text.strip() and not has_drawing(para):
                        continue
                    set_para_font(para, font_name)
                    clear_pPr_sz(para)
                    set_pPr_sz(para, int(base_size * 2))

                    # Apply line spacing
                    pPr = para._p.get_or_add_pPr()
                    spacing = pPr.find(qn('w:spacing'))
                    if spacing is None:
                        spacing = OxmlElement('w:spacing')
                        pPr.append(spacing)
                    try:
                        ls = float(line_spacing)
                    except Exception:
                        ls = 1.5
                    if ls == 1.0:
                        spacing.set(qn('w:lineRule'), 'auto')
                        spacing.set(qn('w:line'), '240')
                    elif ls == 2.0:
                        spacing.set(qn('w:lineRule'), 'auto')
                        spacing.set(qn('w:line'), '480')
                    else:
                        spacing.set(qn('w:lineRule'), 'auto')
                        spacing.set(qn('w:line'), str(int(ls * 240)))

                    for run in para.runs:
                        if run_has_drawing(run):
                            continue
                        was_bold = run.bold
                        set_font_properly(run, font_name, base_size)
                        run.bold = was_bold
                        run.font.color.rgb = black


# def format_table_cells(doc, font_name, base_size, line_spacing, black):
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for para in cell.paragraphs:
#                     if not para.text.strip() and not has_drawing(para):
#                         continue
#                     set_para_font(para, font_name)
#                     clear_pPr_sz(para)
#                     set_pPr_sz(para, int(base_size * 2))
#                     for run in para.runs:
#                         if run_has_drawing(run):
#                             continue
#                         was_bold   = run.bold
#                         set_font_properly(run, font_name, base_size)
#                         run.bold   = was_bold
#                         run.font.color.rgb = black


def center_all_tables(doc):
    for table in doc.tables:
        tbl   = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        jc = tblPr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            tblPr.append(jc)
        jc.set(qn('w:val'), 'center')


def set_para_text_formatted(para, new_text, font_size_pt, bold, color, font_name=None):
    p = para._p
    for r in p.findall(qn('w:r')):
        p.remove(r)
    run = para.add_run(new_text)
    run.bold = True if bold else False
    run.font.size = Pt(font_size_pt)
    run.font.color.rgb = color
    if font_name:
        set_font_properly(run, font_name, font_size_pt)


def strip_list_numbering(para):
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return
    numPr = pPr.find(qn('w:numPr'))
    if numPr is not None:
        pPr.remove(numPr)


def apply_caps_upper(para, krutidev_mode=False):
    if krutidev_mode:
        return
    for run in para.runs:
        if run.text:
            run.text = run.text.upper()


# ═══════════════════════════
# FULL-DOCUMENT KRUTI DEV CONVERSION
# ═══════════════════════════

ENGLISH_SPECIAL_CHARS = [
    ('\u201c', '"'), ('\u201d', '"'), ('\u2018', "'"), ('\u2019', "'"),
    ('\u2013', '-'), ('\u2014', '--'), ('\u2026', '...'), ('\u00a0', ' '),
    ('\u2022', '*'), ('\u00b7', '*'),
]

def _fix_english_special(text):
    for uni, repl in ENGLISH_SPECIAL_CHARS:
        text = text.replace(uni, repl)
    return text


# Punctuation/symbols that must pass through as-is in KrutiDev docs.
# These chars conflict with KrutiDev glyph positions if left in a KrutiDev-font run,
# so non-Hindi segments get a fallback font (Times New Roman) instead.
FALLBACK_FONT = 'Times New Roman'

def _segment_text(text):
    """Split text into (is_hindi, chunk) pairs."""
    segments = []
    current_hindi = None
    current_chunk = []
    for ch in text:
        ch_is_hindi = '\u0900' <= ch <= '\u097F'
        if current_hindi is None:
            current_hindi = ch_is_hindi
        if ch_is_hindi == current_hindi:
            current_chunk.append(ch)
        else:
            segments.append((current_hindi, ''.join(current_chunk)))
            current_hindi = ch_is_hindi
            current_chunk = [ch]
    if current_chunk:
        segments.append((current_hindi, ''.join(current_chunk)))
    return segments


def convert_run_to_krutidev(run):
    """Legacy single-run conversion (no font splitting). Used when caller handles font."""
    text = run.text
    if not text:
        return
    if not has_unicode_hindi(text):
        fixed = _fix_english_special(text)
        if fixed != text:
            run.text = fixed
        return
    segments = _segment_text(text)
    run.text = ''.join(
        unicode_to_krutidev(seg) if is_h else _fix_english_special(seg)
        for is_h, seg in segments
    )


def _copy_run_props(src_r, dst_r):
    """Copy rPr from src run element to dst run element, except rFonts."""
    from docx.oxml.ns import qn as _qn
    src_rPr = src_r.find(_qn('w:rPr'))
    if src_rPr is None:
        return
    dst_rPr = dst_r.get_or_add_rPr()
    for child in list(src_rPr):
        tag = child.tag
        if tag == _qn('w:rFonts'):
            continue  # font set separately per segment
        # replace existing child of same tag
        existing = dst_rPr.find(tag)
        if existing is not None:
            dst_rPr.remove(existing)
        import copy
        dst_rPr.append(copy.deepcopy(child))


def _set_run_font_fallback(run_elem, font_name):
    """Set fallback (Times New Roman) font on a non-Hindi run element."""
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement
    rPr = run_elem.get_or_add_rPr()
    rFonts = rPr.find(_qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    # Clear all attrs first
    for attr in list(rFonts.attrib.keys()):
        del rFonts.attrib[attr]
    for attr in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
        rFonts.set(_qn(f'w:{attr}'), FALLBACK_FONT)
    # lang: en-US for English/punctuation runs
    lang = rPr.find(_qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        rPr.append(lang)
    lang.set(_qn('w:val'), 'en-US')
    for a in ['ascii', 'hAnsi']:
        lang.set(_qn(f'w:{a}'), 'en-US')


def _split_and_convert_run(run, font_name, pre_font=None):
    """
    Split a mixed Hindi+English run into multiple runs:
    - Hindi segments  -> KrutiDev font, unicode_to_krutidev() conversion
    - Non-Hindi segs  -> Times New Roman font, as-is (special chars normalized)
    Returns list of new run elements to insert, or None if no split needed.
    """
    from lxml import etree
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement
    import copy

    text = run.text
    if not text:
        return None

    # Normalize special chars first
    text = _fix_english_special(text)

    if not has_unicode_hindi(text):
        # Pure ASCII run — could be already-converted KrutiDev OR English/punctuation.
        # Use python-docx's inherited font resolution to check effective font.
        run.text = text  # already fixed special chars above
        effective_font = pre_font or ''
        if not effective_font:
            try:
                effective_font = run.font.name or ''
            except Exception:
                effective_font = ''
        # if is_krutidev(effective_font):
        #     return None  # caller will set_font_properly with KrutiDev
        # else:
        #     # Check for KrutiDev-encoded digit/special chars
        #     KRUTIDEV_ENCODED = set('!@#$%^&*(')
        #     if any(c in KRUTIDEV_ENCODED for c in text):
        #         return None  # treat as KrutiDev — caller sets KrutiDev font
        #     _set_run_font_fallback(run._element, font_name)
        #     return None


        if is_krutidev(effective_font):
            # Already KrutiDev-encoded run — keep KrutiDev font
            return None  # caller will set_font_properly with KrutiDev
        else:
            # Non-Hindi, non-KrutiDev run — use fallback font
            _set_run_font_fallback(run._element, font_name)
            return None

    segments = _segment_text(text)
    if len(segments) == 1:
        # Pure Hindi run — normal KrutiDev conversion
        run.text = unicode_to_krutidev(text)
        return None  # caller sets KrutiDev font

    # Mixed run — need to split
    src_r = run._element
    new_runs = []

    for is_hindi, seg in segments:
        if not seg:
            continue
        new_r = OxmlElement('w:r')
        _copy_run_props(src_r, new_r)

        t = OxmlElement('w:t')
        if seg.startswith(' ') or seg.endswith(' '):
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

        if is_hindi:
            t.text = unicode_to_krutidev(seg)
            new_r.append(t)
            # Build a temporary run object to call set_font_properly
            from docx.text.run import Run
            from docx.oxml.ns import qn as _qn2
            rPr = new_r.get_or_add_rPr()
            rFonts = rPr.find(_qn2('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            for attr in list(rFonts.attrib.keys()):
                del rFonts.attrib[attr]
            formal = FONT_NAME_MAP.get(font_name, font_name)
            rFonts.set(_qn2('w:ascii'), formal)
            rFonts.set(_qn2('w:hAnsi'), formal)
            ea = _qn2('w:eastAsia')
            if rFonts.get(ea): del rFonts.attrib[ea]
            cs = _qn2('w:cs')
            if rFonts.get(cs): del rFonts.attrib[cs]
            for ta in ['w:asciiTheme','w:hAnsiTheme','w:eastAsiaTheme','w:cstheme']:
                ta2 = _qn2(ta)
                if rFonts.get(ta2): del rFonts.attrib[ta2]
            lang = rPr.find(_qn2('w:lang'))
            if lang is None:
                lang = OxmlElement('w:lang')
                rPr.append(lang)
            lang.set(_qn2('w:val'), 'x-none')
            lang.set(_qn2('w:ascii'), 'x-none')
            lang.set(_qn2('w:hAnsi'), 'x-none')
            bidi = _qn2('w:bidi')
            if lang.get(bidi): del lang.attrib[bidi]

        # else:
        #     t.text = _fix_english_special(seg)
        #     new_r.append(t)
        #     # KrutiDev digit/encoded chars must stay in KrutiDev font
        #     KRUTIDEV_ENCODED = set('!@#$%^&*(')
        #     if any(c in KRUTIDEV_ENCODED for c in seg):
        #         from docx.oxml import OxmlElement as _OE2
        #         rPr2 = new_r.get_or_add_rPr()
        #         rFonts2 = rPr2.find(_qn('w:rFonts'))
        #         if rFonts2 is None:
        #             rFonts2 = _OE2('w:rFonts')
        #             rPr2.insert(0, rFonts2)
        #         for attr in list(rFonts2.attrib.keys()):
        #             del rFonts2.attrib[attr]
        #         formal = FONT_NAME_MAP.get(font_name, font_name)
        #         rFonts2.set(_qn('w:ascii'), formal)
        #         rFonts2.set(_qn('w:hAnsi'), formal)
        #     else:
        #         _set_run_font_fallback(new_r, font_name)

        else:
            t.text = _fix_english_special(seg)
            new_r.append(t)
            _set_run_font_fallback(new_r, font_name)





        new_runs.append(new_r)

    return new_runs


def convert_doc_runs(doc, font_name):
    if not is_krutidev(font_name):
        return

    def process_para(para):
        if has_drawing(para):
            return
        set_para_font(para, font_name)

        # Collect runs to process (list snapshot — we'll modify XML directly)
        runs_to_process = [r for r in para.runs if not run_has_drawing(r)]

        for run in runs_to_process:
            r_elem = run._element
            # Capture effective font BEFORE clearing (for inheritance-aware detection)
            try:
                run_effective_font = run.font.name or font_name
            except Exception:
                run_effective_font = font_name
            # Clear existing font attrs
            from docx.oxml.ns import qn as _qn
            rPr = r_elem.get_or_add_rPr()
            rFonts = rPr.find(_qn('w:rFonts'))
            if rFonts is None:
                from docx.oxml import OxmlElement as _OE
                rFonts = _OE('w:rFonts')
                rPr.insert(0, rFonts)
            for attr in list(rFonts.attrib.keys()):
                del rFonts.attrib[attr]

            new_runs = _split_and_convert_run(run, font_name, run_effective_font)

            if new_runs is not None:
                # Replace original run with split runs
                parent = r_elem.getparent()
                idx = list(parent).index(r_elem)
                parent.remove(r_elem)
                for i, nr in enumerate(new_runs):
                    parent.insert(idx + i, nr)
            else:
                # Pure run (single language) — set appropriate font
                if run.text and has_unicode_hindi(run.text):
                    # Devanagari text -> KrutiDev
                    set_font_properly(run, font_name)
                elif is_krutidev(run_effective_font):
                    # Already-converted KrutiDev ASCII -> keep KrutiDev
                    set_font_properly(run, font_name)
                else:
                    # Genuine English/punctuation -> fallback font
                    _set_run_font_fallback(r_elem, font_name)

    # 1. Main paragraphs
    for para in doc.paragraphs:
        process_para(para)

    # 2. Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_para(para)

    # 3. Headers & Footers (Fix for missing title/header conversion)
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for para in header.paragraphs:
                    process_para(para)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    process_para(para)

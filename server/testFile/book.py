import re
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from utils import (
    has_drawing, run_has_drawing, is_all_bold, is_bullet_para,
    apply_para_formatting, set_para_text_formatted, strip_list_numbering,
    apply_clean_justify, format_table_cells, add_run_with_font,
    set_font_properly, is_krutidev, CHAPTER_HEADING_RE, CHAPTER_HEADING_LOOSE_RE,
    inject_heading_number
)


# ═══════════════════════════
# TITLE PAGE — BOOK
# ═══════════════════════════

def insert_title_page(doc, opts, font_name):
    black = RGBColor(0, 0, 0)
    gray  = RGBColor(100, 100, 100)
    title       = opts.get('title', '').strip()
    author      = opts.get('author', '').strip()
    volume      = opts.get('volume', '').strip()
    isbn        = opts.get('isbn', '').strip()
    website     = opts.get('website_url', '').strip()
    footer_text = opts.get('footer', '').strip()

    if not title and not author:
        return

    insert_paras = []

    def make_para(text, align, size, bold=False, space_before=0, space_after=12, color=None):
        from docx.enum.text import WD_LINE_SPACING
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        if text:
            add_run_with_font(p, text, font_name, size, bold=bold, color=color or black)
        return p

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(72)
    spacer.paragraph_format.space_after  = Pt(0)
    insert_paras.append(spacer)

    if title:
        insert_paras.append(make_para(title, WD_ALIGN_PARAGRAPH.CENTER, 28, bold=True, space_after=20))
    if volume:
        insert_paras.append(make_para(volume, WD_ALIGN_PARAGRAPH.CENTER, 14, space_after=14, color=gray))
    if author:
        insert_paras.append(make_para(author, WD_ALIGN_PARAGRAPH.CENTER, 16, bold=True, space_before=10, space_after=14))

    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.space_before = Pt(20)
    sep.paragraph_format.space_after  = Pt(20)
    r = sep.add_run('\u2015 \u2015 \u2015')
    set_font_properly(r, font_name)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(180, 180, 180)
    insert_paras.append(sep)

    if footer_text:
        insert_paras.append(make_para(footer_text, WD_ALIGN_PARAGRAPH.CENTER, 11, space_after=8, color=gray))
    if website:
        insert_paras.append(make_para(website, WD_ALIGN_PARAGRAPH.CENTER, 10, space_after=8, color=gray))
    if isbn:
        insert_paras.append(make_para(f'ISBN: {isbn}', WD_ALIGN_PARAGRAPH.CENTER, 10, space_after=0, color=gray))

    pb_para = doc.add_paragraph()
    pb_para.paragraph_format.space_before = Pt(0)
    pb_para.paragraph_format.space_after  = Pt(0)
    run = pb_para.add_run()
    br  = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    insert_paras.append(pb_para)

    body = doc.element.body
    for p in reversed(insert_paras):
        body.remove(p._element)
        body.insert(0, p._element)


# ═══════════════════════════
# STRUCTURE DETECTION — BOOK / RESEARCH
# ═══════════════════════════

def _is_conclusion_heading(text):
    CONCLUSION_WORDS = {
        'conclusion', 'conclusions', 'summary', 'chapter summary',
        'concluding remarks', 'unit summary', 'let us sum up',
        'let us sumup', 'key points', 'review questions',
        'निष्कर्ष', 'सारांश', 'समापन', 'अध्याय सारांश',
    }
    t = text.lower().strip().rstrip(':').strip()
    return t in CONCLUSION_WORDS or t.startswith('conclusion') or t.startswith('निष्कर्ष')


def detect_structure(para, index, doc=None):
    """Book-aware structure detection."""
    text  = para.text.strip()
    words = text.split()
    wc    = len(words)

    if wc == 0:
        return 'empty'
    if has_drawing(para):
        return 'drawing'

    is_bold = is_all_bold(para)

    if is_bullet_para(para):
        if is_bold and wc <= 15:
            return 'sub_heading'
        return 'bullet'

    if index < 5 and text.isupper() and wc <= 15 and is_bold:
        return 'book_title'

    if CHAPTER_HEADING_RE.match(text) and wc <= 20:
        return 'chapter_heading'

    # 'षष्ठम अध्याय: ...' style — word before अध्याय
    if CHAPTER_HEADING_LOOSE_RE.match(text) and wc <= 20:
        return 'chapter_heading'

    if doc and index > 0:
        prev_text = doc.paragraphs[index - 1].text.strip()
        if CHAPTER_HEADING_RE.match(prev_text) and wc <= 20:
            return 'chapter_heading'

    TABLE_PAT  = r'(table|तालिका|सारणी)'
    FIGURE_PAT = r'(figure|fig\.?|चित्र|आकृति)'

    if re.match(TABLE_PAT, text, re.IGNORECASE) and wc <= 25:
        return 'table_caption'
    if re.match(FIGURE_PAT, text, re.IGNORECASE) and wc <= 25:
        return 'figure_caption'

    if re.match(r'^\d+(\.\d+)?\s+' + TABLE_PAT, text, re.IGNORECASE) and wc <= 25:
        return 'table_caption'
    if re.match(r'^\d+(\.\d+)?\s+' + FIGURE_PAT, text, re.IGNORECASE) and wc <= 25:
        return 'figure_caption'

    CONTENT_SECTION_WORDS = {
        'example', 'examples', 'exercise', 'exercises', 'activity', 'activities',
        'practice', 'practices', 'problem', 'problems', 'question', 'questions',
        'solution', 'solutions', 'answer', 'answers', 'task', 'tasks',
        'assignment', 'assignments', 'note', 'notes', 'tip', 'tips',
        'hint', 'hints', 'remark', 'remarks', 'illustration', 'illustrations',
        'case study', 'case studies', 'sample', 'samples',
        'उदाहरण', 'अभ्यास', 'प्रश्न', 'उत्तर', 'समाधान', 'कार्य', 'टिप्पणी',
    }

    bare = re.sub(r'^\d+(\.\d+)*\.?\s+', '', text).strip().rstrip(':').lower()
    if bare in CONTENT_SECTION_WORDS:
        return 'body'

    if re.match(r'^\d+\.\d+\.?\s+\S', text) and is_bold and wc <= 20:
        return 'sub_heading'

    # ASCII digit OR Devanagari digit starting numbered heading
    if re.match(r'^[1-9]\d*\.?\s+\S', text) and is_bold and wc <= 20:
        return 'main_heading'

    # Devanagari numeral starting heading (१., २. etc.)
    if re.match(r'^[१-९][०-९]*\.?\s+\S', text) and is_bold and wc <= 20:
        return 'main_heading'

    if is_bold and wc <= 15:
        return 'sub_heading'

    return 'body'


# ═══════════════════════════
# BOOK / RESEARCH BODY FORMATTING
# ═══════════════════════════

def format_book_body(doc, opts, font_name):
    black         = RGBColor(0, 0, 0)
    krutidev_mode = is_krutidev(font_name)
    base_size     = float(opts.get('font_size', 14 if krutidev_mode else 12))
    line_spacing  = float(opts.get('line_spacing', 1.5))

    heading_font = 'Kruti Dev 010' if krutidev_mode else 'Times New Roman'
    heading_counters = [0, 0]

    i          = 0
    prev_etype = None

    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]

        if has_drawing(para):
            para.paragraph_format.space_after = Pt(5)
            pPr_d = para._p.get_or_add_pPr()
            sp_d = pPr_d.find(qn('w:spacing'))
            if sp_d is None:
                sp_d = OxmlElement('w:spacing')
                pPr_d.append(sp_d)
            sp_d.set(qn('w:after'), '100')
            sp_d.set(qn('w:afterAutospacing'), '0')
            i += 1
            continue

        text = para.text.strip()
        if not text:
            i += 1
            continue

        etype = detect_structure(para, i, doc)
        if etype in ('empty', 'drawing'):
            i += 1
            continue

        space_after  = 5.0
        space_before = 5.0

        if etype == 'book_title':
            apply_para_formatting(para, etype, heading_font,
                font_size_pt=24, bold=True, color=black,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                space_before_pt=15, space_after_pt=10,
                line_spacing=line_spacing)
            set_para_text_formatted(para, text.upper(), 20, True, black, heading_font)

        elif etype == 'chapter_heading':
            heading_counters[0] = 0
            heading_counters[1] = 0

            if ':' in text and re.match(r'^(chapter|unit|part|lesson)\s*[-–—]?\s*\S+', text, re.IGNORECASE):
                parts         = text.split(':', 1)
                chapter_label = parts[0].strip()
                chapter_title = parts[1].strip()

                apply_para_formatting(para, etype, heading_font,
                    font_size_pt=24, bold=True, color=black,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before_pt=15, space_after_pt=0,
                    line_spacing=line_spacing)
                set_para_text_formatted(para,
                    chapter_label.upper() if not krutidev_mode else chapter_label,
                    24, True, black, heading_font)

                title_para = doc.add_paragraph()
                para._p.addnext(title_para._p)
                apply_para_formatting(title_para, 'chapter_title', heading_font,
                    font_size_pt=18, bold=True, color=black,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before_pt=0, space_after_pt=10,
                    line_spacing=line_spacing)
                set_para_text_formatted(title_para,
                    chapter_title.upper() if not krutidev_mode else chapter_title,
                    18, True, black, heading_font)
                i += 2
                prev_etype = 'chapter_heading'
                continue
            else:
                next_is_title = False
                if i + 1 < len(doc.paragraphs):
                    nxt      = doc.paragraphs[i + 1]
                    nxt_text = nxt.text.strip()
                    nxt_etype = detect_structure(nxt, i + 1, doc) if nxt_text else 'empty'
                    if nxt_etype == 'chapter_heading' and not re.match(
                            r'^(chapter|unit|part|lesson)\s*[-–—]?\s*\S+',
                            nxt_text, re.IGNORECASE):
                        next_is_title = True

                apply_para_formatting(para, etype, heading_font,
                    font_size_pt=24, bold=True, color=black,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before_pt=15, space_after_pt=0 if next_is_title else 10,
                    line_spacing=line_spacing)
                set_para_text_formatted(para,
                    text.upper() if not krutidev_mode else text,
                    24, True, black, heading_font)
                prev_etype = etype
                i += 1

                if next_is_title and i < len(doc.paragraphs):
                    title_para  = doc.paragraphs[i]
                    title_text  = title_para.text.strip()
                    apply_para_formatting(title_para, 'chapter_title', heading_font,
                        font_size_pt=18, bold=True, color=black,
                        align=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before_pt=0, space_after_pt=10,
                        line_spacing=line_spacing)
                    set_para_text_formatted(title_para,
                        title_text.upper() if not krutidev_mode else title_text,
                        18, True, black, heading_font)
                    prev_etype = 'chapter_title'
                    i += 1
                continue

        elif etype == 'main_heading':
            strip_list_numbering(para)
            heading_counters[0] += 1
            heading_counters[1]  = 0
            inject_heading_number(para, heading_counters[0], krutidev_mode=krutidev_mode)
            apply_para_formatting(para, etype, font_name,
                font_size_pt=14, bold=True, color=black,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_before_pt=4, space_after_pt=4,
                left_indent=0.0, first_indent=0.0,
                line_spacing=line_spacing)

        elif etype == 'sub_heading':
            strip_list_numbering(para)
            heading_counters[1] += 1
            m = re.match(r'^(\d+)\.(\d+)\.?\s+', text)
            if m:
                heading_counters[0] = int(m.group(1))
                heading_counters[1] = int(m.group(2))
            else:
                inject_heading_number(para, heading_counters[0], heading_counters[1], krutidev_mode=krutidev_mode)

            apply_para_formatting(para, etype, font_name,
                font_size_pt=14, bold=True, color=black,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_before_pt=4, space_after_pt=4,
                left_indent=0.0, first_indent=0.0,
                line_spacing=line_spacing)

        elif etype == 'table_caption':
            apply_para_formatting(para, etype, font_name,
                font_size_pt=12, bold=True, color=black,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                space_before_pt=6, space_after_pt=4,
                line_spacing=1.0)
            for run in para.runs:
                if not run_has_drawing(run):
                    run.italic = True

        elif etype == 'figure_caption':
            apply_para_formatting(para, etype, font_name,
                font_size_pt=12, bold=True, color=black,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                space_before_pt=4, space_after_pt=6,
                line_spacing=1.0)
            for run in para.runs:
                if not run_has_drawing(run):
                    run.italic = True

        elif etype == 'bullet':
            is_bold_para = is_all_bold(para)
            apply_para_formatting(para, etype, font_name,
                font_size_pt=base_size, bold=is_bold_para, color=black,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before_pt=0, space_after_pt=space_after,
                left_indent=0.25, first_indent=-0.25,
                line_spacing=line_spacing)

        else:  # body
            apply_clean_justify(para)
            final_align = para.alignment if para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY else WD_ALIGN_PARAGRAPH.JUSTIFY

            apply_para_formatting(para, etype, font_name,
                font_size_pt=base_size, bold=False, color=black,
                align=final_align,
                space_before_pt=0, space_after_pt=space_after,
                left_indent=0.0, first_indent=0.0,
                line_spacing=line_spacing)

        prev_etype = etype
        i += 1

    format_table_cells(doc, font_name, base_size, line_spacing, black)

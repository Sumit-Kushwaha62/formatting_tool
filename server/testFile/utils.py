import re
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ═══════════════════════════
# DRAWING / IMAGE DETECTION
# ═══════════════════════════

WP_NS  = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
MC_NS  = 'http://schemas.openxmlformats.org/markup-compatibility/2006'
W_NS   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def has_drawing(para):
    """Return True if paragraph contains any image/chart/drawing/object element."""
    p = para._p
    for tag in [qn('w:drawing'), qn('w:pict'), qn('w:object')]:
        if p.find('.//' + tag) is not None:
            return True
    if p.find(f'{{{MC_NS}}}AlternateContent') is not None:
        return True
    return False


# ═══════════════════════════
# PRE-CLEANING
# ═══════════════════════════

def clean_runs_in_para(para):
    for run in para.runs:
        cleaned = run.text.replace('\t', '').replace('\n', ' ')
        cleaned = re.sub(r' {2,}', ' ', cleaned)
        run.text = cleaned


def remove_proof_errors(para):
    """Remove <w:proofErr> elements that cause word splits during run iteration."""
    p = para._p
    for proof_err in p.findall(qn('w:proofErr')):
        p.remove(proof_err)


def run_has_drawing(run):
    """Return True if this run contains any drawing/image element."""
    r = run._r
    if r.find(qn('w:drawing')) is not None:
        return True
    if r.find(qn('w:pict')) is not None:
        return True
    if r.find(qn('w:object')) is not None:
        return True
    return False


def merge_runs_in_para(para):
    """Merge adjacent runs with identical formatting. NEVER merges runs that contain drawings."""
    if len(para.runs) <= 1:
        return

    i = 0
    while i < len(para.runs) - 1:
        r1 = para.runs[i]
        r2 = para.runs[i + 1]

        if run_has_drawing(r1) or run_has_drawing(r2):
            i += 1
            continue

        def fmt_sig(run):
            rPr = run._element.find(qn('w:rPr'))
            if rPr is None:
                return ('', None, None, None, None)
            bold   = rPr.find(qn('w:b'))
            italic = rPr.find(qn('w:i'))
            sz     = rPr.find(qn('w:sz'))
            color  = rPr.find(qn('w:color'))
            rFonts = rPr.find(qn('w:rFonts'))
            b_val     = None if bold   is None else bold.get(qn('w:val'), 'true')
            i_val     = None if italic is None else italic.get(qn('w:val'), 'true')
            sz_val    = None if sz     is None else sz.get(qn('w:val'))
            color_val = None if color  is None else color.get(qn('w:val'))
            font_val  = None if rFonts is None else rFonts.get(qn('w:ascii'))
            return (b_val, i_val, sz_val, color_val, font_val)

        if fmt_sig(r1) == fmt_sig(r2):
            r1.text = (r1.text or '') + (r2.text or '')
            r2._r.getparent().remove(r2._r)
        else:
            i += 1


def is_all_bold(para):
    runs = [r for r in para.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)


def is_bullet_para(para):
    """True if paragraph has a numbering/bullet list marker in XML."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return False
    return pPr.find(qn('w:numPr')) is not None


def apply_bold_before_colon(para, font_name, krutidev_mode):
    """If para has 'Label: rest' pattern, make text before ':' bold."""
    text = para.text
    colon_idx = text.find(':')
    if colon_idx <= 0 or colon_idx > 80:
        return

    label = text[:colon_idx + 1]
    rest  = text[colon_idx + 1:]

    first_run = para.runs[0] if para.runs else None
    size_pt = None
    if first_run:
        size_pt = first_run.font.size

    for run in list(para.runs):
        run._r.getparent().remove(run._r)

    r_bold = para.add_run(label)
    r_bold.bold = True
    if not krutidev_mode:
        set_font_properly(r_bold, font_name)
        if size_pt:
            r_bold.font.size = size_pt
        r_bold.font.color.rgb = RGBColor(0, 0, 0)

    if rest:
        r_rest = para.add_run(rest)
        r_rest.bold = False
        if not krutidev_mode:
            set_font_properly(r_rest, font_name)
            if size_pt:
                r_rest.font.size = size_pt
            r_rest.font.color.rgb = RGBColor(0, 0, 0)


def merge_split_paragraphs(doc):
    pass  # Disabled — causes duplicate rendering with mixed fonts


def clear_pPr_sz(para):
    """Remove font size override from paragraph-level rPr."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        return
    for tag in [qn('w:sz'), qn('w:szCs')]:
        el = rPr.find(tag)
        if el is not None:
            rPr.remove(el)


def set_pPr_sz(para, half_pts):
    """Set font size at paragraph-level rPr."""
    pPr = para._p.get_or_add_pPr()
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)
    for tag_name in ['w:sz', 'w:szCs']:
        el = rPr.find(qn(tag_name))
        if el is None:
            el = OxmlElement(tag_name)
            rPr.append(el)
        el.set(qn('w:val'), str(half_pts))


def clear_para_indent(para):
    """Remove all left/first-line indent from paragraph XML."""
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is not None:
        pPr.remove(ind)
    para.paragraph_format.left_indent        = None
    para.paragraph_format.first_line_indent  = None


def preprocess_document(doc):
    for para in doc.paragraphs:
        if has_drawing(para):
            continue
        remove_proof_errors(para)
        clean_runs_in_para(para)
        merge_runs_in_para(para)
    merge_split_paragraphs(doc)


# ═══════════════════════════
# HELPERS
# ═══════════════════════════

KRUTIDEV_FONTS = {'Kruti Dev 010', 'Kruti Dev 011', 'Krutidev010', 'Krutidev011',
                  'KrutiDev010', 'KrutiDev011', 'Kruti Dev 010', 'Kruti Dev 011'}

# ═══════════════════════════
# SHARED DETECTION CONSTANTS
# ═══════════════════════════

# Hindi + English chapter/unit words, ASCII + Devanagari numerals
CHAPTER_HEADING_RE = re.compile(
    r'^(chapter|unit|part|lesson|अध्याय|इकाई|भाग|पाठ)'
    r'\s*[-–—]?\s*(\d+|[ivxlcdmIVXLCDM]+|[०-९]+)\b',
    re.IGNORECASE
)

# Matches headings where chapter word appears ANYWHERE in first 3 words
# e.g. 'षष्ठम अध्याय: ...' or 'तृतीय अध्याय: ...'
CHAPTER_HEADING_LOOSE_RE = re.compile(
    r'^[\w\u0900-\u097F]+\s+(अध्याय|chapter|unit|part|lesson|इकाई|भाग|पाठ)'
    r'\s*[-:–—]',
    re.IGNORECASE
)


def inject_heading_number(para, sec, sub=None, krutidev_mode=False):
    """
    Inject section number prefix into heading para.
    Skipped in krutidev_mode because Hindi documents already carry
    Devanagari numerals in the text — adding an ASCII counter would
    produce garbage like '0१' or '1२'.
    Also skipped if the paragraph already starts with a digit.
    """
    if krutidev_mode:
        return
    text = para.text.strip()
    # Skip if already starts with ASCII or Devanagari digit
    if re.match(r'^[\d०-९]', text):
        return
    prefix = f"{sec}. " if sub is None else f"{sec}.{sub} "
    if para.runs:
        para.runs[0].text = prefix + para.runs[0].text.lstrip()
    else:
        para.add_run(prefix)

FONT_NAME_MAP = {
    'Krutidev010': 'Kruti Dev 010',
    'Krutidev011': 'Kruti Dev 011',
    'Mangal':      'Mangal',
    'Kokila':      'Kokila',
    'Utsaah':      'Utsaah',
    'Aparajita':   'Aparajita',
    'Nirmala UI':  'Nirmala UI',
}


# ═══════════════════════════
# HINDI CONVERSION
# ═══════════════════════════

def unicode_to_krutidev(text):
    """
    Convert Unicode Devanagari text to Kruti Dev 010 ASCII encoding.
    Uses a robust syllable-based reordering algorithm.
    """
    if not text:
        return ""
    if not re.search(r'[\u0900-\u097F]', text):
        return text

    import unicodedata
    # NFC normalize
    text = unicodedata.normalize('NFC', text)

    halant = '\u094D'
    reph = '\u0930' + halant
    
    # Syllable regex:
    # (Consonant + Halant)* + Consonant + (Matra | Halant)? + (Anusvara|Chandrabindu|Visarga)?
    syll_pattern = r'((?:[\u0905-\u0939\u0958-\u0961]\u094D)*[\u0905-\u0939\u0958-\u0961][\u093E-\u094D\u0901-\u0903\u094E-\u094F\u0955-\u0957]*)'
    
    def process_syllable(m):
        s = m.group(0)
        # 1. Handle Reph (र + ् at the beginning of a syllable)
        has_reph = False
        if s.startswith(reph) and len(s) > 2:
            has_reph = True
            s = s[2:]
        
        # 2. Handle short 'i' (ि) - move to front of syllable
        has_short_i = False
        if 'ि' in s:
            has_short_i = True
            s = s.replace('ि', '')
        
        # Reconstruct syllable in KrutiDev order
        res = ""
        if has_short_i:
            res += 'f'
        res += s
        if has_reph:
            # Reph 'Z' goes after matras but before Anusvara/chandrabindu/visarga
            # At this point text has been partially converted: check ASCII equivalents
            reph_triggers = ('a', '%')  # anusvara='a', visarga='%'
            has_nasal = any(res.endswith(t) or t in res for t in reph_triggers)
            if has_nasal:
                # Find the first nasal/visarga ASCII char position
                idx = -1
                for _i, _c in enumerate(res):
                    if res[_i:].startswith('a~') or res[_i] in ('a', '%'):
                        idx = _i
                        break
                if idx >= 0:
                    res = res[:idx] + 'Z' + res[idx:]
                else:
                    res += 'Z'
            else:
                res += 'Z'
        return res

    # Apply reordering to all syllables
    text = re.sub(syll_pattern, process_syllable, text)

    # Mapping for direct character replacement
    C = {
        'अ': 'v',  'आ': 'vk', 'इ': 'b',  'ई': 'bZ',
        'उ': 'm',  'ऊ': 'Å',
        'ए': ',',  'ऐ': ',s',
        'ओ': 'vks','औ': 'vkS','ऋ': '_',  'ॠ': '__',
        'ऑ': 'vkW',
        'ा': 'k',  'ि': 'f',  'ी': 'h',  'ु': 'q',
        'ू': 'w',  'ृ': '`',  'े': 's',  'ै': 'S',
        'ो': 'ks', 'ौ': 'kS', 'ं': 'a',  'ः': '%',  'ँ': 'p',
        'ॉ': 'W',  'ॊ': 'ks',
        'क': 'd',  'ख': '[k', 'ग': 'x',  'घ': '?k', 'ङ': 'M~',
        'च': 'p',  'छ': 'N',  'ज': 't',  'झ': '>k', 'ञ': '¥',
        'ट': 'V',  'ठ': 'B',  'ड': 'M',  'ढ': '<',  'ण': '.k',
        'त': 'r',  'थ': 'Fk', 'द': 'n',  'ध': '/k', 'न': 'u',
        'प': 'i',  'फ': 'Q',  'ब': 'c',  'भ': 'Hk', 'म': 'e',
        'य': ';',  'र': 'j',  'ल': 'y',  'व': 'o',
        'श': "'k", 'ष': '"k', 'स': 'l',  'ह': 'g',
        'ॐ': 'ks',
        'ऽ': '\'', # avagraha
        '।': 'A',  '॥': 'AA',
        '०': '0',  '१': '1',  '२': '2',  '३': '3',  '४': '4',
        '५': '5',  '६': '6',  '७': '7',  '८': '8',  '९': '9',
    }

    HALF = {
        'क': 'D',  'ख': '[',  'ग': 'X',  'घ': '?',
        'च': 'P',  'ज': 'T',  'झ': '>',
        'ट': 'V~', 'ठ': 'B~', 'ड': 'M~', 'ढ': '<~', 'ण': '.k~',
        'त': 'R',  'थ': 'F',  'द': 'n~', 'ध': '/',
        'न': 'U',  'प': 'I',  'ब': 'C',  'भ': 'H',
        'म': 'E',  'य': 'Y',  'र': 'z',
        'ल': 'y~', 'व': 'O',
        'श': "'",  'ष': '"',  'स': 'L',  'ह': 'g~',
    }

    CONJUNCTS = [
        ('\u0915\u094D\u0937', '{k'), # क्ष
        ('\u0924\u094D\u0930', '='),   # त्र
        ('\u091C\u094D\u091E', 'K'),   # ज्ञ
        ('\u0936\u094D\u0930', "'J"),  # श्र
        ('\u092A\u094D\u0930', 'iz'),  # प्र
        ('\u0917\u094D\u0930', 'xz'),  # ग्र
        ('\u0915\u094D\u0930', 'dz'),  # क्र
        ('\u092C\u094D\u0930', 'cz'),  # ब्र
        ('\u092E\u094D\u0930', 'ez'),  # म्र
        ('\u0926\u094D\u0930', 'nz'),  # द्र
        ('\u0927\u094D\u0930', '/z'),  # ध्र
        ('\u092D\u094D\u0930', 'Hkz'), # भ्र
        ('\u0939\u094D\u0930', 'gz'),  # ह्र
        ('\u0938\u094D\u0924\u094D\u0930', 'L='), # स्त्र
        ('\u0926\u094D\u092F', '|'),    # द्य
        ('\u0926\u094D\u0927', '/~/k'), # द्ध
        ('\u0926\u094D\u0935', 'n~o'),  # द्व
        ('\u0924\u094D\u0924', 'Ùk'),   # त्त
        ('\u091F\u094D\u0930', 'Vz'),   # ट्र
        ('\u0921\u094D\u0930', 'Mz'),   # ड्र
        ('रू', 'tw'),                  # रू
        ('रु', 'rq'),                  # रु
    ]

    for uni, kd in CONJUNCTS:
        text = text.replace(uni, kd)

    for uni, kd in HALF.items():
        text = text.replace(uni + halant, kd)

    res_final = []
    for char in text:
        res_final.append(C.get(char, char))
    
    return "".join(res_final)


def is_krutidev(font_name):
    return font_name and any(k.lower() in font_name.lower() for k in ['kruti', 'krutidev'])


def has_unicode_hindi(text):
    return bool(re.search(r'[\u0900-\u097F]', text))


def set_font_properly(run, font_name, size_pt=None):
    formal_name = FONT_NAME_MAP.get(font_name, font_name)
    run.font.name = formal_name

    r    = run._element
    rPr  = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()

    if is_krutidev(formal_name):
        rFonts.set(qn('w:hint'), 'default')
        for attr in ['ascii', 'hAnsi', 'eastAsia']:
            rFonts.set(qn(f'w:{attr}'), formal_name)
        # Remove cs and theme fonts thoroughly
        cs_attr = qn('w:cs')
        if rFonts.get(cs_attr):
            del rFonts.attrib[cs_attr]
        for theme_attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme']:
            ta = qn(theme_attr)
            if rFonts.get(ta):
                del rFonts.attrib[ta]
        # Remove rtl/cs/bidi markers
        for cs_tag in ['w:rtl', 'w:cs', 'w:bidi']:
            el = rPr.find(qn(cs_tag))
            if el is not None:
                rPr.remove(el)
        
        # Force language to en-US manually
        lang = rPr.find(qn('w:lang'))
        if lang is None:
            lang = OxmlElement('w:lang')
            rPr.append(lang)
        lang.set(qn('w:val'),   'en-US')
        lang.set(qn('w:ascii'), 'en-US')
        lang.set(qn('w:hAnsi'), 'en-US')
        
        # Add NoProofing attribute
        no_proof = rPr.find(qn('w:noProof'))
        if no_proof is None:
            no_proof = OxmlElement('w:noProof')
            rPr.append(no_proof)
    else:
        rFonts.set(qn('w:hint'), 'complex')
        for attr in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
            rFonts.set(qn(f'w:{attr}'), formal_name)
            
        lang = rPr.find(qn('w:lang'))
        if lang is None:
            lang = OxmlElement('w:lang')
            rPr.append(lang)
        lang.set(qn('w:val'), 'hi-IN')
        lang.set(qn('w:cs'),  'hi-IN')

    if size_pt:
        run.font.size = Pt(size_pt)
        if not is_krutidev(formal_name):
            sz_cs = rPr.find(qn('w:szCs'))
            if sz_cs is None:
                sz_cs = OxmlElement('w:szCs')
                rPr.append(sz_cs)
            sz_cs.set(qn('w:val'), str(int(size_pt * 2)))


def set_para_font(para, font_name):
    """Set font at paragraph-level rPr."""
    formal_name = FONT_NAME_MAP.get(font_name, font_name)
    pPr  = para._p.get_or_add_pPr()
    
    # CT_PPr does not have get_or_add_rPr, use manual creation
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)

    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)

    if is_krutidev(formal_name):
        rFonts.set(qn('w:hint'), 'default')
        for attr in ['ascii', 'hAnsi', 'eastAsia']:
            rFonts.set(qn(f'w:{attr}'), formal_name)
        cs_attr = qn('w:cs')
        if rFonts.get(cs_attr):
            del rFonts.attrib[cs_attr]
        for theme_attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme']:
            ta = qn(theme_attr)
            if rFonts.get(ta):
                del rFonts.attrib[ta]
        for cs_tag in ['w:rtl', 'w:cs', 'w:bidi']:
            el = rPr.find(qn(cs_tag))
            if el is not None:
                rPr.remove(el)
    else:
        for attr in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
            rFonts.set(qn(f'w:{attr}'), formal_name)

    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        rPr.append(lang)
        
    if is_krutidev(formal_name):
        lang.set(qn('w:val'), 'en-US')
    else:
        lang.set(qn('w:val'), 'hi-IN')


def add_run_with_font(para, text, font_name, size_pt, bold=False, color=None):
    run = para.add_run(text)
    run.bold = bold
    set_font_properly(run, font_name)
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def add_fld_char(run, fld_type):
    fc = OxmlElement('w:fldChar')
    fc.set(qn('w:fldCharType'), fld_type)
    run._r.append(fc)


def add_instr_text(run, instr):
    it = OxmlElement('w:instrText')
    it.text = instr
    run._r.append(it)


# ═══════════════════════════
# JUSTIFY HELPER
# ═══════════════════════════

def apply_clean_justify(para):
    """Justify only long paragraphs. Short lines stay LEFT."""
    text  = para.text.strip()
    words = text.split()
    if len(words) < 12 or len(text) < 100 or text.endswith(('?', ':', '!', ';')):
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pPr = para._p.get_or_add_pPr()
    for jc in pPr.findall(qn('w:jc')):
        pPr.remove(jc)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'both')
    pPr.append(jc)


def get_original_alignment(para):
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return None
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        return None
    val = jc.get(qn('w:val'))
    mapping = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right':  WD_ALIGN_PARAGRAPH.RIGHT,
        'both':   WD_ALIGN_PARAGRAPH.JUSTIFY,
        'left':   WD_ALIGN_PARAGRAPH.LEFT,
    }
    return mapping.get(val)


# ═══════════════════════════
# CORE FORMATTING ENGINE
# ═══════════════════════════

def apply_para_formatting(para, etype, font_name, font_size_pt, bold, color, align,
                           space_before_pt, space_after_pt,
                           first_indent=None, left_indent=None,
                           line_spacing=1.15):
    set_para_font(para, font_name)
    clear_pPr_sz(para)
    set_pPr_sz(para, int(font_size_pt * 2))

    # Set paragraph-level bold in pPr>rPr so all runs inherit it
    pPr_b = para._p.get_or_add_pPr()
    rPr_b = pPr_b.find(qn('w:rPr'))
    if rPr_b is None:
        rPr_b = OxmlElement('w:rPr')
        pPr_b.append(rPr_b)
    # Set/clear w:b at paragraph level
    b_ppr = rPr_b.find(qn('w:b'))
    if bold:
        if b_ppr is None:
            b_ppr = OxmlElement('w:b')
            rPr_b.insert(0, b_ppr)
        b_ppr.attrib.pop(qn('w:val'), None)
    else:
        if b_ppr is not None:
            rPr_b.remove(b_ppr)
    # Always remove bCs at paragraph level — overrides run bold=False otherwise
    bcs_ppr = rPr_b.find(qn('w:bCs'))
    if bcs_ppr is not None:
        rPr_b.remove(bcs_ppr)

    para.paragraph_format.space_before = Pt(space_before_pt)
    para.paragraph_format.space_after  = Pt(space_after_pt)
    
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'),            str(int(space_before_pt * 20)))
    spacing.set(qn('w:after'),             str(int(space_after_pt  * 20)))
    spacing.set(qn('w:beforeAutospacing'), '0')
    spacing.set(qn('w:afterAutospacing'),  '0')

    try:
        ls = float(line_spacing)
    except Exception:
        ls = 1.15

    if ls == 1.0:
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif ls == 2.0:
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    else:
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para.paragraph_format.line_spacing = ls

    ind = pPr.find(qn('w:ind'))
    if ind is not None:
        pPr.remove(ind)

    if first_indent is not None or left_indent is not None:
        ind = OxmlElement('w:ind')
        if left_indent is not None:
            ind.set(qn('w:left'), str(int(left_indent * 1440)))
        if first_indent is not None:
            twips = int(first_indent * 1440) if isinstance(first_indent, float) else int(first_indent.inches * 1440)
            ind.set(qn('w:firstLine'), str(twips))
        pPr.append(ind)

    para.alignment = align
    pPr2 = para._p.get_or_add_pPr()
    for jc_el in pPr2.findall(qn('w:jc')):
        pPr2.remove(jc_el)
    jc_new = OxmlElement('w:jc')
    align_val_map = {
        WD_ALIGN_PARAGRAPH.JUSTIFY: 'both',
        WD_ALIGN_PARAGRAPH.CENTER:  'center',
        WD_ALIGN_PARAGRAPH.LEFT:    'left',
        WD_ALIGN_PARAGRAPH.RIGHT:   'right',
    }
    jc_new.set(qn('w:val'), align_val_map.get(align, 'left'))
    pPr2.append(jc_new)

    for run in para.runs:
        if run_has_drawing(run):
            continue
        run.bold = bold
        run.italic = False
        run.underline = False
        r = run._element
        rPr = r.get_or_add_rPr()
        # Force bold XML — remove bCs always to prevent bold bleed from original doc
        b_el = rPr.find(qn('w:b'))
        if bold:
            if b_el is None:
                b_el = OxmlElement('w:b')
                rPr.insert(0, b_el)
            b_el.attrib.pop(qn('w:val'), None)  # clear val=false if present
        else:
            if b_el is not None:
                rPr.remove(b_el)
        # Always remove bCs — it independently forces bold for complex scripts
        bcs_el = rPr.find(qn('w:bCs'))
        if bcs_el is not None:
            rPr.remove(bcs_el)
        for tag in ['w:strike', 'w:dstrike', 'w:highlight', 'w:shd',
                    'w:em', 'w:outline', 'w:shadow', 'w:emboss', 'w:imprint']:
            el = rPr.find(qn(tag))
            if el is not None:
                rPr.remove(el)
        set_font_properly(run, font_name, font_size_pt)
        run.font.color.rgb = color


# ═══════════════════════════
# TABLE HELPERS
# ═══════════════════════════

def format_table_cells(doc, font_name, base_size, line_spacing, black):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if not para.text.strip() and not has_drawing(para):
                        continue
                    set_para_font(para, font_name)
                    clear_pPr_sz(para)
                    set_pPr_sz(para, int(base_size * 2))
                    for run in para.runs:
                        if run_has_drawing(run):
                            continue
                        was_bold   = run.bold
                        set_font_properly(run, font_name, base_size)
                        run.bold   = was_bold
                        run.font.color.rgb = black


def center_all_tables(doc):
    for table in doc.tables:
        tbl   = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        jc = tblPr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            tblPr.append(jc)
        jc.set(qn('w:val'), 'center')


def set_para_text_formatted(para, new_text, font_size_pt, bold, color, font_name=None):
    p = para._p
    for r in p.findall(qn('w:r')):
        p.remove(r)
    run = para.add_run(new_text)
    run.bold = True if bold else False
    run.font.size = Pt(font_size_pt)
    run.font.color.rgb = color
    if font_name:
        set_font_properly(run, font_name, font_size_pt)


def strip_list_numbering(para):
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return
    numPr = pPr.find(qn('w:numPr'))
    if numPr is not None:
        pPr.remove(numPr)


def apply_caps_upper(para, krutidev_mode=False):
    if krutidev_mode:
        return
    for run in para.runs:
        if run.text:
            run.text = run.text.upper()


# ═══════════════════════════
# FULL-DOCUMENT KRUTI DEV CONVERSION
# ═══════════════════════════

ENGLISH_SPECIAL_CHARS = [
    ('\u201c', '"'), ('\u201d', '"'), ('\u2018', "'"), ('\u2019', "'"),
    ('\u2013', '-'), ('\u2014', '--'), ('\u2026', '...'), ('\u00a0', ' '),
    ('\u2022', '*'), ('\u00b7', '*'),
]

def _fix_english_special(text):
    for uni, repl in ENGLISH_SPECIAL_CHARS:
        text = text.replace(uni, repl)
    return text


def convert_run_to_krutidev(run):
    """Convert a run to KrutiDev encoding.
    
    Hindi text → KrutiDev ASCII encoding with KrutiDev font.
    English/non-Hindi segments → kept as-is with Times New Roman font
    (so punctuation and English words render correctly in Word).
    If the run has mixed Hindi+English, it is split into multiple sibling runs.
    """
    text = run.text
    if not text:
        return
    if not has_unicode_hindi(text):
        # Pure English/numeric run — keep text as-is, switch font to Times New Roman
        run.font.name = 'Times New Roman'
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        for attr in list(rFonts.attrib.keys()):
            del rFonts.attrib[attr]
        for a in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
            rFonts.set(qn(f'w:{a}'), 'Times New Roman')
        return

    # Segment into Hindi and non-Hindi parts
    segments = []
    current_hindi = None
    current_chunk = []
    for ch in text:
        ch_is_hindi = '\u0900' <= ch <= '\u097F'
        if current_hindi is None:
            current_hindi = ch_is_hindi
        if ch_is_hindi == current_hindi:
            current_chunk.append(ch)
        else:
            segments.append((current_hindi, ''.join(current_chunk)))
            current_hindi = ch_is_hindi
            current_chunk = [ch]
    if current_chunk:
        segments.append((current_hindi, ''.join(current_chunk)))

    if len(segments) == 1 and segments[0][0]:
        # Pure Hindi run — convert and keep KrutiDev font
        run.text = unicode_to_krutidev(segments[0][1])
        return

    # Mixed run — split into sibling runs
    # First run reuses the existing run element
    para = run._r.getparent()  # w:p element
    r_elem = run._r
    r_idx = list(para).index(r_elem)

    # Build list of (text, is_hindi) for new runs
    new_runs = []
    for is_h, seg in segments:
        if is_h:
            new_runs.append((unicode_to_krutidev(seg), True))
        else:
            new_runs.append((seg, False))

    # Modify existing run to be the first segment
    first_text, first_hindi = new_runs[0]
    run.text = first_text
    # Explicitly set font for the first run segment
    if first_hindi:
        set_font_properly(run, run.font.name or 'Kruti Dev 010')
    else:
        # English segment — switch font to Times New Roman
        run.font.name = 'Times New Roman'
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        for attr in list(rFonts.attrib.keys()):
            del rFonts.attrib[attr]
        for a in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
            rFonts.set(qn(f'w:{a}'), 'Times New Roman')
        # Remove any language/bidi overrides that might force KrutiDev
        for tag in ['w:lang', 'w:rtl', 'w:cs']:
            el = rPr.find(qn(tag))
            if el is not None:
                rPr.remove(el)

    # Insert remaining runs after the first
    import copy
    for i, (seg_text, seg_hindi) in enumerate(new_runs[1:], start=1):
        new_r = copy.deepcopy(r_elem)
        # Set text
        t_els = new_r.findall(qn('w:t'))
        if t_els:
            t_els[0].text = seg_text
            if seg_text and (seg_text[0] == ' ' or seg_text[-1] == ' '):
                t_els[0].set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        else:
            from docx.oxml import OxmlElement as _OE
            t_el = _OE('w:t')
            t_el.text = seg_text
            if seg_text and (seg_text[0] == ' ' or seg_text[-1] == ' '):
                t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            new_r.append(t_el)

        # Re-wrap in python-docx Run object to use set_font_properly
        from docx.text.run import Run
        new_run_obj = Run(new_r, run._parent)

        if seg_hindi:
            set_font_properly(new_run_obj, run.font.name or 'Kruti Dev 010')
        else:
            # English segment
            new_run_obj.font.name = 'Times New Roman'
            r_new = new_run_obj._element
            rPr_new = r_new.get_or_add_rPr()
            rFonts_new = rPr_new.get_or_add_rFonts()
            for attr in list(rFonts_new.attrib.keys()):
                del rFonts_new.attrib[attr]
            for a in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
                rFonts_new.set(qn(f'w:{a}'), 'Times New Roman')

        para.insert(r_idx + i, new_r)


def convert_doc_runs(doc, font_name):
    if not is_krutidev(font_name):
        return

    def process_para(para):
        if has_drawing(para):
            return
        set_para_font(para, font_name)
        # Use index-based loop because convert_run_to_krutidev inserts sibling runs
        j = 0
        while j < len(para.runs):
            run = para.runs[j]
            if not run_has_drawing(run):
                # We must NOT call set_font_properly(run, font_name) here because 
                # convert_run_to_krutidev already handles font setting for each segment.
                # If we do, we'll overwrite English TNR segments with KrutiDev.
                orig_run_count = len(para.runs)
                convert_run_to_krutidev(run)
                # Skip any newly inserted runs so we don't process them twice
                new_run_count = len(para.runs)
                j += (new_run_count - orig_run_count)
            j += 1

    # 1. Main paragraphs
    for para in doc.paragraphs:
        process_para(para)

    # 2. Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_para(para)

    # 3. Headers & Footers (Fix for missing title/header conversion)
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for para in header.paragraphs:
                    process_para(para)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    process_para(para)

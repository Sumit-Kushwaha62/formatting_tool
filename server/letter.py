import re
import copy
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from utils import (
    has_drawing, is_all_bold, is_bullet_para,
    apply_para_formatting, apply_bold_before_colon, apply_clean_justify,
    set_font_properly, set_para_font, is_krutidev
)


# ═══════════════════════════
# LETTER HEADER
# ═══════════════════════════

def insert_letter_header(doc, opts, font_name):
    black = RGBColor(0, 0, 0)
    gray  = RGBColor(0, 0, 0)   # forced black
    dark  = RGBColor(0, 0, 0)   # forced black

    org_name     = opts.get('org_name',     '').strip()
    ref_no       = opts.get('ref_no',       '').strip()
    date         = opts.get('date',         '').strip()
    subject      = opts.get('subject',      '').strip()
    letter_title = opts.get('letter_title', '').strip()

    if not org_name and not subject and not letter_title:
        return

    insert_paras = []

    def make_para(text, align, size, bold=False, italic=False,
                  space_before=0, space_after=8, color=None):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        if text:
            r = p.add_run(text)
            r.bold   = bold
            r.italic = italic
            set_font_properly(r, font_name)
            r.font.size = Pt(size)
            r.font.color.rgb = color or black
        return p

    # 1. Org name (top center, 16pt)
    if org_name:
        insert_paras.append(make_para(org_name, WD_ALIGN_PARAGRAPH.CENTER,
                                       16, bold=True, color=dark, space_after=4))

    # 2. Top HR line
    hr = doc.add_paragraph()
    hr.paragraph_format.space_before = Pt(4)
    hr.paragraph_format.space_after  = Pt(6)
    pPr   = hr._p.get_or_add_pPr()
    pBdr  = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2222AA')
    pBdr.append(bottom)
    pPr.append(pBdr)
    insert_paras.append(hr)

    # 3. Ref. No. (top-left) + Date (top-right) on same line
    if ref_no or date:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        # Ref. No. always top-left; Date always top-right via right-aligned tab stop
        pPr = p._p.get_or_add_pPr()
        tabs_elem = OxmlElement('w:tabs')
        tab_elem  = OxmlElement('w:tab')
        tab_elem.set(qn('w:val'), 'right')
        tab_elem.set(qn('w:pos'), '9026')  # right edge of A4 content (1" margins)
        tabs_elem.append(tab_elem)
        pPr.append(tabs_elem)
        if ref_no:
            r1 = p.add_run(f'Ref. No.: {ref_no}')
            r1.bold = True
            set_font_properly(r1, font_name)
            r1.font.size = Pt(12)
            r1.font.color.rgb = black
        if date:
            # Tab pushes date to the right edge (top-right)
            tab_r = p.add_run('\t')
            set_font_properly(tab_r, font_name)
            tab_r.font.size = Pt(12)
            r2 = p.add_run(f'Date: {date}')
            r2.bold = True
            set_font_properly(r2, font_name)
            r2.font.size = Pt(12)
            r2.font.color.rgb = black
        insert_paras.append(p)

    # 4. Letter title — center below ref/date (14pt bold, only element at 14pt)
    if letter_title:
        insert_paras.append(make_para(letter_title, WD_ALIGN_PARAGRAPH.CENTER,
                                       14, bold=True, color=dark,
                                       space_before=12, space_after=8))

    # 5. Bottom HR line
    hr2  = doc.add_paragraph()
    hr2.paragraph_format.space_before = Pt(4)
    hr2.paragraph_format.space_after  = Pt(10)
    pPr2  = hr2._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    b2 = OxmlElement('w:bottom')
    b2.set(qn('w:val'),   'single')
    b2.set(qn('w:sz'),    '4')
    b2.set(qn('w:space'), '1')
    b2.set(qn('w:color'), 'AAAACC')
    pBdr2.append(b2)
    pPr2.append(pBdr2)
    insert_paras.append(hr2)

    # 6. Subject line
    if subject:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(12)
        sp.paragraph_format.space_after  = Pt(12)
        sp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r_lbl = sp.add_run('Subject: ')
        r_lbl.bold = True
        set_font_properly(r_lbl, font_name)
        r_lbl.font.size = Pt(12)
        r_lbl.font.color.rgb = black
        r_sub = sp.add_run(subject)
        r_sub.bold      = True
        r_sub.underline = True
        set_font_properly(r_sub, font_name)
        r_sub.font.size = Pt(12)
        r_sub.font.color.rgb = dark
        insert_paras.append(sp)

    body = doc.element.body
    for p in reversed(insert_paras):
        body.remove(p._element)
        body.insert(0, p._element)


# ═══════════════════════════
# LETTER STRUCTURE DETECTION
# ═══════════════════════════

# ═══════════════════════════════════════════════════════
# OPTIMAL 2-PASS LETTER STRUCTURE DETECTION ALGORITHM
#
# Pass 1 (scan_letter_header_block):
#   Scans top N paragraphs to locate and extract:
#     - ref_line_index  : paragraph index containing Ref. No.
#     - date_line_index : paragraph index containing Date
#     - title_index     : paragraph index of letter title (ALL-CAPS short line)
#   These three are ALWAYS treated independently regardless of where they appear.
#
# Pass 2 (detect_letter_structure):
#   Each paragraph is classified using context from Pass 1.
#   ref/date/title indices are excluded from generic detection so they
#   never get misclassified as label/body/salutation.
# ═══════════════════════════════════════════════════════

# ── Regex patterns (compiled once) ──────────────────────
_RE_REF  = re.compile(r'^ref\.?(?:\s*no\.?)?\s*:', re.IGNORECASE)
_RE_DATE = re.compile(r'^date\s*:', re.IGNORECASE)
# Ref and Date on same line: "Ref. No.: X   Date: Y"
_RE_REF_INLINE  = re.compile(r'ref\.?(?:\s*no\.?)?\s*:\s*([^\t]+?)(?:\s{2,}|\t)', re.IGNORECASE)
_RE_DATE_INLINE = re.compile(r'date\s*:\s*(\S+)', re.IGNORECASE)
# Title: ALL-CAPS words only, 1–6 words, optionally bold
_RE_TITLE = re.compile(r'^[A-Z][A-Z\s]{2,49}$')


def is_ref_line(text):
    """True if paragraph starts with a Ref. No. pattern."""
    return bool(_RE_REF.match(text))


def is_date_line(text):
    """True if paragraph starts with a Date: pattern."""
    return bool(_RE_DATE.match(text))


def is_title_line(text):
    """
    True if text looks like a standalone letter title:
      - ALL CAPS
      - 1–6 words
      - Only alpha + spaces (no digits, punctuation)
      - Not a known salutation or closing keyword
    """
    words = text.split()
    if not (1 <= len(words) <= 6):
        return False
    if not _RE_TITLE.match(text):
        return False
    # Exclude salutations/closings that happen to be all-caps
    _skip = {'DEAR', 'TO', 'RESPECTED', 'YOURS', 'SINCERELY', 'REGARDS'}
    if text.strip().upper() in _skip or words[0].upper() in _skip:
        return False
    return True


def scan_letter_header_block(doc, scan_limit=20):
    """
    Pass 1 — scans first `scan_limit` paragraphs and returns a dict:
      {
        'ref_idx':   int or None,   # para index with Ref. No.
        'date_idx':  int or None,   # para index with Date (may equal ref_idx if same line)
        'title_idx': int or None,   # para index of letter title
        'ref_val':   str,           # extracted ref number value
        'date_val':  str,           # extracted date value
      }
    Handles three layouts:
      A) Ref and Date on separate lines
      B) Ref and Date on the same line (tab/space separated)
      C) Only one of them present
    Title is detected independently — it can appear before or after ref/date.
    """
    result = {
        'ref_idx':   None,
        'date_idx':  None,
        'title_idx': None,
        'ref_val':   '',
        'date_val':  '',
    }

    paras = doc.paragraphs[:scan_limit]

    for i, para in enumerate(paras):
        text = para.text.strip()
        if not text:
            continue

        # ── Ref. No. detection (dedicated line) ──
        if result['ref_idx'] is None and is_ref_line(text):
            result['ref_idx'] = i
            # Try to also find date on same line
            date_m = _RE_DATE_INLINE.search(text)
            if date_m:
                result['date_idx'] = i   # same paragraph
                result['date_val'] = date_m.group(1).strip()
            # Extract ref value
            ref_m = re.search(r'ref\.?(?:\s*no\.?)?\s*:\s*(.+?)(?:\s{2,}|\t|date\s*:|$)',
                               text, re.IGNORECASE)
            if ref_m:
                result['ref_val'] = ref_m.group(1).strip()
            continue

        # ── Date detection (dedicated line) ──
        if result['date_idx'] is None and is_date_line(text):
            result['date_idx'] = i
            date_m = re.search(r'date\s*:\s*(.+)', text, re.IGNORECASE)
            if date_m:
                result['date_val'] = date_m.group(1).strip()
            continue

        # ── Title detection (ALL-CAPS standalone line) ──
        if result['title_idx'] is None and is_title_line(text):
            result['title_idx'] = i

    return result


def has_existing_letter_header(doc):
    """Quick check — returns True if doc already has a formatted Ref. No. line."""
    for para in doc.paragraphs[:10]:
        if is_ref_line(para.text.strip()):
            return True
    return False


def is_ref_date_line(para):
    """
    Legacy helper — kept for backward compatibility.
    Returns True if paragraph is a Ref. No. line (with or without inline Date).
    """
    return is_ref_line(para.text.strip())


def detect_letter_structure(para, index, header_block=None):
    """
    Pass 2 — classify a single paragraph.

    header_block (from scan_letter_header_block) lets us skip ref/date/title
    paragraphs so they are never misclassified.

    Classification priority (highest → lowest):
      1. empty       — blank paragraph
      2. bullet      — list item
      3. ref_line    — Ref. No. paragraph (handled separately)
      4. date_line   — Date paragraph (handled separately, if on own line)
      5. title       — ALL-CAPS letter title
      6. salutation  — starts with Dear/To/Respected/Subject
      7. closing     — starts with Yours/Regards/etc. (≤5 words)
      8. signature   — bold short line late in document (index > 5)
      9. label       — bold side-heading (≤12 words)
     10. body        — everything else
    """
    text  = para.text.strip()
    words = text.split()
    wc    = len(words)

    if wc == 0:
        return 'empty'
    if is_bullet_para(para):
        return 'bullet'

    # If header_block provided, skip paragraphs already claimed as ref/date/title
    if header_block:
        if index == header_block.get('ref_idx'):
            return 'ref_line'
        if (index == header_block.get('date_idx')
                and header_block.get('date_idx') != header_block.get('ref_idx')):
            return 'date_line'
        if index == header_block.get('title_idx'):
            return 'title'

    is_bold = is_all_bold(para)

    # Salutation
    if re.match(r'^(dear|to|respected|sub|subject)\b', text, re.IGNORECASE):
        return 'salutation'

    # Closing
    _closing = ['yours', 'sincerely', 'faithfully', 'regards', 'thanking',
                'with regards', 'best regards', 'warm regards']
    if any(text.lower().startswith(w) for w in _closing) and wc <= 5:
        return 'closing'

    # Title (fallback when no header_block passed)
    if is_title_line(text):
        return 'title'

    # Signature — bold, short, appears later in document
    if is_bold and wc <= 8 and index > 5:
        return 'signature'

    # Side heading / label — bold, ≤12 words
    if is_bold and wc <= 12:
        return 'label'

    # Numbered side heading pattern: "1. Heading Text" or "Assigned Responsibilities" (all-caps-ish short line)
    # Even if NOT bold in source — these should be treated as labels and made bold
    if re.match(r'^\d+\.\s+\w', text) and wc <= 12:
        return 'label'

    # Short non-bold heading-like line (e.g. "Assigned Responsibilities")
    # Detect: title-case or all-caps, short (≤6 words), no sentence-ending punctuation
    if wc <= 6 and not text.endswith('.') and not text.endswith(',') and re.match(r'^[A-Z]', text):
        # All words start with capital or are short connectors — looks like a heading
        words_check = text.split()
        if all(w[0].isupper() or len(w) <= 3 for w in words_check):
            return 'label'

    return 'body'


def preserve_para_indent(para):
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return None
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        return None
    return copy.deepcopy(ind)


def restore_para_indent(para, saved_ind):
    if saved_ind is None:
        return
    pPr = para._p.get_or_add_pPr()
    existing = pPr.find(qn('w:ind'))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(saved_ind)


# ═══════════════════════════
# LETTER BODY FORMATTING
# ═══════════════════════════

def _reorder_title_after_ref(doc, header_block):
    """
    If title para appears BEFORE ref/date para in the document body,
    physically move it to just after the ref/date para in XML.
    Ensures correct order: ref/date line → title → rest of letter body.
    """
    title_idx = header_block.get('title_idx')
    ref_idx   = header_block.get('ref_idx')
    date_idx  = header_block.get('date_idx')

    if title_idx is None or ref_idx is None:
        return  # nothing to reorder

    # Only reorder if title physically precedes ref/date in doc
    if title_idx >= ref_idx:
        return

    # anchor = last of ref_idx / date_idx (whichever appears later)
    anchor_idx = ref_idx
    if date_idx is not None and date_idx != ref_idx:
        anchor_idx = max(anchor_idx, date_idx)

    paras = doc.paragraphs
    if title_idx >= len(paras) or anchor_idx >= len(paras):
        return

    title_elem  = paras[title_idx]._p
    anchor_elem = paras[anchor_idx]._p
    body        = doc.element.body

    # Remove title from current position, insert right after anchor
    body.remove(title_elem)
    anchor_elem.addnext(title_elem)


def format_letter_body(doc, opts, font_name):
    black         = RGBColor(0, 0, 0)
    dark          = RGBColor(0, 0, 0)   # forced black
    krutidev_mode = is_krutidev(font_name)

    # ── Pass 1: locate ref/date/title positions ──────────────
    header_block = scan_letter_header_block(doc)

    # ── Reorder: ensure title appears AFTER ref/date in body ─
    _reorder_title_after_ref(doc, header_block)
    # Re-scan after reorder so Pass 2 indices are fresh/correct
    header_block = scan_letter_header_block(doc)

    # ── Remove leading empty paras before ref/date ───────────
    ref_idx = header_block.get('ref_idx')
    if ref_idx:
        paras_snap = list(doc.paragraphs)
        for j in range(ref_idx):
            p = paras_snap[j]
            if not p.text.strip() and not has_drawing(p):
                p._element.getparent().remove(p._element)
        # Re-scan after cleanup
        header_block = scan_letter_header_block(doc)

    # ── Pass 2: format each paragraph ───────────────────────
    for i, para in enumerate(doc.paragraphs):
        if has_drawing(para):
            continue
        text = para.text.strip()
        if not text:
            continue

        etype = detect_letter_structure(para, i, header_block=header_block)

        if etype == 'empty':
            continue

        # ── Ref. No. line: always top-left; Date always top-right ──
        if etype == 'ref_line':
            ref_val  = header_block.get('ref_val',  '') or ''
            date_val = header_block.get('date_val', '') or ''

            # If date is on a different paragraph, still use its value (we'll suppress the date_line para)
            # Only try to re-extract inline if header_block has no date_val at all
            if header_block.get('date_idx') != i and not date_val:
                dm = _RE_DATE_INLINE.search(para.text)
                if dm:
                    date_val = dm.group(1).strip()

            p_elem = para._p
            for child in list(p_elem):
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag in ('r', 'bookmarkStart', 'bookmarkEnd', 'proofErr'):
                    p_elem.remove(child)

            pPr = p_elem.get_or_add_pPr()
            for old in pPr.findall(qn('w:tabs')):
                pPr.remove(old)
            tabs_elem = OxmlElement('w:tabs')
            tab_elem  = OxmlElement('w:tab')
            tab_elem.set(qn('w:val'), 'right')
            tab_elem.set(qn('w:pos'), '9026')
            tabs_elem.append(tab_elem)
            pPr.append(tabs_elem)
            sp = pPr.find(qn('w:spacing'))
            if sp is not None:
                pPr.remove(sp)
            sp_new = OxmlElement('w:spacing')
            sp_new.set(qn('w:before'), '0')
            sp_new.set(qn('w:after'),  '80')
            sp_new.set(qn('w:line'),   '240')
            sp_new.set(qn('w:lineRule'), 'auto')
            pPr.append(sp_new)

            if ref_val:
                r_ref = para.add_run(f'Ref. No.: {ref_val}')
                r_ref.bold = True
                if not krutidev_mode:
                    set_font_properly(r_ref, font_name)
                    r_ref.font.size = Pt(12)
                    r_ref.font.color.rgb = RGBColor(0, 0, 0)
            if date_val:
                r_tab = para.add_run('\t')
                if not krutidev_mode:
                    set_font_properly(r_tab, font_name)
                    r_tab.font.size = Pt(12)
                r_date = para.add_run(f'Date: {date_val}')
                r_date.bold = True
                if not krutidev_mode:
                    set_font_properly(r_date, font_name)
                    r_date.font.size = Pt(12)
                    r_date.font.color.rgb = RGBColor(0, 0, 0)
            continue

        # ── Date-only line (Date on its own paragraph, separate from Ref) ──
        if etype == 'date_line':
            # If ref_idx also exists, date shown in ref_line para. Suppress this para.
            if header_block and header_block.get('ref_idx') is not None:
                p_elem = para._p
                for child in list(p_elem):
                    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag in ('r', 'bookmarkStart', 'bookmarkEnd', 'proofErr'):
                        p_elem.remove(child)
                pPr = p_elem.get_or_add_pPr()
                sp = pPr.find(qn('w:spacing'))
                if sp is not None:
                    pPr.remove(sp)
                sp_new = OxmlElement('w:spacing')
                sp_new.set(qn('w:before'), '0')
                sp_new.set(qn('w:after'),  '0')
                sp_new.set(qn('w:line'),   '240')
                sp_new.set(qn('w:lineRule'), 'auto')
                pPr.append(sp_new)
                continue

            date_val = header_block.get('date_val', '').strip()
            p_elem = para._p
            for child in list(p_elem):
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag in ('r', 'bookmarkStart', 'bookmarkEnd', 'proofErr'):
                    p_elem.remove(child)

            pPr = p_elem.get_or_add_pPr()
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                pPr.append(jc)
            jc.set(qn('w:val'), 'right')

            sp = pPr.find(qn('w:spacing'))
            if sp is not None:
                pPr.remove(sp)
            sp_new = OxmlElement('w:spacing')
            sp_new.set(qn('w:before'), '80')
            sp_new.set(qn('w:after'),  '80')
            sp_new.set(qn('w:line'),   '240')
            sp_new.set(qn('w:lineRule'), 'auto')
            pPr.append(sp_new)

            r_date = para.add_run(f'Date: {date_val}')
            r_date.bold = True
            if not krutidev_mode:
                set_font_properly(r_date, font_name)
                r_date.font.size = Pt(12)
                r_date.font.color.rgb = RGBColor(0, 0, 0)
            continue

        # ── Letter title: 14pt bold center ──
        if etype == 'title':
            apply_para_formatting(para, etype, font_name,
                font_size_pt=14, bold=True, color=black,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                space_before_pt=12, space_after_pt=12)
            continue

        # ── Remaining types ──────────────────────────────────
        if etype == 'salutation':
            apply_para_formatting(para, etype, font_name,
                font_size_pt=12, bold=True, color=black,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before_pt=8, space_after_pt=8)

        elif etype == 'closing':
            apply_para_formatting(para, etype, font_name,
                font_size_pt=12, bold=False, color=black,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before_pt=16, space_after_pt=4)

        elif etype == 'signature':
            apply_para_formatting(para, etype, font_name,
                font_size_pt=12, bold=True, color=black,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before_pt=2, space_after_pt=2)

        elif etype == 'label':
            # Side headings: 12pt bold black — always bold
            apply_para_formatting(para, etype, font_name,
                font_size_pt=12, bold=True, color=black,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before_pt=8, space_after_pt=2)
            if not krutidev_mode and ': ' in para.text:
                apply_bold_before_colon(para, font_name, krutidev_mode)

        elif etype == 'bullet':
            is_bold_para = is_all_bold(para)
            apply_para_formatting(para, etype, font_name,
                font_size_pt=12, bold=is_bold_para, color=black,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before_pt=0, space_after_pt=4)
            if not krutidev_mode and ': ' in para.text and not is_bold_para:
                apply_bold_before_colon(para, font_name, krutidev_mode)

        else:  # body — 12pt, 1.15 line spacing
            if krutidev_mode:
                apply_para_formatting(para, etype, font_name,
                    font_size_pt=12, bold=False, color=black,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before_pt=0, space_after_pt=4)
            else:
                apply_clean_justify(para)
                apply_para_formatting(para, etype, font_name,
                    font_size_pt=12, bold=False, color=black,
                    align=para.alignment,
                    space_before_pt=0, space_after_pt=4)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = 1.15